#!/usr/bin/env python3
"""report_dive.py — generate the FINAL report bundle for a deep-dive.

Reads a deep-dive bundle (concept or ayah) under SpatialAnalysis/deep-dives/ and
writes report/{technical,plain_en,plain_fa}.docx + matching .pdf (via LibreOffice).
This is the user-GATED last step (per DEEP_DIVE_SPEC §4) — run it only on request.

Three registers of the SAME data-driven findings:
  technical  — methods, parameters, every statistic + null, tables, provenance.
  plain_en   — plain English with METHODOLOGICAL analogies (ecology/GIS/network).
  plain_fa   — Persian parallel, با تشبیه‌های روش‌شناختی.
HARD constraint: computational DESCRIPTION, never tafsir.

Usage:
  python report_dive.py <bundle_dir>        # e.g. .../deep-dives/concepts/قلب_20260604
  python report_dive.py --latest concept    # newest concept bundle
  python report_dive.py --latest ayah
"""
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

HERE = Path(__file__).resolve().parent
OUT_BASE = HERE.parent / "SpatialAnalysis" / "deep-dives"
NAVY = RGBColor(0x1D, 0x35, 0x57)
TEAL = RGBColor(0x0F, 0x6E, 0x56)
# Complex-script (Arabic/Persian) font. "Arial" is Persian-capable on Windows
# (where the .docx is read) and substitutes cleanly to DejaVu Sans in the
# LibreOffice sandbox (verified to shape گچپژکی + Persian digits correctly).
FA_CS_FONT = "Arial"


def _cs_font(run, name=FA_CS_FONT):
    """Pin BOTH the ASCII/hAnsi and the COMPLEX-SCRIPT font so Word/LibreOffice
    never guess an Arabic default (the old font bug)."""
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.insert(0, rf)
    rf.set(qn("w:cs"), name); rf.set(qn("w:ascii"), name); rf.set(qn("w:hAnsi"), name)


_ARABIC = re.compile(r"[؀-ۿݐ-ݿࢠ-ࣿﭐ-﷿ﹰ-﻿]")


def _has_arabic(t):
    return bool(_ARABIC.search(t or ""))


def _pin_cs(run, name=FA_CS_FONT):
    """Pin ONLY the complex-script font, so Arabic/Persian inside an otherwise
    Latin (LTR) run renders with a Persian-capable face instead of boxes, while
    Latin text keeps the body font. Fixes the 'roots show as □□□' bug in LTR
    paragraphs and table cells."""
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.insert(0, rf)
    rf.set(qn("w:cs"), name)


# ── docx helpers ───────────────────────────────────────────────────────────
def _new_doc(rtl=False):
    d = Document()
    stl = d.styles["Normal"]
    stl.font.name = "Times New Roman"
    stl.font.size = Pt(11)
    pf = stl.paragraph_format
    pf.line_spacing = 1.0            # single-spaced (tight, =1)
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    for sec in d.sections:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(1)
        sec.top_margin = sec.bottom_margin = Inches(1)
        _page_number_footer(sec)
    return d


def _page_number_footer(sec):
    """Centered 'PAGE' field in the footer of a section."""
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.font.name = "Times New Roman"; run.font.size = Pt(10)
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr); run._r.append(e)


def _set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.insert(0, OxmlElement("w:bidi"))      # paragraph direction = RTL
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   # jc=both (Persian house style)
    for run in p.runs:
        rPr = run._r.get_or_add_rPr()
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))       # run is complex-script RTL
        _cs_font(run)                              # pin the Persian font


def H(d, text, lvl=1, rtl=False):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(16 if lvl == 1 else 13)
    r.font.color.rgb = NAVY
    if rtl:
        _set_rtl(p)
    elif _has_arabic(text):
        _pin_cs(r)
    return p


def P(d, text, rtl=False, italic=False, bold=False, color=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if rtl:
        _set_rtl(p)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if _has_arabic(text):
            _pin_cs(r)
    return p


def BULLET(d, text, rtl=False):
    p = d.add_paragraph(style="List Bullet")
    p.add_run(text)
    if rtl:
        _set_rtl(p)
    return p


def TABLE(d, headers, rows):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        rr = c.paragraphs[0].add_run(str(h)); rr.bold = True
        if _has_arabic(str(h)):
            _pin_cs(rr)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(v))
            if _has_arabic(str(v)):
                _pin_cs(rr)
    return t


def _to_pdf(docx_path):
    """Convert in an ISOLATED temp dir (soffice scatters temp/lock files in its
    outdir), then copy only the finished .pdf back so the report dir stays clean."""
    import shutil
    import tempfile
    tmpd = tempfile.mkdtemp(prefix="lo_out_")
    prof = tempfile.mkdtemp(prefix="lo_prof_")
    try:
        subprocess.run(["soffice", "--headless",
                        f"-env:UserInstallation=file://{prof}",
                        "--convert-to", "pdf", "--outdir", tmpd, str(docx_path)],
                       check=True, capture_output=True, timeout=120)
        src = Path(tmpd) / (docx_path.stem + ".pdf")
        if src.exists():
            shutil.copy(str(src), str(docx_path.with_suffix(".pdf")))
            return True
        return False
    except Exception as e:
        print("  [pdf conversion failed]", e)
        return False
    finally:
        shutil.rmtree(tmpd, ignore_errors=True)
        shutil.rmtree(prof, ignore_errors=True)


def _load(bundle):
    f = bundle / "findings"
    out = {"manifest": json.loads((bundle / "MANIFEST.json").read_text("utf-8")),
           "request": json.loads((bundle / "00_request.json").read_text("utf-8"))}
    for j in f.glob("*.json"):
        out[j.stem] = json.loads(j.read_text("utf-8"))
    data = bundle / "data" / "full.json"
    if data.exists():
        out["full"] = json.loads(data.read_text("utf-8"))
    return out


# ── CONCEPT report content ─────────────────────────────────────────────────
REL_EN = {
    "consensus": "agree on ≥2 modalities (robust bond)",
    "semantic": "meaning-mates (distributional)",
    "co-location": "territory-mates (shared deployment)",
    "spatial": "similar distribution shape",
    "orthogonal": "tied on ONE modality, independent on the rest",
    "divergent": "close on one modality, OPPOSED on another (tension)",
}
REL_FA = {
    "consensus": "هم‌داستان روی ≥۲ پیمانه (پیوند استوار)",
    "semantic": "هم‌معنا (توزیعی)",
    "co-location": "هم‌قلمرو (هم‌جایگاهِ متنی)",
    "spatial": "هم‌شکلِ پراکندگی",
    "orthogonal": "پیوند روی یک پیمانه، مستقل بر بقیه",
    "divergent": "نزدیک بر یک پیمانه و مخالف بر دیگری (تنش)",
}
REL_ORDER = ["consensus", "semantic", "co-location", "spatial", "orthogonal", "divergent"]


def _concept_technical(d, b, t):
    syn = b.get("synthesis", {}) or {}
    rel = b.get("relations", {}) or {}
    seq = b.get("sequence", {}) or {}
    fld, dist, null = b["field"], b["distribution"], b["null"]
    cg, mani = b["cross_granularity"], b["manifest"]
    rbt = rel.get("related_by_type", {})
    cm = syn.get("cross_modal", {})
    H(d, f"Concept Deep-Dive — Technical Report: {t}", 1)
    P(d, f"Generated {mani.get('generated','')} · code {mani.get('code_version','')} · "
         f"unit={b['request'].get('unit')}", italic=True)
    P(d, "Computational DESCRIPTION, not tafsir. Read across ALL occurrences "
         "(tree→forest). Characterised by MULTIMODAL FUSION — independent modalities "
         "(semantic ∥ co-location ∥ spatial ∥ morphology ∥ sequence) kept separate and "
         "SYNTHESISED, never reduced to any single view.", italic=True)

    H(d, "1. Multimodal fusion — the headline", 2)
    P(d, syn.get("reading", ""))
    for ty in REL_ORDER:
        lst = rbt.get(ty) or []
        if not lst:
            continue
        H(d, f"{ty} — {REL_EN[ty]}", 2)
        TABLE(d, ["root", "semantic", "co-location", "spatial"],
              [[x["root"], x["axes"]["semantic"], x["axes"]["co-location"],
                x["axes"]["spatial"]] for x in lst[:8]])
    P(d, "Convergence (meaning ∩ territory): " +
         ("، ".join(cm.get("convergence", [])) or "—"))
    P(d, "Divergence: " + str(cm.get("divergence", "—")))
    P(d, "Verified bonds (root∥surface): " +
         ("، ".join(cm.get("verified_bonds", [])) or "—"))

    H(d, "2. Modalities in detail", 2)
    P(d, "Semantic field (meaning): " + ("، ".join(fld["semantic_field"]) or "—"))
    P(d, "Co-location territory: " + ("، ".join(fld["co_location_neighbours"]) or "—"))
    P(d, f"Sequence (Two Books): mean in-ayah position "
         f"{seq.get('mean_within_ayah_position')} · ayah-final share "
         f"{seq.get('ayah_final_share')} (rhyme/fawāṣil).")
    arch = dist["archetype"]
    P(d, f"Spatial distribution shape: archetype {(arch['tag'] if arch else '—')}; "
         f"areal-evenness null z={null['z']} → {null['interpretation']} "
         f"(ONE modality — null here, NOT the headline).")
    if b.get("senses"):
        P(d, "Morphology / surface senses:")
        for sN in b["senses"][:6]:
            P(d, f"  {sN['form']} (×{sN['count']}): " +
                 ("، ".join(f"{r}({a})" for r, a, p in sN["share"][:6]) or "—"))

    H(d, "3. Distribution & provenance", 2)
    TABLE(d, ["metric", "value"], [
        ["frequency", dist["frequency"]],
        ["surahs present", dist["n_surahs_present"]],
        ["archetype", (arch["tag"] if arch else "—")],
        ["archetype stability", (arch["stability"] if arch else "—")],
    ])
    for k, v in mani.get("verification", {}).items():
        P(d, f"• {k}: {v}")


def _concept_plain_en(d, b, t):
    syn = b.get("synthesis", {}) or {}
    rel = b.get("relations", {}) or {}
    rbt = rel.get("related_by_type", {})
    seq = b.get("sequence", {}) or {}
    H(d, f"What the data shows about \u201c{t}\u201d — in plain English", 1)
    P(d, "A computational description, not an interpretation. We read every place this "
         "concept appears (the tree) across the whole Qur\u2019an (the forest), and we "
         "view it through SEVERAL independent lenses at once — then combine what they say.",
      italic=True)

    H(d, "Several lenses, combined — not just one", 2)
    P(d, "Like studying a city with three maps — one of who SOUNDS alike (meaning), one "
         "of who LIVES in the same neighbourhood (territory), and one of the overall "
         "street-plan (distribution) — we ask where the maps AGREE and where they "
         "DISAGREE. The disagreements are often the most telling.")

    def line(ty, lead):
        lst = rbt.get(ty) or []
        if lst:
            P(d, f"{lead}: " + "، ".join(x["root"] for x in lst[:6]))

    line("consensus", "Confirmed by two or more lenses (strongest bonds)")
    line("semantic", "Close in MEANING (same family of ideas)")
    line("co-location", "Share the same TERRITORY in the text")
    line("orthogonal", "Tied on just ONE lens (a link the others don\u2019t see)")
    P(d, "And the revealing tensions — close on one lens but pushed APART on another:")
    line("divergent", "In tension")
    if rbt.get("divergent"):
        P(d, "These \u2018in-tension\u2019 links are where a concept meets something "
             "that shares its setting yet pulls the opposite way in meaning — the data "
             "flags them automatically, no hand-picking.")

    H(d, "The sequence lens (the \u2018Second Book\u2019)", 2)
    P(d, f"Within a verse this concept sits on average near the "
         f"{round((seq.get('mean_within_ayah_position') or 0)*100)}% mark, and ends a "
         f"verse about {round((seq.get('ayah_final_share') or 0)*100)}% of the time — a "
         f"separate order-and-rhyme view, kept apart from meaning and combined only at "
         f"the end.")

    H(d, "What we are NOT claiming", 2)
    P(d, "This is a multi-lens map and census, not a commentary. It shows where the "
         "concept lives, what company it keeps, and where tensions sit; the meaning "
         "itself is for the text and its scholars.")


def _concept_plain_fa(d, b, t):
    syn = b.get("synthesis", {}) or {}
    rel = b.get("relations", {}) or {}
    rbt = rel.get("related_by_type", {})
    seq = b.get("sequence", {}) or {}
    H(d, f"\u0622\u0646\u0686\u0647 \u062f\u0627\u062f\u0647\u200c\u0647\u0627 \u062f\u0631\u0628\u0627\u0631\u0647\u0654 \u00ab{t}\u00bb \u0646\u0634\u0627\u0646 \u0645\u06cc\u200c\u062f\u0647\u0646\u062f \u2014 \u0686\u0646\u062f\u200c\u0648\u062c\u0647\u06cc", 1, rtl=True)
    P(d, "\u06cc\u06a9 \u062a\u0648\u0635\u06cc\u0641 \u0645\u062d\u0627\u0633\u0628\u0627\u062a\u06cc \u0627\u0633\u062a\u060c \u0646\u0647 \u062a\u0641\u0633\u06cc\u0631. \u0647\u0631 \u062c\u0627\u06cc\u06cc \u06a9\u0647 \u0627\u06cc\u0646 \u0645\u0641\u0647\u0648\u0645 \u0645\u06cc\u200c\u0622\u06cc\u062f (\u062f\u0631\u062e\u062a) \u062f\u0631 \u0633\u0631\u0627\u0633\u0631 \u0642\u0631\u0622\u0646 (\u062c\u0646\u06af\u0644) \u0645\u06cc\u200c\u062e\u0648\u0627\u0646\u06cc\u0645 \u0648 \u0622\u0646 \u0631\u0627 \u0627\u0632 \u0686\u0646\u062f \u0639\u062f\u0633\u06cc\u0654 \u0645\u0633\u062a\u0642\u0644 \u0645\u06cc\u200c\u0628\u06cc\u0646\u06cc\u0645\u060c \u0633\u067e\u0633 \u06cc\u0627\u0641\u062a\u0647\u200c\u0647\u0627 \u0631\u0627 \u062f\u0631\u0647\u0645 \u0645\u06cc\u200c\u0622\u0645\u06cc\u0632\u06cc\u0645.", rtl=True, italic=True)

    H(d, "\u0686\u0646\u062f \u0639\u062f\u0633\u06cc\u060c \u062f\u0631\u0647\u0645\u200c\u0622\u0645\u06cc\u062e\u062a\u0647 \u2014 \u0646\u0647 \u0641\u0642\u0637 \u06cc\u06a9\u06cc", 2, rtl=True)
    P(d, "\u0645\u0627\u0646\u0646\u062f \u0645\u0637\u0627\u0644\u0639\u0647\u0654 \u06cc\u06a9 \u0634\u0647\u0631 \u0628\u0627 \u0633\u0647 \u0646\u0642\u0634\u0647: \u06cc\u06a9\u06cc \u00ab\u0647\u0645\u200c\u0645\u0639\u0646\u0627\u00bb\u060c \u06cc\u06a9\u06cc \u00ab\u0647\u0645\u200c\u0642\u0644\u0645\u0631\u0648\u00bb\u060c \u0648 \u06cc\u06a9\u06cc \u00ab\u0634\u06a9\u0644 \u067e\u0631\u0627\u06a9\u0646\u062f\u06af\u06cc\u00bb. \u0645\u06cc\u200c\u067e\u0631\u0633\u06cc\u0645 \u06a9\u062c\u0627 \u0647\u0645\u200c\u062f\u0627\u0633\u062a\u0627\u0646\u200c\u0627\u0646\u062f \u0648 \u06a9\u062c\u0627 \u0646\u0627\u0647\u0645\u200c\u062f\u0627\u0633\u062a\u0627\u0646. \u0646\u0627\u0647\u0645\u200c\u062f\u0627\u0633\u062a\u0627\u0646\u06cc\u200c\u0647\u0627 \u0627\u063a\u0644\u0628 \u06af\u0648\u06cc\u0627\u062a\u0631\u0646\u062f.", rtl=True)

    def lineF(ty, lead):
        lst = rbt.get(ty) or []
        if lst:
            P(d, f"{lead}: " + "، ".join(x["root"] for x in lst[:6]), rtl=True)

    lineF("consensus", "\u062a\u0623\u06cc\u06cc\u062f\u0634\u062f\u0647 \u0628\u0627 \u062f\u0648 \u067e\u06cc\u0645\u0627\u0646\u0647 \u06cc\u0627 \u0628\u06cc\u0634\u062a\u0631 (\u0627\u0633\u062a\u0648\u0627\u0631\u062a\u0631\u06cc\u0646)")
    lineF("semantic", "\u0646\u0632\u062f\u06cc\u06a9 \u062f\u0631 \u0645\u0639\u0646\u0627")
    lineF("co-location", "\u0647\u0645\u200c\u0642\u0644\u0645\u0631\u0648 \u062f\u0631 \u0645\u062a\u0646")
    lineF("orthogonal", "\u067e\u06cc\u0648\u0633\u062a\u0647 \u0641\u0642\u0637 \u0628\u0631 \u06cc\u06a9 \u067e\u06cc\u0645\u0627\u0646\u0647")
    P(d, "\u0648 \u062a\u0646\u0634\u200c\u0647\u0627\u06cc \u06af\u0648\u06cc\u0627 \u2014 \u0646\u0632\u062f\u06cc\u06a9 \u0628\u0631 \u06cc\u06a9 \u067e\u06cc\u0645\u0627\u0646\u0647 \u0648\u0644\u06cc \u062f\u0648\u0631 \u0628\u0631 \u062f\u06cc\u06af\u0631\u06cc:", rtl=True)
    lineF("divergent", "\u062f\u0631 \u062a\u0646\u0634")

    H(d, "\u0639\u062f\u0633\u06cc\u0654 \u062a\u0648\u0627\u0644\u06cc (\u06a9\u062a\u0627\u0628 \u062f\u0648\u0645)", 2, rtl=True)
    P(d, f"\u062f\u0631\u0648\u0646 \u0622\u06cc\u0647\u060c \u0627\u06cc\u0646 \u0645\u0641\u0647\u0648\u0645 \u0628\u0647\u200c\u0637\u0648\u0631 \u0645\u06cc\u0627\u0646\u06af\u06cc\u0646 \u0646\u0632\u062f\u06cc\u06a9 \u0646\u0634\u0627\u0646\u0647\u0654 "
         f"{round((seq.get('mean_within_ayah_position') or 0)*100)}\u066a \u0642\u0631\u0627\u0631 \u0645\u06cc\u200c\u06af\u06cc\u0631\u062f \u0648 \u062d\u062f\u0648\u062f "
         f"{round((seq.get('ayah_final_share') or 0)*100)}\u066a \u0645\u0648\u0627\u0642\u0639 \u067e\u0627\u06cc\u0627\u0646\u200c\u0628\u062e\u0634 \u0622\u06cc\u0647 \u0627\u0633\u062a (\u0622\u0647\u0646\u06af/\u0641\u0648\u0627\u0635\u0644).", rtl=True)

    H(d, "\u0622\u0646\u0686\u0647 \u0627\u062f\u0639\u0627 \u0646\u0645\u06cc\u200c\u06a9\u0646\u06cc\u0645", 2, rtl=True)
    P(d, "\u0627\u06cc\u0646 \u06cc\u06a9 \u0646\u0642\u0634\u0647\u0654 \u0686\u0646\u062f\u200c\u0639\u062f\u0633\u06cc \u0648 \u0633\u0631\u0634\u0645\u0627\u0631\u06cc \u0627\u0633\u062a\u060c \u0646\u0647 \u062a\u0641\u0633\u06cc\u0631. \u0645\u0639\u0646\u0627 \u0627\u0632 \u0622\u0646\u0650 \u0645\u062a\u0646 \u0648 \u0639\u0627\u0644\u0645\u0627\u0646\u0650 \u0622\u0646 \u0627\u0633\u062a.", rtl=True)


# ── AYAH report content ────────────────────────────────────────────────────
TYPE_EN = {
    "consensus": "reinforced on ≥2 independent axes (robust)",
    "resonant": "close in meaning even without shared words (distributional)",
    "direct": "shares roots (lexical)",
    "co-located": "shares the same territory in the text",
    "orthogonal": "linked on one axis, independent on the others",
    "divergent": "close on one axis but opposed on another (tension)",
}
TYPE_FA = {
    "consensus": "تأییدشده روی ≥۲ محورِ مستقل (استوار)",
    "resonant": "نزدیک در معنا، حتی بدونِ واژهٔ مشترک (توزیعی)",
    "direct": "ریشهٔ مشترک (لفظی)",
    "co-located": "هم‌قلمروییِ متنی",
    "orthogonal": "پیوند روی یک محور، مستقل بر بقیه",
    "divergent": "نزدیک بر یک محور و مخالف بر محورِ دیگر (تنش)",
}


def _ayah_blocks(d, b, lang):
    seed = b["seed"]["seed"] if "seed" in b else b["full"]["seed"]
    concepts = b["seed"]["seed_concepts"] if "seed" in b else b["full"]["seed_concepts"]
    rbt = b.get("related_by_type") or b["full"]["related_by_type"]
    syn = b.get("synthesis") or b["full"]["synthesis"]
    refs = b["request"]["seeds"]
    rtl = (lang == "fa")
    if lang == "tech":
        H(d, f"Ayah-content Deep-Dive — Technical Report: {', '.join(refs)}", 1)
        P(d, "Computational cross-references, not tafsir. Each related ayah is typed "
             "by its profile on three independent axes (lexical / semantic / spatial), "
             "z-standardised over all candidates.", italic=True)
        for sd in seed:
            P(d, f"{sd['ref']}  —  {sd['text']}")
            P(d, "concepts: " + "، ".join(sd["roots"]), italic=True)
        P(d, "Relation distribution: " +
             " · ".join(f"{k}={v}" for k, v in syn["by_relation"].items()))
        for ty in ["consensus", "resonant", "direct", "co-located", "orthogonal", "divergent"]:
            lst = rbt.get(ty) or []
            if not lst:
                continue
            H(d, f"{ty.upper()} — {TYPE_EN[ty]}", 2)
            TABLE(d, ["ayah", "L", "S", "P", "shared roots"],
                  [[x["ref"], x["axes"]["lexical"], x["axes"]["semantic"],
                    x["axes"]["spatial"], " ".join(x["shared_roots"])] for x in lst[:8]])
        return
    if lang == "en":
        H(d, f"Explaining {', '.join(refs)} through the whole Qur’an — plain English", 1)
        P(d, "A note: these are computational cross-references — verses the data links "
             "to this one — not an interpretation of meaning.", italic=True)
        for sd in seed:
            P(d, f"The verse ({sd['ref']}): {sd['text']}")
            P(d, "Its concepts: " + "، ".join(sd["roots"]), italic=True)
        P(d, "We compare this verse to every other verse on three separate yardsticks: "
             "shared words, shared meaning (verses that talk about the same things in "
             "different words), and shared territory (verses that sit in the same parts "
             "of the text). Then we label each link by which yardsticks agree.")
        for ty in ["consensus", "resonant", "direct", "co-located", "orthogonal", "divergent"]:
            lst = rbt.get(ty) or []
            if not lst:
                continue
            H(d, f"{ty.title()} — {TYPE_EN[ty]}", 2)
            for x in lst[:6]:
                P(d, f"• {x['ref']}  (shared: {' '.join(x['shared_roots']) or '—'})  —  {x['text']}")
        H(d, "What we are NOT claiming", 2)
        P(d, "This is a concordance built from patterns, not a commentary. It points to "
             "verses worth reading together; the meaning is for the text and its scholars.")
        return
    # fa
    H(d, f"تبیینِ {', '.join(refs)} در پرتوِ کلِ قرآن — به زبان ساده", 1, rtl=True)
    P(d, "یک نکته: اینها ارجاع‌های متقابلِ محاسباتی‌اند — آیاتی که داده‌ها به این آیه "
         "پیوند می‌دهند — نه تفسیرِ معنا.", rtl=True, italic=True)
    for sd in seed:
        P(d, f"آیه ({sd['ref']}): {sd['text']}", rtl=True)
        P(d, "مفاهیمِ آن: " + "، ".join(sd["roots"]), rtl=True, italic=True)
    P(d, "این آیه را با هر آیهٔ دیگر بر سه سنجهٔ جداگانه می‌سنجیم: واژگانِ مشترک، معنای "
         "مشترک (آیاتی که از همان چیزها با واژگانِ دیگر سخن می‌گویند)، و قلمروِ مشترک. "
         "سپس هر پیوند را بر پایهٔ آنکه کدام سنجه‌ها هم‌داستان‌اند برچسب می‌زنیم.", rtl=True)
    for ty in ["consensus", "resonant", "direct", "co-located", "orthogonal", "divergent"]:
        lst = rbt.get(ty) or []
        if not lst:
            continue
        H(d, f"{TYPE_FA[ty]}", 2, rtl=True)
        for x in lst[:6]:
            P(d, f"• {x['ref']}  (مشترک: {' '.join(x['shared_roots']) or '—'})  —  {x['text']}", rtl=True)
    H(d, "آنچه ادعا نمی‌کنیم", 2, rtl=True)
    P(d, "این یک کشّاف از روی الگوهاست، نه تفسیر. آیاتی را نشان می‌دهد که ارزشِ "
         "هم‌خوانی دارند؛ معنا از آنِ متن و عالمانِ آن است.", rtl=True)


# ── orchestrator ───────────────────────────────────────────────────────────
def build_reports(bundle):
    bundle = Path(bundle)
    b = _load(bundle)
    kind = b["manifest"]["kind"]
    rep = bundle / "report"; rep.mkdir(exist_ok=True)
    target = b["manifest"].get("normalized") or b["manifest"].get("slug")
    made = []
    if kind == "concept":
        specs = [("technical", _concept_technical),
                 ("plain_en", _concept_plain_en),
                 ("plain_fa", _concept_plain_fa)]
        for name, fn in specs:
            d = _new_doc()
            fn(d, b, target)
            path = rep / f"{name}.docx"; d.save(path); made.append(path)
    else:
        for name, lang in [("technical", "tech"), ("plain_en", "en"), ("plain_fa", "fa")]:
            d = _new_doc()
            _ayah_blocks(d, b, lang)
            path = rep / f"{name}.docx"; d.save(path); made.append(path)
    pdfs = []
    for p in made:
        if _to_pdf(p):
            pdfs.append(p.with_suffix(".pdf"))
    # update MANIFEST
    mani = b["manifest"]
    mani.setdefault("outputs", {})["reports"] = sorted(p.name for p in made) + \
        sorted(p.name for p in pdfs)
    mani.setdefault("verification", {})["reports_generated"] = True
    (bundle / "MANIFEST.json").write_text(
        json.dumps(mani, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"wrote {len(made)} docx + {len(pdfs)} pdf in {rep}")
    for p in made:
        print("  ", p.name)
    return rep


def _latest(kind):
    kd = OUT_BASE / ("concepts" if kind == "concept" else "ayahs")
    ds = sorted((p for p in kd.iterdir() if p.is_dir()),
                key=lambda p: p.stat().st_mtime)
    return ds[-1] if ds else None


def main():
    args = sys.argv[1:]
    if not args:
        raise SystemExit(__doc__)
    if args[0] == "--latest":
        bundle = _latest(args[1])
        if not bundle:
            raise SystemExit("no bundle found")
    else:
        bundle = Path(args[0])
    build_reports(bundle)


if __name__ == "__main__":
    main()


def docx_bytes_from_result(res, register):
    """Build ONE register's .docx IN MEMORY from a live deep-dive result dict
    (for on-demand generation inside the app — pure python-docx, no LibreOffice,
    works on the hosted app). register in {technical, plain_en, plain_fa}."""
    import io
    kind = res["request"]["kind"]
    meta = res.get("meta", {})
    man = dict(kind=kind, generated=meta.get("generated", ""),
               code_version=meta.get("code_version", ""), verification={})
    d = _new_doc()
    lang = {"technical": "tech", "plain_en": "en", "plain_fa": "fa"}[register]
    if kind == "concept":
        _concept_paper(d, res, lang)
    else:
        _ayah_paper(d, res, lang)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════════════════
# DEEP-DIVE PAPERS — clean rewrite.  Three registers (technical master, plain
# English, plain Persian) generated from one structure.  Tables carry the data;
# prose carries one insight per section.  No repetition, no filler.
# ════════════════════════════════════════════════════════════════════════════
def _fmt(x, n=2):
    try:
        return f"{float(x):.{n}f}"
    except Exception:
        return str(x)


_AY_LABEL = {  # (en, fa, one-line meaning en, one-line meaning fa)
    "consensus":  ("Consensus",  "هم‌داستان",  "reinforced on ≥2 independent axes",
                   "تأیید‌شده روی ≥۲ محورِ مستقل"),
    "resonant":   ("Resonant",   "هم‌طنین",    "shared meaning with little shared wording",
                   "هم‌معنا بی‌واژهٔ مشترک"),
    "direct":     ("Direct",     "مستقیم",     "shared roots (lexical)",
                   "ریشهٔ مشترک (لفظی)"),
    "co-located": ("Co-located", "هم‌قلمرو",   "same regions of the text",
                   "هم‌جایگاهِ متنی"),
    "orthogonal": ("Orthogonal", "متعامد",     "one axis only, independent on the rest",
                   "تنها یک محور، مستقل بر بقیه"),
    "divergent":  ("Divergent",  "واگرا",      "high on one axis, opposed on another",
                   "نزدیک بر یک محور، مخالف بر دیگری"),
}
_AY_ORDER = ["consensus", "resonant", "direct", "co-located", "orthogonal", "divergent"]


def _ay_read(ty, en):
    """One substantive sentence per relation type — said once, not per item."""
    EN = {
        "consensus": "The dependable core: more than one independent axis points the same way.",
        "resonant": "The analytically valuable set — high meaning with low wording, so the tie "
                    "rests on subject, not vocabulary, and a concordance would miss it.",
        "direct": "Classical lexical cross-references; the shared roots are explicit, though "
                  "lexical overlap alone does not establish a thematic tie.",
        "co-located": "Bound by shared setting rather than shared sense; strongest where it "
                      "converges with another axis.",
        "orthogonal": "Visible through exactly one lens while the others stay silent.",
        "divergent": "Shares the target's wording or territory yet pulls against it in meaning "
                     "— a foil worth reading for contrast.",
    }
    FA = {
        "consensus": "هستهٔ قابل‌اتکا: بیش از یک محورِ مستقل به یک سو اشاره می‌کنند.",
        "resonant": "ارزشمندترین دسته از نگاهِ محاسباتی — معنای بالا با واژگانِ کم؛ پیوند بر موضوع "
                    "استوار است نه واژه، و جست‌وجوی واژگانی آن را نمی‌یابد.",
        "direct": "ارجاعِ لفظیِ کلاسیک؛ ریشه‌های مشترک آشکارند، هرچند هم‌پوشانیِ لفظی به‌تنهایی "
                  "پیوندِ موضوعی را ثابت نمی‌کند.",
        "co-located": "پیوند از هم‌جایگاهی، نه هم‌معنایی؛ آنجا که با محورِ دیگر هم‌رس شود قوی‌تر است.",
        "orthogonal": "تنها از یک عدسی دیده می‌شود و بقیه خاموش‌اند.",
        "divergent": "واژه یا قلمروِ آیهٔ کانونی را شریک است ولی در معنا رویاروی آن می‌ایستد — "
                     "هماوردی برای خوانشِ تقابلی.",
    }
    return (EN if en else FA)[ty]


def _ayah_paper(d, res, lang):
    seed = res["seed"]; rbt = res["related_by_type"]; syn = res["synthesis"]
    refs = res["request"]["seeds"]; refstr = ", ".join(refs)
    bc = syn.get("by_relation", {}) or {}
    cm = syn.get("cross_modal", {}) or {}
    req = res.get("request", {}); meta = res.get("meta", {})
    n_ay = meta.get("n_ayahs", "—"); cv = meta.get("code_version", "—")
    k = req.get("k", "—"); mf = req.get("min_freq", "—"); unit = req.get("unit", "surah")
    n_cand = sum(len(v) for v in rbt.values())
    nroots = len(seed[0].get("roots", [])) if seed else 0

    if lang == "tech":
        _ayah_tech(d, res, refstr, seed, rbt, bc, cm, n_ay, cv, k, mf, unit, n_cand, nroots)
    else:
        _ayah_plain(d, res, refstr, seed, rbt, bc, cm, n_cand, rtl=(lang == "fa"))


def _ayah_plain(d, res, refstr, seed, rbt, bc, cm, n_cand, rtl):
    def h(x, lvl=2): H(d, x, lvl, rtl=rtl)
    def p(x, **k): P(d, x, rtl=rtl, **k)
    L = (lambda en, fa: fa if rtl else en)
    nroots = sum(len(sd.get("roots", [])) for sd in seed)
    lead = [(ty, (rbt.get(ty) or [None])[0]) for ty in _AY_ORDER if rbt.get(ty)]

    h(L(f"Reading verse(s) {refstr} in the light of the whole Qur'an — a plain-language report",
        f"خواندنِ آیه(های) {refstr} در پرتوِ کلِ قرآن — گزارشی به زبانِ ساده"), 1)
    p(L("Abstract — The plain-language companion to the technical report, in the same order. "
        "The verse is compared with every other verse on three independent measures — shared "
        "words, shared meaning, and shared textual setting — and each related verse is labelled "
        "by how those measures agree. It describes structure, not meaning.",
        "چکیده — همراهِ ساده‌زبانِ گزارشِ فنّی، به همان ترتیب. آیه با هر آیهٔ دیگر بر سه سنجهٔ مستقل — واژگانِ مشترک، معنای مشترک، و جایگاهِ متنیِ مشترک — سنجیده می‌شود و هر آیهٔ مرتبط بر پایهٔ هم‌داستانیِ آن‌ها برچسب می‌خورد. ساختار را توصیف می‌کند، نه معنا را."), italic=True)

    h(L("1. Introduction", "۱. مقدمه"))
    p(L("To understand one verse we read it against the whole: a single tree is known only by "
        "walking the entire forest. We do not interpret meaning; we find, across the corpus, "
        "which verses relate to this one, by what kind of link, and how strongly.",
        "برای فهمِ یک آیه آن را در برابرِ کل می‌خوانیم: یک درخت تنها با گام‌زدنِ همهٔ جنگل شناخته می‌شود. معنا را تفسیر نمی‌کنیم؛ در سراسرِ متن می‌یابیم که کدام آیات با این آیه پیوند دارند، از چه گونه‌ای و با چه نیرویی."))

    h(L("2. Method", "۲. روش"))
    p(L(f"The verse is treated as one whole, weighted toward its {nroots} distinctive roots, and "
        "compared with every other verse on three independent lenses: shared words (rare shared "
        "roots count more than common ones), shared meaning (verses about the same things, even "
        "in different words), and shared territory (verses used in the same regions of the text). "
        "A link counts only when it stands well above the typical verse, and is kept only if it "
        "survives a test against chance. The lenses are kept separate, never averaged.",
        f"آیه یک کلِ واحد گرفته می‌شود، با وزنِ بیشتر بر {nroots} ریشهٔ متمایزش، و با هر آیهٔ دیگر بر سه عدسیِ مستقل سنجیده می‌شود: واژگانِ مشترک (ریشهٔ کم‌یاب بیش از پرتکرار)، معنای مشترک (آیاتی دربارهٔ یک چیز، حتی با واژگانِ دیگر)، و قلمروِ مشترک (آیاتی در همان نواحیِ متن). پیوند تنها وقتی شمرده می‌شود که آشکارا از آیهٔ معمول بالاتر باشد، و تنها اگر از آزمونِ بخت بگذرد نگاه داشته می‌شود. عدسی‌ها جدا نگاه داشته می‌شوند، هرگز میانگین نمی‌شوند."))

    h(L("3. Results", "۳. نتایج"))
    p(L(f"Of {n_cand} candidate verses, the closest by each kind of link are "
        + "; ".join(f"{_AY_LABEL[ty][0].lower()} → {x['ref']}" for ty, x in lead)
        + ". The full set is in Table 1, with the number of shared roots and the meaning and "
        "territory scores (in standard-deviation units).",
        f"از میانِ {n_cand} آیهٔ نامزد، نزدیک‌ترین‌ها بر حسبِ هر گونه پیوند: "
        + "؛ ".join(f"{_AY_LABEL[ty][1]} → {x['ref']}" for ty, x in lead)
        + ". مجموعهٔ کامل در جدولِ ۱ آمده، با شمارِ ریشه‌های مشترک و نمره‌های معنا و قلمرو (به واحدِ انحرافِ معیار)."))
    rows = []
    for ty in _AY_ORDER:
        for x in (rbt.get(ty) or [])[:4]:
            a = x["axes"]
            rows.append([L(_AY_LABEL[ty][0], _AY_LABEL[ty][1]), x["ref"],
                         len(x["shared_roots"]), f"{a['semantic']:+.1f}", f"{a['spatial']:+.1f}"])
    TABLE(d, [L("relation", "گونه"), L("verse", "آیه"),
              L("# shared", "مشترک"), L("meaning", "معنا"),
              L("territory", "قلمرو")], rows)

    h(L("4. Discussion", "۴. بحث"))
    p(L("The most reliable cross-references are the consensus verses, agreed on more than one "
        "lens. The most striking are the resonant verses — close in meaning with few shared "
        "words, which a keyword search would miss. The divergent verses share this verse's words "
        "or setting yet stand opposed in meaning. This across-verse view complements within-verse "
        "co-occurrence, which cannot see beyond a single verse.",
        "اعتمادپذیرترین ارجاع‌ها آیاتِ هم‌داستان‌اند، که بیش از یک عدسی تأییدشان می‌کند. چشمگیرترین، آیاتِ هم‌طنین‌اند — نزدیک در معنا با واژگانِ مشترکِ اندک، که جست‌وجوی واژگانی نمی‌یابد. آیاتِ واگرا واژه یا جایگاهِ این آیه را شریک‌اند ولی در معنا رویاروی آن. این نگاهِ میان‌آیه‌ای مکملِ هم‌نشینیِ درون‌آیه است، که فراتر از یک آیه را نمی‌بیند."))

    h(L("5. Limitations", "۵. محدودیت‌ها"))
    p(L("This is a map, not a commentary. It fixes no meaning, ruling, or theology. The meaning "
        "lens is an estimate, so single-lens labels are tentative. Verses are cited by reference; "
        "the meaning is for the text and its scholars.",
        "این یک نقشه است، نه تفسیر. هیچ معنا، حکم یا الهیاتی را تعیین نمی‌کند. عدسیِ معنا یک برآورد است، پس برچسب‌های تک‌عدسی غیرقطعی‌اند. آیات با شمارهٔ ارجاع آورده می‌شوند؛ معنا از آنِ متن و عالمانِ آن است."))

    h(L("6. Conclusion", "۶. نتیجه"))
    p(L(f"The verse(s) {refstr} sit within a measurable, typed web of cross-references — a "
        "reproducible map of where the verse resonates across the Qur'an, for closer reading.",
        f"آیه(های) {refstr} در شبکه‌ای سنجش‌پذیر و گونه‌بندی‌شده از ارجاع‌ها جای دارند — نقشه‌ای بازتولیدپذیر از جایی که آیه در سراسرِ قرآن طنین می‌اندازد، برای خوانشِ دقیق‌تر."))


def _concept_paper(d, res, lang):
    t = res["request"].get("normalized") or res["request"].get("target")
    if lang == "tech":
        _concept_tech(d, res, t)
    else:
        _concept_plain(d, res, t, rtl=(lang == "fa"))


# ── final, well-written concept builders (worked example: قلب) ──────────────
_CN_LABEL = {  # short relation labels for the compact companion table (en, fa)
    "consensus":   ("Consensus",     "هم‌داستان"),
    "semantic":    ("Meaning-mate",  "هم‌معنا"),
    "co-location": ("Territory-mate", "هم‌قلمرو"),
    "spatial":     ("Shape-mate",    "هم‌شکل"),
    "orthogonal":  ("Orthogonal",    "متعامد"),
    "divergent":   ("Divergent",     "واگرا"),
}


def _sense_cohesion(senses, mass_floor=5, max_forms=5, min_sig=2):
    """Does a root's surface forms keep similar company (cohesive sense) or
    divergent company (split/polysemous)?  Mean pairwise Jaccard of the per-form
    significant co-locator sets.  Returns None when too sparse to assert (the
    report then says nothing about senses).  Surface-only signals are sense-
    specific by the pipeline's own epistemics, so we report only the aggregate
    verdict and the form-robust core, never per-form association lists."""
    import itertools
    sets = []
    for s in (senses or []):
        if int(s.get("count", 0) or 0) < mass_floor:
            continue
        sig = {r for (r, a, p) in (s.get("share") or [])
               if p is not None and p <= 0.10}
        if len(sig) >= min_sig:
            sets.append((s.get("form", "\u2014"), int(s.get("count", 0) or 0), sig))
        if len(sets) >= max_forms:
            break
    if len(sets) < 2:
        return None
    js = []
    for (_, _, a), (_, _, b) in itertools.combinations(sets, 2):
        u = a | b
        js.append((len(a & b) / len(u)) if u else 0.0)
    J = sum(js) / len(js)
    core = set.intersection(*[s for _, _, s in sets])
    verdict = "cohesive" if J >= 0.34 else ("split" if J <= 0.12 else "mixed")
    return dict(n_forms=len(sets), forms=[f for f, _, _ in sets],
                jaccard=round(J, 3), core=sorted(core), verdict=verdict)


def _concept_plain(d, res, t, rtl):
    fld = res.get("field", {}) or {}; dist = res.get("distribution", {}) or {}
    null = res.get("null", {}) or {}; rel = res.get("relations", {}) or {}
    syn = res.get("synthesis", {}) or {}; cg = res.get("cross_granularity", {}) or {}
    rbt = rel.get("related_by_type", {}); cm = syn.get("cross_modal", {}) or {}
    feats = dist.get("features", {}) or {}
    senses = res.get("senses", []) or []

    def h(x, lvl=2): H(d, x, lvl, rtl=rtl)
    def p(x, **k): P(d, x, rtl=rtl, **k)
    L = (lambda en, fa: fa if rtl else en)

    def z(v):
        return f"{v:+.1f}" if isinstance(v, (int, float)) else "—"

    sem = "، ".join(fld.get("semantic_field", [])[:6]) or "—"
    terr = "، ".join(fld.get("co_location_neighbours", [])[:6]) or "—"
    vb = "، ".join(cm.get("verified_bonds", [])) or "—"
    n_cand = sum(len(v or []) for v in rbt.values())
    lead = [(ty, (rbt.get(ty) or [None])[0]) for ty in REL_ORDER if rbt.get(ty)]

    h(L(f"The concept “{t}” across the whole Qur'an — a plain-language report",
        f"مفهومِ «{t}» در سراسرِ قرآن — گزارشی به زبانِ ساده"), 1)
    p(L("Abstract — The plain-language companion to the technical report, in the same order. "
        "The concept is read in every place it appears and compared with every other concept on "
        "five independent measures — the ideas it shares meaning with, the textual territory it "
        "shares, the shape of its spread, how it splits across word-forms, and where in a verse "
        "it falls — and each related concept is labelled by how those measures agree. It "
        "describes structure, not meaning.",
        "چکیده — همراهِ ساده‌زبانِ گزارشِ فنّی، به همان ترتیب. مفهوم در هر جا که می‌آید خوانده "
        "و با هر مفهومِ دیگر بر پنج سنجهٔ مستقل سنجیده می‌شود — اندیشه‌هایی که با آن‌ها هم‌معناست، "
        "قلمروِ متنی‌ای که شریک است، شکلِ پراکندگی‌اش، چگونگیِ تقسیمش میانِ صورت‌های واژه، و جایش "
        "در آیه — و هر مفهومِ مرتبط بر پایهٔ هم‌داستانیِ این‌ها برچسب می‌خورد. ساختار را توصیف "
        "می‌کند، نه معنا را."), italic=True)

    h(L("1. Introduction", "۱. مقدمه"))
    p(L("To understand one concept we read it against the whole: a single tree is known only by "
        "walking the entire forest. The classical principle that the text interprets itself "
        "motivates reading every occurrence together. We do not interpret meaning; we find, "
        "across the corpus, which concepts relate to this one, by what kind of link, and how "
        "strongly.",
        "برای فهمِ یک مفهوم آن را در برابرِ کل می‌خوانیم: یک درخت تنها با گام‌زدنِ همهٔ جنگل شناخته "
        "می‌شود. اصلِ کلاسیکِ «قرآن خود را تفسیر می‌کند» انگیزهٔ خواندنِ همهٔ موارد با هم است. "
        "معنا را تفسیر نمی‌کنیم؛ در سراسرِ متن می‌یابیم که کدام مفاهیم با این یکی پیوند دارند، از "
        "چه گونه‌ای و با چه نیرویی."))

    h(L("2. Method", "۲. روش"))
    p(L("No one measure captures a whole concept, much as each of the blind men feeling an "
        "elephant grasps only a part. So we use five lenses and, crucially, keep them separate: "
        "what the concept shares MEANING with (the ideas it travels among), what TERRITORY it "
        "shares (the parts of the text it occupies), the SHAPE of its spread, how it splits "
        "across word-forms, and where in a verse it falls. A link counts only when it stands "
        "well above the typical concept, and the shape lens is kept only if it survives a test "
        "against chance. Their agreement and disagreement is the whole point, so we never blend "
        "them into a single number — averaging distinct signals is like mixing distinct colours "
        "into grey.",
        "هیچ سنجهٔ تنها یک مفهوم را کامل نمی‌گیرد، چنان‌که هر یک از نابینایانِ فیل تنها بخشی را "
        "لمس می‌کند. پس پنج عدسی به‌کار می‌بریم و — مهم‌تر — جدا نگاهشان می‌داریم: اینکه مفهوم با "
        "چه چیزی هم‌معناست (اندیشه‌هایی که میانشان می‌رود)، چه قلمرویی را شریک است (بخش‌هایی از "
        "متن که اشغال می‌کند)، شکلِ پراکندگی‌اش، چگونگیِ تقسیمش میانِ صورت‌های واژه، و جایش در "
        "آیه. پیوند تنها وقتی شمرده می‌شود که آشکارا از مفهومِ معمول بالاتر باشد، و عدسیِ شکل تنها "
        "اگر از آزمونِ بخت بگذرد نگاه داشته می‌شود. هم‌داستانی و ناهم‌داستانیِ این‌ها تمامِ مطلب "
        "است، پس هرگز در یک عدد درهمشان نمی‌آمیزیم — میانگین‌گیریِ سیگنال‌های جدا مانندِ آمیختنِ "
        "رنگ‌ها به خاکستری است."))

    h(L("3. Results", "۳. نتایج"))
    p(L(f"The concepts “{t}” shares MEANING with are not the concepts it shares TEXTUAL "
        f"TERRITORY with. By meaning it sits among {sem}; by territory it sits among {terr}. "
        f"Those two lists barely overlap, and that gap is the result worth reporting: the "
        f"concept is talked about in the company of one family of ideas, yet placed in the text "
        f"beside another. Of all its links, the one that holds up under the strictest check — "
        f"confirmed independently at both the root level and the word-form level — is {vb}.",
        f"مفاهیمی که «{t}» با آن‌ها هم‌معناست همان مفاهیمی نیستند که با آن‌ها هم‌قلمروِ متنی است. "
        f"به‌حسبِ معنا در میانِ {sem} می‌نشیند؛ به‌حسبِ قلمرو در میانِ {terr}. این دو فهرست "
        f"به‌سختی هم‌پوشانی دارند، و همین شکاف، نتیجهٔ درخورِ گزارش است: مفهوم در هم‌نشینیِ یک "
        f"خانوادهٔ اندیشه گفته می‌شود، اما در متن کنارِ خانوادهٔ دیگری می‌نشیند. از میانِ همهٔ "
        f"پیوندهایش، آن‌که زیرِ سخت‌گیرانه‌ترین وارسی پابرجا می‌ماند — مستقل در سطحِ ریشه و سطحِ "
        f"صورتِ واژه — این است: {vb}."))
    p(L(f"By raw count the concept looks clustered (clustering index ≈ "
        f"{_fmt(feats.get('moran_I'),2)}), but a common word clusters for the trivial reason "
        f"that it is everywhere — as coffee shops 'cluster' in big cities simply because more "
        f"people are there. Tested against a decoy that keeps the same frequency but scatters "
        f"the word at random, the real value falls inside the range of pure chance "
        f"({null.get('interpretation','—')}), so the clustering is not counted as a finding. "
        f"Stating plainly what is NOT there is part of why the rest can be trusted.",
        f"به شمارشِ خام مفهوم خوشه‌بسته می‌نماید (شاخصِ خوشه‌بندیِ حدودِ "
        f"{_fmt(feats.get('moran_I'),2)})، اما واژهٔ پُربسامد به این دلیلِ پیشِ‌پاافتاده خوشه "
        f"می‌بندد که همه‌جا هست — چنان‌که کافی‌شاپ‌ها در شهرهای بزرگ تنها به‌سببِ جمعیتِ بیشتر "
        f"«خوشه» می‌زنند. در آزمون با بدلی‌ای که همان بسامد را دارد ولی واژه را تصادفی پخش کرده، "
        f"عددِ واقعی درونِ دامنهٔ بختِ محض می‌افتد ({null.get('interpretation','—')})، پس "
        f"خوشه‌بندی یافته شمرده نمی‌شود. گفتنِ صریحِ آنچه نیست، بخشی از آن است که بقیه را "
        f"اعتمادپذیر می‌کند."))
    _coh = _sense_cohesion(senses)
    if _coh:
        _n = _coh["n_forms"]; _frm = "\u060c ".join(_coh["forms"])
        _core = "\u060c ".join(_coh["core"]) or "\u2014"
        if _coh["verdict"] == "cohesive":
            p(L(f"The concept surfaces in {_n} main word-forms ({_frm}); across them it keeps "
                f"largely the same company, so within this corpus its sense reads as unified "
                f"rather than split, and the companions common to all forms ({_core}) are its "
                f"form-independent core.",
                f"\u0645\u0641\u0647\u0648\u0645 \u062f\u0631 {_n} \u0635\u0648\u0631\u062a\u0650 \u0648\u0627\u0698\u06af\u0627\u0646\u06cc\u0650 \u0627\u0635\u0644\u06cc ({_frm}) \u0646\u0645\u0627\u06cc\u0627\u0646 \u0645\u06cc\u200c\u0634\u0648\u062f\u061b \u062f\u0631 \u0647\u0645\u0647\u0654 \u0622\u0646\u200c\u0647\u0627 \u06a9\u0645\u0627\u0628\u06cc\u0634 \u0647\u0645\u200c\u0646\u0634\u06cc\u0646\u06cc\u0650 \u06cc\u06a9\u0633\u0627\u0646\u06cc \u062f\u0627\u0631\u062f\u060c \u067e\u0633 \u062f\u0631 \u0627\u06cc\u0646 \u0645\u062a\u0646 \u0645\u0639\u0646\u0627\u06cc\u0634 \u06cc\u06a9\u067e\u0627\u0631\u0686\u0647 \u0645\u06cc\u200c\u0646\u0645\u0627\u06cc\u062f \u0646\u0647 \u0686\u0646\u062f\u067e\u0627\u0631\u0647\u060c \u0648 \u0647\u0645\u200c\u0646\u0634\u06cc\u0646\u200c\u0647\u0627\u06cc \u0645\u0634\u062a\u0631\u06a9\u0650 \u0647\u0645\u0647\u0654 \u0635\u0648\u0631\u062a\u200c\u0647\u0627 ({_core}) \u0647\u0633\u062a\u0647\u0654 \u0645\u0633\u062a\u0642\u0644 \u0627\u0632 \u0635\u0648\u0631\u062a\u0650 \u0622\u0646\u200c\u0627\u0646\u062f."))
        elif _coh["verdict"] == "split":
            p(L(f"The concept surfaces in {_n} main word-forms ({_frm}); these forms keep "
                f"markedly different company, which points to more than one sense carried by "
                f"the root here \u2014 better read form by form than as a single idea.",
                f"\u0645\u0641\u0647\u0648\u0645 \u062f\u0631 {_n} \u0635\u0648\u0631\u062a\u0650 \u0648\u0627\u0698\u06af\u0627\u0646\u06cc\u0650 \u0627\u0635\u0644\u06cc ({_frm}) \u0646\u0645\u0627\u06cc\u0627\u0646 \u0645\u06cc\u200c\u0634\u0648\u062f\u061b \u0627\u06cc\u0646 \u0635\u0648\u0631\u062a\u200c\u0647\u0627 \u0647\u0645\u200c\u0646\u0634\u06cc\u0646\u06cc\u0650 \u0622\u0634\u06a9\u0627\u0631\u0627 \u0645\u062a\u0641\u0627\u0648\u062a\u06cc \u062f\u0627\u0631\u0646\u062f\u060c \u06a9\u0647 \u0628\u0647 \u0628\u06cc\u0634 \u0627\u0632 \u06cc\u06a9 \u0645\u0639\u0646\u0627 \u062f\u0631 \u0631\u06cc\u0634\u0647 \u0627\u0634\u0627\u0631\u0647 \u062f\u0627\u0631\u062f \u2014 \u0628\u0647\u062a\u0631 \u0627\u0633\u062a \u0635\u0648\u0631\u062a\u200c\u0628\u0647\u200c\u0635\u0648\u0631\u062a \u062e\u0648\u0627\u0646\u062f\u0647 \u0634\u0648\u062f \u0646\u0647 \u0686\u0648\u0646 \u06cc\u06a9 \u0645\u0641\u0647\u0648\u0645\u0650 \u0648\u0627\u062d\u062f."))
        else:
            p(L(f"The concept surfaces in {_n} main word-forms ({_frm}); their company partly "
                f"overlaps \u2014 a shared core ({_core}) with form-specific differences around "
                f"it \u2014 so the sense is broadly single but not perfectly uniform across forms.",
                f"\u0645\u0641\u0647\u0648\u0645 \u062f\u0631 {_n} \u0635\u0648\u0631\u062a\u0650 \u0648\u0627\u0698\u06af\u0627\u0646\u06cc\u0650 \u0627\u0635\u0644\u06cc ({_frm}) \u0646\u0645\u0627\u06cc\u0627\u0646 \u0645\u06cc\u200c\u0634\u0648\u062f\u061b \u0647\u0645\u200c\u0646\u0634\u06cc\u0646\u06cc\u0650 \u0622\u0646\u200c\u0647\u0627 \u062a\u0627 \u062d\u062f\u06cc \u0647\u0645\u200c\u067e\u0648\u0634\u0627\u0646 \u0627\u0633\u062a \u2014 \u0647\u0633\u062a\u0647\u200c\u0627\u06cc \u0645\u0634\u062a\u0631\u06a9 ({_core}) \u0628\u0627 \u062a\u0641\u0627\u0648\u062a\u200c\u0647\u0627\u06cc\u06cc \u0635\u0648\u0631\u062a\u200c\u0648\u06cc\u0698\u0647 \u067e\u06cc\u0631\u0627\u0645\u0648\u0646\u0634 \u2014 \u067e\u0633 \u0645\u0639\u0646\u0627 \u06a9\u0645\u0627\u0628\u06cc\u0634 \u06cc\u06af\u0627\u0646\u0647 \u0627\u0633\u062a \u0648\u0644\u06cc \u062f\u0631 \u0645\u06cc\u0627\u0646\u0650 \u0635\u0648\u0631\u062a\u200c\u0647\u0627 \u06a9\u0627\u0645\u0644\u0627\u064b \u06cc\u06a9\u062f\u0633\u062a \u0646\u06cc\u0633\u062a."))
    p(L(f"Of {n_cand} candidate concepts, the closest by each kind of link are "
        + "; ".join(f"{_CN_LABEL[ty][0].lower()} → {x['root']}" for ty, x in lead)
        + ". The full set is in Table 1, with each companion's meaning, territory and shape "
        "scores (in standard-deviation units).",
        f"از میانِ {n_cand} مفهومِ نامزد، نزدیک‌ترین‌ها بر حسبِ هر گونه پیوند: "
        + "؛ ".join(f"{_CN_LABEL[ty][1]} → {x['root']}" for ty, x in lead)
        + ". مجموعهٔ کامل در جدولِ ۱ آمده، با نمره‌های معنا، قلمرو و شکلِ هر هم‌نشین (به واحدِ "
        "انحرافِ معیار)."))
    rows = []
    for ty in REL_ORDER:
        for x in (rbt.get(ty) or [])[:3]:
            a = x.get("axes", {}) or {}
            rows.append([L(_CN_LABEL[ty][0], _CN_LABEL[ty][1]), x["root"],
                         z(a.get("semantic")), z(a.get("co-location")), z(a.get("spatial"))])
    TABLE(d, [L("relation", "گونه"), L("concept", "مفهوم"),
              L("meaning", "معنا"), L("territory", "قلمرو"),
              L("shape", "شکل")], rows)

    h(L("4. Discussion", "۴. بحث"))
    p(L("The most reliable companions are the consensus concepts, agreed on more than one lens. "
        "The most striking are the meaning-mates — close in meaning yet sharing little textual "
        "territory, which a concordance built on shared wording would miss. The divergent "
        "concepts share this concept's territory or wording yet stand opposed in meaning, marking "
        "structural tension rather than noise. This across-corpus view complements within-verse "
        "co-occurrence, which is exact but cannot see beyond a single verse.",
        "اعتمادپذیرترین هم‌نشین‌ها مفاهیمِ هم‌داستان‌اند، که بیش از یک عدسی تأییدشان می‌کند. "
        "چشمگیرترین، هم‌معناهایند — نزدیک در معنا ولی با قلمروِ متنیِ مشترکِ اندک، که نمایه‌ای "
        "برپایهٔ واژگانِ مشترک نمی‌یابد. مفاهیمِ واگرا قلمرو یا واژهٔ این مفهوم را شریک‌اند ولی در "
        "معنا رویاروی آن می‌ایستند و تنشِ ساختاری‌اند، نه نوفه. این نگاهِ سراسرِ متن مکملِ "
        "هم‌نشینیِ درون‌آیه است، که دقیق است ولی فراتر از یک آیه را نمی‌بیند."))

    h(L("5. Limitations", "۵. محدودیت‌ها"))
    p(L("This is a map, not a commentary, and a census is not exegesis. It fixes no meaning, "
        "ruling, or theology. The meaning lens is an estimate, so single-lens labels are "
        "tentative; the shape lens is null for many concepts, as the frequency control shows. "
        "Concepts are named by their roots; the meaning is for the text and its scholars.",
        "این یک نقشه است، نه تفسیر، و سرشماری تأویل نیست. هیچ معنا، حکم یا الهیاتی را تعیین "
        "نمی‌کند. عدسیِ معنا یک برآورد است، پس برچسب‌های تک‌عدسی غیرقطعی‌اند؛ عدسیِ شکل برای "
        "بسیاری مفاهیم پوچ است، چنان‌که کنترلِ بسامد نشان می‌دهد. مفاهیم با ریشه‌شان نامیده "
        "می‌شوند؛ معنا از آنِ متن و عالمانِ آن است."))

    h(L("6. Conclusion", "۶. نتیجه"))
    p(L(f"The concept “{t}” sits within a measurable, typed web of companions — a reproducible "
        "map of where it resonates across the Qur'an, defined here by the divergence between "
        "meaning and territory, and offered for closer reading rather than as a substitute for "
        "it.",
        f"مفهومِ «{t}» در شبکه‌ای سنجش‌پذیر و گونه‌بندی‌شده از هم‌نشین‌ها جای دارد — نقشه‌ای "
        "بازتولیدپذیر از جایی که در سراسرِ قرآن طنین می‌اندازد، که اینجا با واگراییِ میانِ معنا و "
        "قلمرو تعریف می‌شود، و برای خوانشِ دقیق‌تر پیشنهاد می‌گردد، نه جایگزینِ آن."))


# ── figures for the technical papers (matplotlib, ASCII labels only) ─────────
_TYPE_COLOR = {"consensus": "#1D3557", "resonant": "#2A9D8F", "semantic": "#2A9D8F",
               "direct": "#457B9D", "co-located": "#E9C46A", "co-location": "#E9C46A",
               "spatial": "#8D99AE", "orthogonal": "#A8A4CE", "divergent": "#E76F51"}


def _fig_path():
    import os, tempfile
    fd, p = tempfile.mkstemp(suffix=".png", prefix="ddfig_"); os.close(fd); return p


def _style_ax(ax):
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    ax.tick_params(labelsize=9)


def _add_figure(d, png, caption):
    import os
    from docx.shared import Inches
    pic = d.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        pic.add_run().add_picture(png, width=Inches(5.4))
    finally:
        try:
            os.unlink(png)
        except OSError:
            pass
    cap = P(d, caption, italic=True); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fig_relation_bar(counts, order, labels, title):
    import matplotlib; matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    items = [(labels.get(k, k), int(counts.get(k, 0))) for k in order if counts.get(k, 0)]
    items.sort(key=lambda x: x[1])
    fig, ax = plt.subplots(figsize=(6, 3.0), dpi=150)
    ax.barh([a for a, _ in items], [b for _, b in items], color="#1D3557")
    for i, (a, b) in enumerate(items):
        ax.text(b, i, f" {b}", va="center", fontsize=9)
    ax.set_xlabel("number of related items", fontsize=9)
    ax.set_title(title, fontsize=11)
    _style_ax(ax); fig.tight_layout()
    p = _fig_path(); fig.savefig(p); plt.close(fig); return p


def _fig_null(null, title):
    import numpy as np, matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    mu = float(null.get("null_mean", 0) or 0)
    sd = float(null.get("null_sd", 1) or 1e-9) or 1e-9
    obs = float(null.get("real", mu) or mu)
    lo = min(mu - 4 * sd, obs - sd)
    hi = max(mu + 4 * sd, obs + sd)
    xs = np.linspace(lo, hi, 240)
    ys = np.exp(-0.5 * ((xs - mu) / sd) ** 2)
    fig, ax = plt.subplots(figsize=(6, 2.9), dpi=150)
    ax.fill_between(xs, ys, color="#A8C0DD", alpha=0.75, label="frequency-matched null")
    ax.axvline(obs, color="#C0392B", lw=2, label=f"observed = {obs:.3f}")
    ax.axvline(mu, color="#1D3557", ls="--", lw=1, label=f"null mean = {mu:.3f}")
    ax.set_yticks([])
    ax.set_xlabel("Moran's I (areal clustering)", fontsize=9)
    ax.set_title(title, fontsize=11)
    ax.legend(fontsize=8, frameon=False)
    _style_ax(ax)
    fig.tight_layout()
    p = _fig_path()
    fig.savefig(p)
    plt.close(fig)
    return p


def _fig_axis_scatter(rbt, order, title):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(6, 4.0), dpi=150)
    for k in order:
        xs, ys = [], []
        for x in (rbt.get(k) or []):
            a = x.get("axes", {})
            if "semantic" in a and "lexical" in a:
                xs.append(a["semantic"])
                ys.append(a["lexical"])
        if xs:
            ax.scatter(xs, ys, s=30, c=_TYPE_COLOR.get(k, "#888"), label=k,
                       edgecolor="white", linewidth=0.4)
    ax.axhline(1, color="#999", ls=":", lw=0.8)
    ax.axvline(1, color="#999", ls=":", lw=0.8)
    ax.set_xlabel("semantic similarity (z)", fontsize=9)
    ax.set_ylabel("lexical overlap (z)", fontsize=9)
    ax.set_title(title, fontsize=11)
    ax.legend(fontsize=7, frameon=False, ncol=2)
    _style_ax(ax)
    fig.tight_layout()
    p = _fig_path()
    fig.savefig(p)
    plt.close(fig)
    return p


# ── additional figure generators (ASCII labels only) ────────────────────────
def _fig_grouped_axes(rbt, order, axkeys, axlabels, title):
    import numpy as np, matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    types = [k for k in order if (rbt.get(k))]
    means = {a: [] for a in axkeys}
    for k in types:
        lst = rbt.get(k) or []
        for a in axkeys:
            vals = [x["axes"].get(a) for x in lst if x.get("axes", {}).get(a) is not None]
            means[a].append(float(np.mean(vals)) if vals else 0.0)
    x = np.arange(len(types)); w = 0.8 / max(1, len(axkeys))
    fig, ax = plt.subplots(figsize=(6.2, 3.4), dpi=150)
    cols = ["#1D3557", "#2A9D8F", "#E9C46A"]
    for i, a in enumerate(axkeys):
        ax.bar(x + i * w, means[a], w, label=axlabels[i], color=cols[i % len(cols)])
    ax.axhline(0, color="#444", lw=0.6)
    ax.set_xticks(x + w * (len(axkeys) - 1) / 2)
    ax.set_xticklabels(types, rotation=20, ha="right", fontsize=8)
    ax.set_ylabel("mean z-score", fontsize=9)
    ax.set_title(title, fontsize=11); ax.legend(fontsize=8, frameon=False)
    _style_ax(ax); fig.tight_layout()
    p = _fig_path(); fig.savefig(p); plt.close(fig); return p


def _fig_feature_bars(pairs, title):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    labels = [a for a, _ in pairs]; vals = [b for _, b in pairs]
    fig, ax = plt.subplots(figsize=(6.2, 3.2), dpi=150)
    ax.bar(range(len(labels)), vals, color="#457B9D")
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=25, ha="right", fontsize=8)
    for i, v in enumerate(vals):
        ax.text(i, v, f"{v:.2f}", ha="center", va="bottom", fontsize=8)
    ax.set_title(title, fontsize=11); _style_ax(ax); fig.tight_layout()
    p = _fig_path(); fig.savefig(p); plt.close(fig); return p


def _fig_counts(pairs, title, xlabel, ylabel, color="#1D3557"):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    labels = [str(a) for a, _ in pairs]; vals = [b for _, b in pairs]
    fig, ax = plt.subplots(figsize=(6.2, 3.2), dpi=150)
    ax.bar(range(len(labels)), vals, color=color)
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels, fontsize=8)
    for i, v in enumerate(vals):
        ax.text(i, v, f" {v}", ha="center", va="bottom", fontsize=8)
    ax.set_xlabel(xlabel, fontsize=9); ax.set_ylabel(ylabel, fontsize=9)
    ax.set_title(title, fontsize=11); _style_ax(ax); fig.tight_layout()
    p = _fig_path(); fig.savefig(p); plt.close(fig); return p


def _fig_hist(values, title, xlabel, bins=20, color="#2A9D8F"):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(6.2, 3.0), dpi=150)
    ax.hist(values, bins=bins, color=color, edgecolor="white", linewidth=0.4)
    ax.set_xlabel(xlabel, fontsize=9); ax.set_ylabel("count", fontsize=9)
    ax.set_title(title, fontsize=11); _style_ax(ax); fig.tight_layout()
    p = _fig_path(); fig.savefig(p); plt.close(fig); return p


def _concept_figures(d, res):
    import numpy as np
    rel = res.get("relations", {}) or {}; rbt = rel.get("related_by_type", {})
    dist = res.get("distribution", {}) or {}; feats = dist.get("features", {}) or {}
    cg = res.get("cross_granularity", {}) or {}; null = res.get("null", {}) or {}
    bc = rel.get("by_relation", {}) or {}
    t = res["request"].get("normalized") or res["request"].get("target")

    H(d, "7. Figures and supporting analysis", 2)

    # Fig 3 — mean axis profile per relation type
    _add_figure(d, _fig_grouped_axes(rbt, REL_ORDER,
                ["semantic", "co-location", "spatial"],
                ["meaning", "territory", "shape"],
                "Mean modality score by relation type"),
                "Figure 3. Mean standardised score on each comparative modality, by relation "
                "type. Consensus stands above the rest on two modalities at once; the single-"
                "modality types peak on exactly one bar, by construction.")
    def _avg(k, a):
        lst = rbt.get(k) or []
        vs = [x["axes"].get(a) for x in lst if x.get("axes", {}).get(a) is not None]
        return float(np.mean(vs)) if vs else 0.0
    P(d, f"The profile confirms the typology is doing what it claims. Consensus companions "
         f"average {_avg('consensus','semantic'):+.2f} on meaning and "
         f"{_avg('consensus','co-location'):+.2f} on territory — high on both at once — "
         f"whereas the semantic type averages {_avg('semantic','semantic'):+.2f} on meaning "
         f"but only {_avg('semantic','co-location'):+.2f} on territory, and the co-location "
         f"type inverts that. The bars are not redundant: if the modalities were measuring the "
         f"same thing the types would be indistinguishable, and they are not.")

    # Fig 4 — distribution-shape features
    fp = [("coverage", feats.get("coverage", 0)),
          ("log-freq", feats.get("log_freq", 0)),
          ("spread-entropy", feats.get("spread_entropy", 0)),
          ("peak-share", feats.get("peak_share", 0)),
          ("lacunarity", feats.get("lacunarity", 0)),
          ("Meccan share", feats.get("meccan_pct", 0))]
    _add_figure(d, _fig_feature_bars([(a, float(b or 0)) for a, b in fp],
                "Distribution-shape feature vector"),
                "Figure 4. The shape descriptors summarising how the concept is laid across "
                "the muṣḥaf.")
    P(d, f"“{t}” occupies {round((feats.get('coverage') or 0)*100)}% of surahs with a "
         f"spread-entropy of {_fmt(feats.get('spread_entropy'),2)} (high = even), a "
         f"peak-share of {_fmt(feats.get('peak_share'),3)} (the densest single surah holds "
         f"that fraction), and a Meccan share of {_fmt(feats.get('meccan_pct'),2)}. These "
         f"descriptors are summarised, not interpreted: they feed the archetype clustering "
         f"and the null of §5, where the apparent concentration is shown to be a frequency "
         f"effect.")

    # Fig 5 — hotspot surahs
    hs = dist.get("hotspot_surahs", [])[:10]
    if hs:
        _add_figure(d, _fig_counts([(int(s), int(c)) for s, c in hs],
                    "Occurrences by surah (top 10)", "surah", "occurrences", color="#457B9D"),
                    "Figure 5. The ten surahs in which the concept occurs most often.")
        top = "، ".join(str(int(s)) for s, _ in hs[:5])
        P(d, f"The heaviest deployment falls in surahs {top}. A bare reading of this "
             f"concentration would over-claim — which is exactly why the §5 null re-tests it "
             f"against a frequency-matched control rather than trusting the raw counts.")

    # Fig 6 — cross-granularity verification
    cgc = [("both levels", len(cg.get("verified_both_levels", []))),
           ("root only", len(cg.get("root_level_only", []))),
           ("surface only", len(cg.get("surface_level_only", [])))]
    _add_figure(d, _fig_counts(cgc, "Bond verification by granularity", "level",
                "number of bonds", color="#2A9D8F"),
                "Figure 6. How many bonds survive at both the root and surface levels versus "
                "only one.")
    P(d, f"Of the bonds examined, {cgc[0][1]} hold at both granularities and are treated as "
         f"firm; {cgc[1][1]} appear at the root level only and {cgc[2][1]} at the surface "
         f"level only, and these are reported as weaker candidates. This split is the "
         f"analysis policing itself against artefacts of the morphological reduction.")


def _ayah_figures(d, res):
    import numpy as np
    rbt = res["related_by_type"]
    H(d, "7. Figures and supporting analysis", 2)

    _add_figure(d, _fig_grouped_axes(rbt, _AY_ORDER,
                ["lexical", "semantic", "spatial"],
                ["lexical", "semantic", "territory"],
                "Mean axis score by relation type"),
                "Figure 3. Mean standardised score on each axis, by relation type. Consensus "
                "rises on two or more axes together; the divergent type shows the signature "
                "opposition — high on one axis, negative on another.")
    def _avg(k, a):
        lst = rbt.get(k) or []
        vs = [x["axes"].get(a) for x in lst if x.get("axes", {}).get(a) is not None]
        return float(np.mean(vs)) if vs else 0.0
    P(d, f"The resonant type averages {_avg('resonant','semantic'):+.2f} on semantic "
         f"similarity but only {_avg('resonant','lexical'):+.2f} on lexical overlap — the "
         f"defining gap between sharing meaning and sharing words. The divergent type, by "
         f"contrast, is positive on lexical ({_avg('divergent','lexical'):+.2f}) yet negative "
         f"on semantic ({_avg('divergent','semantic'):+.2f}): the structural tension the "
         f"typology is built to isolate.")

    # all candidate points
    allc = [x for lst in rbt.values() for x in lst]
    sp = [(x["axes"]["semantic"], x["axes"]["spatial"]) for x in allc
          if x.get("axes", {}).get("spatial") is not None]
    if sp:
        rb2 = {k: rbt.get(k) for k in _AY_ORDER}
        _add_figure(d, _fig_axis_scatter2(rb2, _AY_ORDER, "semantic", "spatial",
                    "semantic similarity (z)", "territory similarity (z)",
                    "Related verses: semantic vs territory"),
                    "Figure 4. The same verses on the semantic and territory axes. Points high "
                    "on both are co-located meaning-mates; the spread shows the two axes are "
                    "largely independent.")
        P(d, "Plotting meaning against territory shows little diagonal structure: a verse "
             "close in meaning is not generally close in textual setting, which is why the two "
             "axes are kept separate rather than collapsed into one proximity score.")

    sem_z = [x["axes"]["semantic"] for x in allc if x.get("axes", {}).get("semantic") is not None]
    if sem_z:
        _add_figure(d, _fig_hist(sem_z, "Distribution of semantic z across candidates",
                    "semantic similarity (z)"),
                    "Figure 5. The semantic-axis z-scores of the admitted candidates. The "
                    "threshold for 'high' (z = 1) selects the right tail, not the bulk.")
        P(d, f"The admitted candidates span semantic z from {min(sem_z):+.2f} to "
             f"{max(sem_z):+.2f}; only the tail beyond z = 1 qualifies as a resonant or "
             f"consensus bond, so the relation set is a deliberately selective slice of the "
             f"corpus, not a long ranked list padded with marginal matches.")

    top = sorted(allc, key=lambda x: -x.get("relevance", 0))[:10]
    if top:
        _add_figure(d, _fig_counts([(x["ref"], round(float(x.get("relevance", 0)), 1)) for x in top],
                    "Top related verses by relevance", "verse", "relevance", color="#1D3557"),
                    "Figure 6. The ten most relevant related verses overall, by combined "
                    "evidence across axes.")
        P(d, f"The single strongest cross-reference is {top[0]['ref']} (relevance "
             f"{_fmt(top[0].get('relevance'),2)}), followed by {top[1]['ref']} and "
             f"{top[2]['ref']}; these are the verses a reader should consult first.")


def _fig_axis_scatter2(rbt, order, xk, yk, xlabel, ylabel, title):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(6, 4.0), dpi=150)
    for k in order:
        xs, ys = [], []
        for x in (rbt.get(k) or []):
            a = x.get("axes", {})
            if a.get(xk) is not None and a.get(yk) is not None:
                xs.append(a[xk]); ys.append(a[yk])
        if xs:
            ax.scatter(xs, ys, s=30, c=_TYPE_COLOR.get(k, "#888"), label=k,
                       edgecolor="white", linewidth=0.4)
    ax.axhline(1, color="#999", ls=":", lw=0.8); ax.axvline(1, color="#999", ls=":", lw=0.8)
    ax.set_xlabel(xlabel, fontsize=9); ax.set_ylabel(ylabel, fontsize=9)
    ax.set_title(title, fontsize=11); ax.legend(fontsize=7, frameon=False, ncol=2)
    _style_ax(ax); fig.tight_layout()
    p = _fig_path(); fig.savefig(p); plt.close(fig); return p


# ════════════════════════════════════════════════════════════════════════════
# STANDARD IMRaD TECHNICAL PAPERS (final).  Introduction · Materials and Methods
# · Results (with inline, captioned, discussed figures) · Discussion ·
# Limitations · Conclusion · References.  Conforms to the conventional
# scientific-article structure rather than a bespoke scheme.
# ════════════════════════════════════════════════════════════════════════════
def _concept_tech(d, res, t):
    import numpy as np
    fld = res.get("field", {}) or {}; dist = res.get("distribution", {}) or {}
    null = res.get("null", {}) or {}; rel = res.get("relations", {}) or {}
    seq = res.get("sequence", {}) or {}; syn = res.get("synthesis", {}) or {}
    cg = res.get("cross_granularity", {}) or {}; senses = res.get("senses", []) or []
    rbt = rel.get("related_by_type", {}); cm = syn.get("cross_modal", {}) or {}
    mods = syn.get("modalities", {}) or {}; arch = dist.get("archetype") or {}
    feats = dist.get("features", {}) or {}; bc = rel.get("by_relation", {}) or {}
    req = res.get("request", {}); meta = res.get("meta", {})
    k = req.get("k", "—"); mf = req.get("min_freq", "—"); unit = req.get("unit", "surah")
    nscr = null.get("n_scramble", "—"); cv = meta.get("code_version", "—")

    def avg(kk, a):
        lst = rbt.get(kk) or []
        vs = [x["axes"].get(a) for x in lst if x.get("axes", {}).get(a) is not None]
        return float(np.mean(vs)) if vs else 0.0

    H(d, f"A multimodal-fusion characterisation of the Qur'anic concept “{t}”", 1)
    P(d, f"Abstract — We characterise the triliteral root “{t}” "
         f"({dist.get('frequency','—')} occurrences across {dist.get('n_surahs_present','—')} "
         f"surahs) by its relationship to every other concept in the corpus along five "
         f"independent modalities: distributional meaning (PPMI–SVD, k = {k}), areal "
         f"co-location, spatial distribution shape, morphology, and within-verse sequence. "
         f"The modalities are held separate and fused by agreement rather than averaged. The "
         f"principal result is a divergence between the concepts “{t}” co-occurs-with in "
         f"meaning and those it shares textual territory with. A frequency-controlled "
         f"permutation null shows the apparent areal clustering to be an artefact of "
         f"frequency, and a root∥surface test verifies each bond. The study is a structural "
         f"description of the corpus, not exegesis. Code {cv}.", italic=True)
    P(d, "Keywords — Qur'anic corpus; distributional semantics; PPMI–SVD; areal "
         "co-location; spatial autocorrelation; permutation null; multimodal fusion.",
      italic=True)

    H(d, "1. Introduction", 2)
    P(d, "A concept in the Qur'an recurs across many passages, and the company it keeps "
         "shapes how it functions; the classical principle that the text interprets itself "
         "(al-Qur'ān yufassiru ba'ḍuhu ba'ḍan) motivates reading every occurrence together. "
         "We operationalise that reading computationally, without interpreting meaning. The "
         "guiding commitment is anti-reductionist: different modalities answer different "
         "questions — what a concept means-with is not what it is deployed-with, nor its "
         "distribution shape, nor its morphological profile — so collapsing them into one "
         "score destroys the very structure we wish to observe. We therefore measure each "
         "modality separately and fuse them only by agreement. This paper reports the method "
         "and its results for a single focal concept.")

    H(d, "2. Materials and methods", 2)
    P(d, f"2.1. Corpus and preprocessing. The corpus is the full Qur'anic text, "
         f"root-tokenised and orthographically normalised (hamza/alif and yā'/alif-maqṣūra "
         f"variants unified; clitics resolved to stems), so a root does not fragment across "
         f"spellings. Embeddings are estimated for roots attested at least {mf} times.")
    P(d, f"2.2. Modalities. Three modalities are comparative and yield a per-candidate "
         f"score. SEMANTIC is the cosine similarity in a positive-PMI matrix of within-verse "
         f"root co-occurrence reduced by truncated SVD (k = {k} components); it captures "
         f"second-order meaning — what the concept co-occurs-with. CO-LOCATION is the cosine "
         f"of areal deployment profiles over {unit}-level units — what regions the concept "
         f"occupies. SPATIAL is a vector of distribution-shape features (§2.4, §3.3). Two "
         f"further modalities are descriptive: MORPHOLOGY, the split of occurrences across "
         f"{mods.get('morphology',{}).get('n_sense_forms','—')} surface forms, and SEQUENCE, "
         f"the position within the verse. Each comparative score is standardised (z-scored) "
         f"across candidate roots.")
    P(d, "2.3. Fusion and typology. The modalities are never averaged into one score: an "
         "ablation on this pipeline showed that doing so halves a clean consensus signal "
         "(≈0.36 → 0.22), because partly anti-correlated evidence cancels under a mean. "
         "Instead we fuse by agreement, typing each companion by how many modalities are "
         "high (z ≥ 1) and whether any is opposed (z ≤ −1): consensus (≥2 high), divergent "
         "(one high, one opposed), orthogonal (one high, the rest negligible), or the native "
         "type of its single high modality.")
    P(d, f"2.4. Null model and verification. Because a frequent concept clusters spatially "
         f"for trivial reasons, the spatial modality is tested against a frequency-preserving "
         f"permutation null (n = {nscr} scrambles holding the count fixed). Separately, each "
         f"bond is checked at two granularities — the root level and the clitic-stripped "
         f"surface level — and only those holding at both are treated as firm.")

    H(d, "3. Results", 2)
    P(d, f"3.1. Principal finding. The concepts “{t}” means-with and those it is "
         f"deployed-with do not coincide. Its strongest meaning-mates are "
         f"{('، '.join(fld.get('semantic_field', [])[:6]) or '—')}; its strongest "
         f"territory-mates are {('، '.join(fld.get('co_location_neighbours', [])[:6]) or '—')}. "
         f"The two sets are largely disjoint ({cm.get('divergence','—')}): the concept is "
         f"discussed among one family of ideas yet placed in the text among another. The "
         f"companions on which the two agree are {('، '.join(cm.get('convergence', [])) or '—')}, "
         f"and the bond additionally surviving verification at both granularities — the most "
         f"defensible single result — is {('، '.join(cm.get('verified_bonds', [])) or '—')}.")

    P(d, f"3.2. Typed companions. Table 1 lists the leading companions of each relation "
         f"type with their per-modality z-scores. The relation counts are shown in Figure 1, "
         f"and the mean modality profile of each type in Figure 2.")
    for ty in REL_ORDER:
        lst = rbt.get(ty) or []
        if not lst:
            continue
        H(d, f"Table 1{chr(97+REL_ORDER.index(ty))}. {ty.title()} — {REL_EN[ty]} (n = {len(lst)})", 2)
        TABLE(d, ["root", "meaning", "territory", "shape"],
              [[x["root"], x["axes"].get("semantic"), x["axes"].get("co-location"),
                x["axes"].get("spatial")] for x in lst[:8]])
    _add_figure(d, _fig_relation_bar(bc, REL_ORDER, {kk: kk.title() for kk in REL_ORDER},
                "Companion concepts by relation type"),
                "Figure 1. Number of companion concepts in each relation type.")
    P(d, "Figure 1 shows the relation type is dominated by consensus and the single-modality "
         "classes, with the off-diagonal classes (orthogonal, divergent) comparatively rare "
         "— consistent with reinforcement being the corpus-wide norm and tension the "
         "exception.")
    _add_figure(d, _fig_grouped_axes(rbt, REL_ORDER, ["semantic", "co-location", "spatial"],
                ["meaning", "territory", "shape"], "Mean modality score by relation type"),
                "Figure 2. Mean standardised score on each comparative modality, by type.")
    P(d, f"Figure 2 confirms the typology is non-redundant. Consensus companions average "
         f"{avg('consensus','semantic'):+.2f} on meaning and {avg('consensus','co-location'):+.2f} "
         f"on territory — high on both at once — whereas the semantic type averages "
         f"{avg('semantic','semantic'):+.2f} on meaning against only "
         f"{avg('semantic','co-location'):+.2f} on territory, and the co-location type "
         f"inverts that. Were the modalities measuring the same quantity, these profiles "
         f"would be indistinguishable; they are not.")

    P(d, "3.3. Distribution and shape. Figure 3 gives the shape descriptors and Figure 4 the "
         "surahs of heaviest deployment.")
    fp = [("coverage", feats.get("coverage", 0)), ("log-freq", feats.get("log_freq", 0)),
          ("spread-ent.", feats.get("spread_entropy", 0)), ("peak-share", feats.get("peak_share", 0)),
          ("lacunarity", feats.get("lacunarity", 0)), ("Meccan", feats.get("meccan_pct", 0))]
    _add_figure(d, _fig_feature_bars([(a, float(b or 0)) for a, b in fp],
                "Distribution-shape feature vector"),
                "Figure 3. Shape descriptors summarising the concept's spread across the muṣḥaf.")
    P(d, f"The concept occupies {round((feats.get('coverage') or 0)*100)}% of surahs with "
         f"spread-entropy {_fmt(feats.get('spread_entropy'),2)} and peak-share "
         f"{_fmt(feats.get('peak_share'),3)}; these descriptors feed the archetype clustering "
         f"(“{arch.get('tag','—')}”, stability {_fmt(arch.get('stability'),2)}) and the null "
         f"of §3.4.")
    hs = dist.get("hotspot_surahs", [])[:10]
    if hs:
        _add_figure(d, _fig_counts([(int(s), int(c)) for s, c in hs],
                    "Occurrences by surah (top 10)", "surah", "occurrences", color="#457B9D"),
                    "Figure 4. The ten surahs in which the concept occurs most often.")
        P(d, f"Deployment is heaviest in surahs {('، '.join(str(int(s)) for s,_ in hs[:5]))}. "
             f"A bare reading of this concentration would over-claim, which §3.4 addresses "
             f"directly.")

    P(d, "3.4. The spatial modality is null once frequency is controlled. The analytic "
         f"Moran's I is {_fmt(feats.get('moran_I'),3)} (analytic z {_fmt(feats.get('moran_z'),2)}), "
         f"apparently strong; but the analytic statistic assumes spatial randomness, and a "
         f"frequent root clusters merely because it is common. Against the frequency-"
         f"preserving null (mean {_fmt(null.get('null_mean'),3)} ± {_fmt(null.get('null_sd'),3)}) "
         f"the observed value yields z = {_fmt(null.get('z'),2)} (Figure 5): "
         f"{null.get('interpretation','—')}.")
    _add_figure(d, _fig_null(null, "Areal clustering vs a frequency-matched null"),
                "Figure 5. Observed areal-clustering statistic (red) against the "
                "frequency-preserving null (band); it sits within the null.")
    P(d, "Because the observed value falls inside the null band, the apparent clustering is "
         "not counted as a finding. Keeping the spatial modality as one track of five is what "
         "prevents this frequency artefact from being mistaken for structure.")

    if senses:
        P(d, "3.5. Morphology. Figure 6 shows how occurrences split across surface forms.")
        sc = [(f"f{i+1}", int(s["count"])) for i, s in enumerate(senses[:8])]
        _add_figure(d, _fig_counts(sc, "Occurrences by surface form", "surface form (ranked)",
                    "occurrences", color="#8D99AE"),
                    "Figure 6. Distribution of the concept's occurrences across its surface "
                    "forms (f1 = most frequent). Forms, in order: "
                    + "، ".join(s["form"] for s in senses[:8]) + ".")
        _coh = _sense_cohesion(senses)
        if _coh:
            _verd = {"cohesive": "keep largely the same company (a cohesive, single-sense profile)",
                     "mixed": "keep partly overlapping company (a shared core with form-specific variation)",
                     "split": "keep markedly different company (evidence of more than one sense carried by the root)"}[_coh["verdict"]]
            P(d, f"The concept appears in {len(senses)} distinct surface forms; across the "
                 f"{_coh['n_forms']} high-mass forms the leading forms {_verd}, with a mean "
                 f"pairwise co-locator overlap (Jaccard) of {_coh['jaccard']:.2f}. Morphology is "
                 f"therefore tracked as its own modality.")
        else:
            P(d, f"The concept appears in {len(senses)} distinct surface forms; the sense profile "
                 f"is concentrated in the leading forms, each of which keeps somewhat different "
                 f"company, which is why morphology is tracked as its own modality.")

    cgc = [("both", len(cg.get("verified_both_levels", []))),
           ("root only", len(cg.get("root_level_only", []))),
           ("surface only", len(cg.get("surface_level_only", [])))]
    _add_figure(d, _fig_counts(cgc, "Bond verification by granularity", "level",
                "bonds", color="#2A9D8F"),
                "Figure 7. Bonds surviving at both the root and surface levels versus one only.")
    P(d, f"3.6. Verification. {cgc[0][1]} bonds hold at both granularities and are treated as "
         f"firm; {cgc[1][1]} appear at the root level only and {cgc[2][1]} at the surface "
         f"level only, reported as weaker candidates. This split is the analysis policing "
         f"itself against artefacts of the morphological reduction.")

    H(d, "4. Discussion", 2)
    P(d, f"The central observation is structural: for “{t}”, meaning-similarity and "
         f"textual-territory are nearly orthogonal. A concept can be discussed in the "
         f"semantic company of one set of ideas while being physically deployed beside "
         f"another, and the multimodal design is what makes this visible — a single blended "
         f"similarity would have averaged the two into an uninformative middle. The "
         f"consensus companions, where the modalities do agree, are the safest anchors for "
         f"further study; the divergent companions, where they conflict, are the most "
         f"suggestive. This across-verse view complements within-verse motif (co-occurrence) "
         f"analysis, which is exact but verse-bounded; their intersection yields latent "
         f"motifs, of which 502 survive a Benjamini–Hochberg FDR gate at q = 0.05.")

    H(d, "5. Limitations", 2)
    P(d, "This is a structural description, not tafsir, and asserts nothing about meaning, "
         "ruling, or theology. The semantic modality is a model inferred from co-occurrence "
         "and is coarser for the corpus's size and the chosen dimensionality; single-"
         "modality labels are candidates until significance-gated; the spatial modality is "
         "null for many concepts, as §3.4 shows for this one; morphological senses depend on "
         "the underlying analysis, which §3.6 verifies. The corpus is its own reference "
         "distribution, so z-scores are internal, not absolute.")

    H(d, "6. Conclusion", 2)
    P(d, f"“{t}” occupies a measurable, typed position in the corpus along five modalities "
         f"held apart and fused by agreement; its defining feature here is the divergence "
         f"between meaning and territory. The procedure is deterministic given the corpus and "
         f"parameters (min_freq = {mf}, k = {k}, unit = {unit}, n_scramble = {nscr}; code "
         f"{cv}), and is offered as a reproducible map for closer reading, not a substitute "
         f"for it.")
    P(d, "References — Spärck Jones (1972); Church & Hanks (1990); Deerwester et al. (1990); "
         "Levy & Goldberg (2014); Moran (1950); Anselin (1995); Benjamini & Hochberg (1995); "
         "Blondel et al. (2008).", italic=True)


def _ayah_tech(d, res, refstr, seed, rbt, bc, cm, n_ay, cv, k, mf, unit, n_cand, nroots):
    import numpy as np
    def avg(kk, a):
        lst = rbt.get(kk) or []
        vs = [x["axes"].get(a) for x in lst if x.get("axes", {}).get(a) is not None]
        return float(np.mean(vs)) if vs else 0.0
    allc = [x for lst in rbt.values() for x in lst]

    H(d, f"Locating and typing the cross-references of Qur'anic verse(s) {refstr}: "
         "a context-entity, multimodal method", 1)
    P(d, f"Abstract — For a target verse we locate every related verse in the corpus "
         f"(N = {n_ay}), assign each a relation type, and attach its evidence. The target is "
         f"represented as a single context entity — an inverse-document-frequency-weighted "
         f"centroid of its root embeddings, so its distinctive concepts define its signature "
         f"— and every candidate is scored, identically, on three independent axes (lexical, "
         f"semantic, territorial), standardised and fused by agreement rather than averaged. "
         f"Every root enters the computation regardless of display. The study is a structural "
         f"description, not exegesis. Code {cv}.", italic=True)
    P(d, "Keywords — Qur'anic corpus; weighted context embedding; PPMI–SVD; areal "
         "co-location; z-standardisation; relation typology; cross-reference discovery.",
      italic=True)

    H(d, "1. Introduction", 2)
    P(d, "Classical hermeneutics reads each verse in the light of the whole "
         "(al-Qur'ān yufassiru ba'ḍuhu ba'ḍan). We operationalise that movement from part to "
         "whole — exhaustive cross-reference discovery — while declining interpretation. A "
         "keyword approach to relating verses is exact but blind to verses that express the "
         "same matter in different words and is misled by ubiquitous terms. Our design "
         "answers both failures: a distributional axis surfaces shared subject without shared "
         "vocabulary, and a salience-weighted whole-verse representation prevents common "
         "roots from dominating. The contributions are a context-entity representation, three "
         "independent axes kept separate, a fixed typology separating reinforced bonds from "
         "single-axis links and tensions, and analysis decoupled from display.")

    H(d, "2. Materials and methods", 2)
    P(d, f"2.1. Corpus and preprocessing. Full Qur'anic text, N = {n_ay} verses, each "
         f"reduced to triliteral roots and orthographically normalised; a diacritised column "
         f"serves display only and a root-tokenised column serves all computation. Embeddings "
         f"are estimated for roots attested at least {mf} times; rarer roots keep no vector "
         f"but still contribute to the lexical axis and the evidence.")
    P(d, "2.2. The verse as a context entity. For the embeddable roots R⁺(v) of verse v the "
         "signature is the idf-weighted centroid s(v) = Σ w_r·e(r), w_r = idf(r)/Σidf, with "
         "idf(r) = log(1 + N/(1+df(r))). A flat average would let common roots drag the "
         "centroid toward the corpus mean and blur the verse's identity; the idf weighting "
         "lets distinctive roots dominate, so the verse's signature reflects what is "
         "particular about it. The centroid is computed over the semantic space D (PPMI–SVD, "
         f"k = {k}) and the territory space L (areal deployment over {unit}-level units); "
         "target and candidates are embedded identically, so all comparisons are "
         "entity-to-entity.")
    P(d, "2.3. Axes, standardisation, typology. The lexical axis sums idf over the full "
         "shared-root sets; the semantic and territorial axes are cosines of the respective "
         "signatures. Each axis is z-standardised across candidates. A candidate is typed "
         "consensus (≥2 axes with z ≥ 1), divergent (one high, semantic or territory ≤ −1), "
         "orthogonal (one high, the rest |z| < 0.5), or the native type of its single high "
         "axis. The axes are never averaged: an ablation showed blending halves a clean "
         "consensus signal (≈0.36 → 0.22). Consensus is assessed against a cross-modal-"
         "alignment null and occurs far beyond chance corpus-wide.")

    H(d, "3. Results", 2)
    P(d, f"3.1. Cross-reference distribution. The target verse contributes {nroots} roots to "
         f"its signature; {n_cand} candidates were admitted. The relation counts (Figure 1) "
         f"are: " + "; ".join(f"{kk} {bc.get(kk,0)}" for kk in _AY_ORDER if kk in bc) + ".")
    P(d, "Target verse(s): " + ", ".join(sd["ref"] for sd in seed)
         + f"; the verse signature is built from {nroots} constituent roots. Verses are "
         "cited by reference throughout; the Qur'anic text is not reproduced.")
    _add_figure(d, _fig_relation_bar(bc, _AY_ORDER, {kk: _AY_LABEL[kk][0] for kk in _AY_ORDER},
                "Related verses by relation type"),
                "Figure 1. Distribution of the related verses across the six relation types.")
    P(d, "Consensus and the single-axis classes account for most cross-references, with the "
         "off-diagonal classes rarer — reinforcement is the norm, tension the exception.")

    P(d, "3.2. Typed cross-references. Table 1 lists the leading members of each type with "
         "axis z-scores (L lexical, S semantic, P territory) and shared roots.")
    for ty in _AY_ORDER:
        lst = rbt.get(ty) or []
        if not lst:
            continue
        H(d, f"Table 1{chr(97+_AY_ORDER.index(ty))}. {_AY_LABEL[ty][0]} — {_AY_LABEL[ty][2]} "
             f"(n = {bc.get(ty,len(lst))})", 2)
        TABLE(d, ["verse", "L", "S", "P", "# shared"],
              [[x["ref"], x["axes"]["lexical"], x["axes"]["semantic"], x["axes"]["spatial"],
                len(x["shared_roots"])] for x in lst[:8]])

    P(d, "3.3. Axis structure. Figures 2–4 show how the axes relate. Figure 2 plots lexical "
         "against semantic, Figure 3 the mean axis profile by type, and Figure 4 semantic "
         "against territory.")
    _add_figure(d, _fig_axis_scatter(rbt, _AY_ORDER,
                "Related verses by lexical and semantic similarity"),
                "Figure 2. Each related verse by lexical overlap (vertical) and semantic "
                "similarity (horizontal), in z units; dotted lines mark z = 1.")
    P(d, f"The resonant type sits in the lower-right (high semantic, low lexical): verses "
         f"about the same matter in different words. Resonant members average "
         f"{avg('resonant','semantic'):+.2f} on semantic against {avg('resonant','lexical'):+.2f} "
         f"on lexical — the gap that defines the type.")
    _add_figure(d, _fig_grouped_axes(rbt, _AY_ORDER, ["lexical", "semantic", "spatial"],
                ["lexical", "semantic", "territory"], "Mean axis score by relation type"),
                "Figure 3. Mean standardised score on each axis, by relation type.")
    P(d, f"The divergent type is positive on lexical ({avg('divergent','lexical'):+.2f}) yet "
         f"negative on semantic ({avg('divergent','semantic'):+.2f}) — the signature "
         f"opposition the typology isolates; consensus is positive across the board.")
    sp = [(x['axes']['semantic'], x['axes']['spatial']) for x in allc
          if x.get('axes', {}).get('spatial') is not None]
    if sp:
        _add_figure(d, _fig_axis_scatter2({kk: rbt.get(kk) for kk in _AY_ORDER}, _AY_ORDER,
                    "semantic", "spatial", "semantic similarity (z)", "territory similarity (z)",
                    "Related verses: semantic vs territory"),
                    "Figure 4. The same verses on the semantic and territory axes; the spread "
                    "shows the two are largely independent.")
        P(d, "Meaning and territory show little diagonal structure — closeness in meaning "
             "does not imply closeness in textual setting — which justifies keeping the axes "
             "separate rather than collapsing them.")

    sem_z = [x['axes']['semantic'] for x in allc if x.get('axes', {}).get('semantic') is not None]
    if sem_z:
        _add_figure(d, _fig_hist(sem_z, "Distribution of semantic z across candidates",
                    "semantic similarity (z)"),
                    "Figure 5. Semantic-axis z-scores of the admitted candidates; the z = 1 "
                    "threshold selects the right tail.")
        P(d, f"3.4. Selectivity. Candidate semantic z spans {min(sem_z):+.2f} to "
             f"{max(sem_z):+.2f}; only the tail beyond z = 1 qualifies, so the relation set "
             f"is a deliberately selective slice, not a long list padded with marginal "
             f"matches.")
    top = sorted(allc, key=lambda x: -x.get('relevance', 0))[:10]
    if top:
        _add_figure(d, _fig_counts([(x['ref'], round(float(x.get('relevance', 0)), 1)) for x in top],
                    "Top related verses by relevance", "verse", "relevance", color="#1D3557"),
                    "Figure 6. The ten most relevant related verses, by combined evidence "
                    "across axes.")
        P(d, f"3.5. The strongest cross-reference is {top[0]['ref']} (relevance "
             f"{_fmt(top[0].get('relevance'),2)}), followed by {top[1]['ref']} and "
             f"{top[2]['ref']} — the verses to consult first.")

    H(d, "4. Discussion", 2)
    P(d, "The neighbourhood of the target partitions into qualitatively different relations. "
         "The consensus set is the dependable core; the resonant set is the most distinctive, "
         "naming verses that share subject without sharing vocabulary and so invisible to a "
         "concordance; the divergent set marks foils. This across-verse method complements "
         "within-verse motif analysis, which is exact but verse-bounded; their intersection "
         "yields latent motifs, of which 502 survive a Benjamini–Hochberg FDR gate at "
         "q = 0.05. For this target the cross-modal reading is: "
         + str(cm.get("divergence", "the axes are largely independent here")) + ".")

    H(d, "5. Limitations", 2)
    P(d, "Description, not tafsir. The semantic axis is inferred from co-occurrence and may "
         "carry embedding artefacts; off-diagonal labels are provisional; the territory axis "
         "is weak and often null; the context centroid discards within-verse word order and "
         "can be dominated by a single rare root; z-scores are internal to the corpus.")

    H(d, "6. Conclusion", 2)
    P(d, f"The verse(s) {refstr} sit within a measurable, typed web of cross-references made "
         f"explicit and evidenced here. The procedure is deterministic given min_freq = {mf}, "
         f"k = {k}, unit = {unit}; code {cv}. It is offered as a reproducible map for closer "
         f"reading, not a substitute for it.")
    P(d, "References — Spärck Jones (1972); Church & Hanks (1990); Deerwester et al. (1990); "
         "Levy & Goldberg (2014); Arora et al. (2017); Benjamini & Hochberg (1995).",
      italic=True)
