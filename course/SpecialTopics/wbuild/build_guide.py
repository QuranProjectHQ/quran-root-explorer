# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
OUT="/sessions/stoic-serene-wozniak/mnt/Downloads/RootCourse/SpecialTopics/SpecialTopics_App_and_Plot_Guide.docx"
ACCENT=RGBColor(0x1F,0x4E,0x79); TEAL=RGBColor(0x0E,0x6D,0x63); GREY=RGBColor(0x55,0x55,0x55)
def setcs(r):
    rPr=r._r.get_or_add_rPr(); rf=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
d=Document(); st=d.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(11)
rf=st.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
def P(text,size=11,bold=False,color=None,after=4,before=0,style=None,italic=False):
    p=d.add_paragraph(style=style); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; r.italic=italic; setcs(r)
    if color: r.font.color.rgb=color
    return p
def B(text,size=10.5): return P(text,size,style="List Bullet",after=2)
def N(text,size=10.5): return P(text,size,style="List Number",after=2)
def H(text): return P(text,12.5,True,TEAL,after=3,before=10)

P("Special Topics - App & Plot Guide",16,True,ACCENT,after=2)
P("Using the Quran Root Explorer app to drive the 27 Special Topics, and how to read every figure",11.5,False,GREY,after=8,italic=True)
P("The Special Topics decks are companions to the live app. Every count and chart in them recomputes from Book6 - the same corpus the app reads - so each topic can be reproduced live. Below: a page-by-page tour of the app, the tasks to run during a session, how to read each chart type with its honest caveats, and a map from chart to topic. Honest spine throughout: the app and the figures PRESENT the structure; the theological reading is labelled and never adjudicated.",11,after=8)

H("1. The app, page by page")
P("ROOT TOOLS - explore a single root or a small set:",11,True,after=2)
B("Per-Root Profile - search a root; read its frequency, surface forms, and co-occurring roots. The workhorse for any concept topic (e.g. غفر in W10 Ghafr, نفق in W08 Hypocrite).")
B("Network - the co-occurrence graph; see which roots cluster with your root (supports W05 Name Pairs, W07 Signs).")
B("Motifs - recurring multi-root patterns (supports the hypocrite trio, the wealth+children pair).")
B("Ayah Browser - pull any verse by address with its vocalized text (supports the Challenges verses, the equity verses 4:7/4:11/2:282).")
B("Compare & Heatmaps - co-occurrence and lift between two roots side by side (the engine behind W05 Name-Pair lift).")
B("Morphology - surface-form breakdown of a root - the §14a concept-verification view (supports W03 mukhlis, act-vs-state, light/darkness number).")
P("STATISTICS & INTERPRET - quantify and sanity-check:",11,True,after=2,before=6)
B("Statistics - corpus totals, frequency tables, distributions (supports W02 units, divine-name frequencies).")
B("Surface Divergence - how a root's surface forms split its senses; the guard against concept-mixing (supports din/islam/quran, ghafr forms).")
B("Topic Map / My Topics / Topic Modeling - data-driven concept fields (supports the concept-spectrum framing).")
B("Calibration / Practical Lens / Interpret - turn a raw count into a calibrated, honestly-framed reading.")
P("TWO BOOKS - the structural battery:",11,True,after=2,before=6)
B("Disjoint Letters, Signal, Biology - the live permutation tests on letters, spectra and composition.")
B("FDR Summary - the Benjamini-Hochberg dashboard (the source of the '6 of 8 survive' figure in the Challenges deck).")
P("HELP & EXPORT - Help (orientation), Export (download tables), Usage (session stats).",11,True,after=2,before=6)

H("2. Live app tasks (run during a session)")
B("Search a root in Per-Root Profile (e.g. غفر) -> read frequency, forms, neighbours (supports the W10 forgiveness topics).")
B("Open Compare & Heatmaps for two Divine Names -> read shared count AND lift (supports W05 Name Pairs - count vs lift).")
B("Open Morphology / Surface Divergence on a polysemous root -> watch the senses separate (supports mukhlis, act-vs-state, sword sense-filtering).")
B("Pull a challenge verse in Ayah Browser by address (17:88, 2:24, 67:3) -> read it verbatim (supports the Challenges deck).")
B("Open the Two Books FDR Summary -> show the live permutation p-values behind the consistency dare (supports Challenges, FDR).")
B("Open Statistics -> corpus totals and length distributions (supports W02 sura/ayah units, mushaf order).")

H("3. Screenshot capture list")
N("sshot-1: a root's Per-Root Profile (frequency + forms + neighbours).")
N("sshot-2: Compare & Heatmaps - two Names, shared count and lift.")
N("sshot-3: Morphology / Surface Divergence - one root's surface forms split by sense.")
N("sshot-4: Ayah Browser - a challenge verse by address, vocalized.")
N("sshot-5: FDR Summary - the Benjamini-Hochberg dashboard.")
P("Tip: capture screenshots at 150 dpi or higher; crop to the chart so it reads from the back of the room.",10.5,True,after=4,before=4)

H("4. Reading the figures - the plot guide")
P("Every Special-Topics figure is one of a small set of chart types. For each: what it shows, how to read it, and the caveat.",11,after=4)
def fig(name,shows,read,caveat):
    P(name,11,True,ACCENT,after=1,before=4)
    B("Shows: "+shows); B("Read: "+read); B("Caveat: "+caveat,10.5)
fig("Frequency bar (ayat per root)","how many ayat contain a root/concept.","longer bar = more ayat; compare fields side by side.","raw root counts can fuse senses - figures use surface forms where a concept must be kept apart (the light topic counts نور as a noun, not نار fire).")
fig("Co-occurrence / shared-ayah bar","how often two roots appear in the same ayah.","taller = the two ideas travel together more.","a big shared count can be mere frequency (two common words overlap by chance) - read it with lift.")
fig("Count-vs-lift scatter","support (x = shared count) against lift (y = times above chance).","top-right = strong AND well-supported; high-left = tight but thin; low-right = frequent but weak.","high lift on a low count is fragile - a leave-one-out would wobble it (W05 gold-standard logic).")
fig("Surface-form split bar","the distinct forms of one root, kept apart.","each bar is a separate concept (e.g. deeds vs a righteous person vs a reformer).","merging the root into one tally would fuse different concepts - the split is the point (§14a).")
fig("Sura distribution bar","a root's occurrences across the 114 suras.","spikes show where a concept lives (e.g. inheritance clusters in sura 4).","raw per-sura counts track sura length - long suras hold more of everything.")
fig("Meccan/Medinan timeline (stacked)","the period split of a concept or verse-set.","teal = Meccan, amber = Medinan; reads whether a theme is early, late, or sustained.","revelation rank is narrated (per sura), not computed - an approximation.")
fig("Length histogram","the distribution of sura or ayah sizes.","the bulk vs the long tail; a log y-axis tames extreme outliers (the 84-token verse 2:282).","length describes the units; it does not define them (they are marked).")
fig("Escalation / FDR q-value bar","a ladder (units demanded) or test q-values against a threshold line.","follow the dashed line - bars past it 'survive' (e.g. 6 of 8 below 5% FDR).","surviving a test means reproducible/non-random, NOT proof of a theological claim.")

H("5. Plot to topic map")
B("Frequency bars - nearly every topic (bashir/nadhir, sword, divine names, equity vocabulary).")
B("Count-vs-lift scatter - W05 Name Pairs, W08 Hypocrite, W04 Wealth & Children.")
B("Co-occurrence bars - W03 din/islam/quran, W07 Signs, W10 Ghafr, intercession.")
B("Surface-form split - W03 mukhlis & act-vs-state, W04 light/darkness number, ghafr forms.")
B("Sura distribution - every topic (where the concept lives in the corpus).")
B("Meccan/Medinan timeline - Challenges, W01 who-addressed, act-vs-state, light/darkness.")
B("Length histograms - W02 sura/ayah units, W02 mushaf order.")
B("FDR q-value bar - the Challenges consistency dare and the Two Books FDR Summary.")

P("Reproducibility: every number shown in the app or the decks recomputes from Book6 (the app reads it live; the decks via the shared wk.py kernel, fixed seed). Roots are normalized (Persian/Arabic letter variants folded) so counts are stable; surface forms are used wherever a concept must be separated from its root family.",10.5,False,GREY,after=4,before=8,italic=True)
d.save(OUT); print("saved",OUT)
