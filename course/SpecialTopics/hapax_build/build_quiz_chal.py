# -*- coding: utf-8 -*-
import os, string
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55)
OUTDIR="/sessions/stoic-serene-wozniak/mnt/Downloads/RootCourse/SpecialTopics"
def setcs(r):
    rPr=r._r.get_or_add_rPr(); rf=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
def newdoc():
    d=Document(); st=d.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(11)
    rf=st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
    return d
def P(d,segs,size=11,bold=False,color=None,after=4,before=0):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if isinstance(segs,str): segs=[(segs,bold)]
    for t,b in segs:
        r=p.add_run(t); r.font.size=Pt(size); r.font.bold=b; setcs(r)
        if color: r.font.color.rgb=color
    return p
QQ=[
("1.  al-Tahaddi refers to the Qur'an's:","challenges / dares to its audience",["names of God","rules of recitation","narrative parables"],"al-tahaddi = the open, falsifiable challenges the text issues."),
("2.  The literary challenge escalates DOWN in this order:","the whole -> ten suras -> one sura -> a single statement",["one word -> one verse -> one sura","ten suras -> the whole","it never changes"],"17:88 (whole) -> 11:13 (ten) -> 2:23/10:38 (one sura) -> 52:34 (a statement)."),
("3.  Which verse PREDICTS the literary challenge will never be met?","2:24  ('and you will never do it')",["17:88","67:3","4:82"],"2:24: 'if you do not - and you will never do it -'. A falsifiable prediction."),
("4.  The internal-consistency challenge (4:82) invites the reader to:","find a contradiction across the whole corpus",["count the letters","memorize the text","find the longest sura"],"4:82: had it been from other than God they would find much discrepancy."),
("5.  The cosmic challenge (67:3-4) dares the reader to:","find a flaw or rupture in the creation",["predict the weather","fast for a month","build a mosque"],"67:3-4: do you see any disparity (tafawut) or rupture (futur)? Look again."),
("6.  The word tafawut ('disparity', root f-w-t) occurs in how many ayat?","5",["148","486","19"],"tafawut is a near-unique, precise term - only 5 ayat in Book6."),
("7.  Of the 8 representative structural tests, how many survive a 5% FDR?","6 of 8",["all 8","2 of 8","none"],"6/8 survive (q <= 0.0067); the two borderline tests correctly drop out."),
("8.  The dare-verb 'bring!' (root a-t-y) appears in how many ayat?","486",["5","148","16"],"the imperative 'bring the like' draws on a high-frequency core verb (486 ayat)."),
("9.  The root for 'the like' (mithl) occurs in 148 ayat; the literary challenges number:","5 of them",["all 148","148","none"],"exactly 5 of the 148 mithl-ayat carry the literary dare - a precise subset."),
("10.  The seven challenge suras range in length from:","al-Mulk (30 ayat) to al-Baqara (286)",["all under 30 ayat","all over 200 ayat","exactly 100 ayat each"],"the dare recurs across short Meccan and long Medinan suras alike."),
("11.  A structural test SURVIVING the FDR correction means the result is:","reproducible and non-random - NOT proof of a miracle",["proof of divine origin","a hidden numeric code","statistically meaningless"],"FDR controls multiplicity; survival shows non-randomness, not a theological proof."),
("12.  How does this topic present the challenges?","quotes each by address and counts the vocabulary, without declaring them met or unmet",["declares them all met","declares them all unmet","ignores the addresses"],"we present and count from Book6; 'met?' is a theological reading we do not adjudicate."),
("13.  How many distinct challenge ayat are collected, and in how many kinds?","9 ayat in 4 kinds (literary, prediction, consistency, cosmic)",["1 ayah, 1 kind","100 ayat, 10 kinds","9 ayat, 1 kind"],"5 literary + 1 prediction (2:24) + 1 consistency (4:82) + 2 cosmic (67:3-4) = 9 ayat, 4 kinds."),
]
d=newdoc(); P(d,[("Special Topic - The Challenges of the Qur'an (al-Tahaddi) - Quiz",True)],size=17,color=ACCENT,after=2)
P(d,"13 questions - ~15 minutes - choose the single best answer. Every value reproducible from Book6.",size=9.5,color=GREY,after=8)
KEY=[]
for qi,(stem,correct,distr,expl) in enumerate(QQ):
    pos=qi%4; opts=list(distr); opts.insert(pos,correct)
    P(d,[(stem,True)],size=10.5,after=2,before=6)
    for i,o in enumerate(opts): P(d,string.ascii_uppercase[i]+")  "+o,size=10,after=1)
    KEY.append((qi+1,string.ascii_uppercase[pos],expl))
d.save(OUTDIR+"/SpecialTopic_Challenges_Quiz.docx")
d=newdoc(); P(d,[("Special Topic - The Challenges of the Qur'an - Quiz Answer Key",True)],size=17,color=ACCENT,after=2)
P(d,"One point each, 13 total. Every value reproducible from Book6.",size=9.5,color=GREY,after=8)
for n,a,ex in KEY: P(d,[(str(n)+".  "+a+"  ",True),("- "+ex,False)],size=10,after=2)
d.save(OUTDIR+"/SpecialTopic_Challenges_Quiz_Answer_Key.docx")
print("quiz+key saved | answer letters:",[a for _,a,_ in KEY])
