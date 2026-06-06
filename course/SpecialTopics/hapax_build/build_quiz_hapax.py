# -*- coding: utf-8 -*-
import os, string
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55)
OUTDIR="/sessions/stoic-serene-wozniak/mnt/Downloads/RootCourse/SpecialTopics"
def setcs(r):
    rPr=r._r.get_or_add_rPr(); rf=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
def newdoc():
    d=Document(); st=d.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(11)
    rf=st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
    return d
def P(d,segs,size=11,bold=False,color=None,after=4,before=0):
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if isinstance(segs,str): segs=[(segs,bold)]
    for t,b in segs:
        r=p.add_run(t); r.font.size=Pt(size); r.font.bold=b; setcs(r)
        if color: r.font.color.rgb=color
    return p
QQ=[
("1.  A hapax legomenon is a word that occurs:","exactly once in the corpus",["most often","never","in every sūra"],"hapax = a single occurrence in the whole corpus."),
("2.  How many ROOTS occur exactly once in the Qur'an (Book6)?","408",["1,701","3,027","114"],"408 of 1,701 distinct roots are once-only."),
("3.  How many surface FORMS occur exactly once?","3,027",["408","6,821","167"],"3,027 of 6,821 distinct forms are hapax."),
("4.  صمد (al-Ikhlāṣ 112:2) is an example of a:","once-only ROOT",["once-only form of a common root","frequent root","letter"],"the root ص-م-د occurs once in the whole Qur'an."),
("5.  كوثر (108:1) is a once-only FORM whose ROOT (كثر) occurs:","167 times",["once","never","408 times"],"كوثر is a unique form; its root كثر is common (167 āyahs)."),
("6.  Raw counts of hapax roots are highest in long sūras (al-Baqara) because:","longer sūras contain more of everything",["they are holier","they are Meccan","of a hidden code"],"raw count tracks length — a size confound."),
("7.  After normalizing by sūra size, the RICHEST sūra in hapax is:","sūra 108 al-Kawthar (~29%)",["al-Baqara","al-Fātiḥa","Ṭā-Hā"],"density flips the ranking: short juz-30 sūras lead; 108 tops it."),
("8.  In al-Baqara 2:61, بقل/قثّاء/فوم/عدس/بصل (herbs/cucumber/garlic/lentils/onions) are hapax because:","they name specific items mentioned only once",["they are miraculous","they are divine names","they are mis-spelled"],"rarity here is lexical specificity, not mystery."),
("9.  Roughly what share of distinct ROOTS are once-only?","about 24% (1 in 4)",["about 1%","about 90%","exactly half"],"408 of 1,701 ≈ 24%."),
("10.  A once-only ROOT vs a once-only FORM differ in that:","a root names a unique concept; a form is a unique wording of a possibly-common idea",["they are identical","forms are rarer than letters","roots are always longer"],"root-hapax = unique concept; form-hapax = unique expression."),
("11.  The honest reading of a hapax is:","a count of one with its address — not a hidden code",["proof of a miracle","a numerical cipher","evidence of authorship"],"we report count + address and decline to over-read."),
("12.  هل ترى من فطور-type once-only words cluster in:","the short sūras of juz 30",["only al-Baqara","the middle sūras","no particular place"],"juz-30 short sūras are densest in hapax."),
("13.  The full hapax inventory is provided as:","Hapax_roots_full.csv and Hapax_forms_full.csv",["a single PDF","nothing — on request only","an image"],"both complete tables (408 roots, 3,027 forms) are exported."),
]
d=newdoc(); P(d,[("Special Topic — Hapax (Once-Only) · Quiz",True)],size=18,color=ACCENT,after=2)
P(d,"13 questions · ~15 minutes · choose the single best answer. Every value reproducible from Book6.",size=9.5,color=GREY,after=8)
KEY=[]
for qi,(stem,correct,distr,expl) in enumerate(QQ):
    pos=qi%4; opts=list(distr); opts.insert(pos,correct)
    P(d,[(stem,True)],size=10.5,after=2,before=6)
    for i,o in enumerate(opts): P(d,f"{string.ascii_uppercase[i]})  {o}",size=10,after=1)
    KEY.append((qi+1,string.ascii_uppercase[pos],expl))
d.save(OUTDIR+"/SpecialTopic_Hapax_Quiz.docx")
d=newdoc(); P(d,[("Special Topic — Hapax · Quiz Answer Key",True)],size=18,color=ACCENT,after=2)
P(d,"One point each, 13 total. Every value reproducible from Book6.",size=9.5,color=GREY,after=8)
for n,a,ex in KEY: P(d,[(f"{n}.  {a}  ",True),("— "+ex,False)],size=10,after=2)
d.save(OUTDIR+"/SpecialTopic_Hapax_Quiz_Answer_Key.docx")
print("quiz+key saved | answer letters:",[a for _,a,_ in KEY])
