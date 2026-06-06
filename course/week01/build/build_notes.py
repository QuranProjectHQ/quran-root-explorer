# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WK = "/sessions/kind-compassionate-feynman/mnt/RootCourse/week01"
OUT = os.path.join(WK, "Week1_Lecture_Notes.docx")
FIG = os.path.join(WK, "fig_freq_ladder.png")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY   = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- defaults: Arial, 11pt, single spacing, small space-after ----
st = doc.styles["Normal"]
st.font.name = "Arial"; st.font.size = Pt(11)
rpr = st.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rpr.set(qn(a),"Arial")
pf = st.paragraph_format
pf.space_after = Pt(4); pf.line_spacing = 1.0

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(0.9)
sec.top_margin = sec.bottom_margin = Inches(0.8)

# ---- page-number footer ----
ftr = sec.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ftr.add_run()
def field(run, instr):
    for t,txt in (("begin",None),("instr",instr),("separate",None),("end",None)):
        e=OxmlElement("w:fldChar") if t!="instr" else OxmlElement("w:instrText")
        if t=="instr":
            e.set(qn("xml:space"),"preserve"); e.text=instr
        else:
            e.set(qn("w:fldCharType"),t)
        run._r.append(e)
r=ftr.add_run("Week 1 — Frequency · Lecture Notes (v3) · page ")
r.font.size=Pt(8); r.font.color.rgb=GREY
field(ftr.add_run(),"PAGE")

def setcs(run):
    rPr=run._r.get_or_add_rPr(); rf=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")

def para(segs, size=11, after=4, before=0, color=None, align=None):
    # segs: list of (text, bold) OR a plain string
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    p.paragraph_format.line_spacing=1.0
    if isinstance(segs,str): segs=[(segs,False)]
    for text,bold in segs:
        run=p.add_run(text); run.bold=bold; run.font.size=Pt(size); setcs(run)
        if color: run.font.color.rgb=color
    return p

def beat(label, text):
    para([(label+"  ", True), (text, False)], size=10.5, after=3)

def module(title, minutes):
    para([(title, True), ("   ("+minutes+")", False)], size=13, before=10, after=4, color=ACCENT)

# ============================ TITLE ============================
para([("Week 1 — Frequency", True)], size=20, after=2, color=ACCENT)
para([("Lecture Notes · v3 — modular skeleton", True)], size=12, after=2, color=GREY)
para("45-minute lecture. Every module carries the same eight beats — what it is · why we do it · "
     "how it's done · what we get · why it matters · in the data · takeaway · bridge — so no concept "
     "is taught in the abstract. Every value is computed from Book6 via the Quran Root Explorer.",
     size=9.5, after=8, color=GREY)

# ====================== MODULE 0 — OPENING ======================
module("Module 0 · Opening & course frame", "0–3 min")
beat("What it is.", "The course's operating method: turning reading into reproducible counting.")
beat("Why we do it.", "Memory and intuition are powerful but unauditable; a count can be wrong, but it can be checked.")
beat("How it's done.", "State a computed fact, then a separately labeled interpretation — never blur the two.")
beat("What we get.", "A shared, checkable language for making claims about the text.")
beat("Why it matters.", "This one rule is what the entire term rests on.")
beat("In the data.", "By the end of today you will defend a claim like “ظلم is named in 290 ayahs” — and keep it separate from what it means.")
beat("Takeaway.", "We measure not to replace reading, but to make claims about the text auditable.")
beat("Bridge.", "Today's measure is the simplest one — how often a root appears.")

# ============ MODULE 1 — WHAT FREQUENCY ANALYSIS IS ============
module("Module 1 · What frequency analysis is", "3–8 min")
beat("What it is.", "Counting how often a meaning-bearing unit occurs in a text — the word-frequency, “bag-of-words” tradition in corpus linguistics.")
beat("Why we do it.", "Frequency is the most objective, reproducible layer; distribution, partners, co-occurrence and motifs all build on these counts.")
beat("How it's done.", "“Bag of words”: temporarily set aside order and grammar; ask only what is present, and how often.")
beat("What we get.", "A signal of emphasis — a measurable trace, not proof, of what a text dwells on.")
beat("Why it matters.", "It gives a baseline that intuition can be tested against, instead of asserted.")
beat("In the data.", "Book6 holds 6,236 ayahs and ~1,701 roots — far more than any reader could tally by hand; the method scales to all of it.")
beat("Takeaway.", "Frequency measures presence, not importance.")
beat("Bridge.", "Before we can count, we must decide what counts as a unit.")

# ===== MODULE 2 — TOKENIZATION & PREPROCESSING =====
module("Module 2 · From text to countable units — tokenization & preprocessing", "8–14 min")
beat("What it is.", "Turning raw ayahs into comparable, countable tokens.")
beat("Why we do it.", "The same concept wears many surface forms; we want to count the concept, not the spelling.")
beat("How it's done.", "Tokenize each ayah → reduce every word to its 3-consonant root → normalize letters and diacritics (hamza variants, alif-maqsura/ya, strip harakat) → drop function words (في، من، الذي، و), which carry grammar, not theme.")
beat("What we get.", "Each ayah as a bag of content roots, counted once per ayah (document frequency).")
beat("Why it matters.", "Counting once-per-ayah stops one long verse from dominating; root-reduction counts concepts, not inflections.")
beat("In the data.", "كَتَبَ / يَكْتُبُ / كِتاب all collapse to the single root كتب; after this reduction the corpus contains 51,044 root-tokens.")
beat("Takeaway.", "The reduction is lossy on purpose — we trade nuance for comparability (we return to exactly what is lost in Module 7).")
beat("Bridge.", "Raw counts still aren't comparable across texts of different sizes — we need a rate.")

# ============ MODULE 3 — NORMALIZATION TO A RATE ============
module("Module 3 · Normalization to a rate", "14–18 min")
beat("What it is.", "Converting a raw count into a size-independent rate.")
beat("Why we do it.", "A larger text yields larger counts regardless of emphasis, so raw counts can't be compared.")
beat("How it's done.", "Rate per 1,000 ayahs = ayah-freq ÷ 6,236 × 1,000.  Rate per 1,000 roots = term-freq ÷ 51,044 × 1,000 (size-true, over all root-tokens).")
beat("What we get.", "Two honest rates: per-1k-ayahs = share of verses; per-1k-roots = share of all root-occurrences.")
beat("Why it matters.", "The denominator can change the ranking, so the choice of rate is not cosmetic.")
beat("In the data.", "صبر → 93 ayahs → 14.9/1k ayahs (103 tokens → 2.02/1k roots). The ranking can flip: ظلم > هدي per ayah (46.5 vs 43.0) but هدي > ظلم per root (6.19 vs 6.17).")
beat("Takeaway.", "Always report both rates; presence is a rate, not a count.")
beat("Bridge.", "With rates in hand, we can finally read the Qur'anic data honestly.")

# ===== MODULE 4 — THE DATA (EXPANDED) =====
module("Module 4 · The Qur'anic data — reading it", "18–30 min")
beat("What it is.", "Applying the rate to a spread of real roots and reading them side by side.")
beat("Why we do it.", "A single number means little; a ladder of them shows the shape of the text's vocabulary.")
beat("How it's done.", "Retrieve each root in the app's Per-Root Profile; record its raw count, per-1k-ayahs and per-1k-roots; then rank.")
beat("What we get.", "A frequency ladder across three descriptive bands — pervasive, mid, rare.")
beat("Why it matters.", "It orients attention: it shows where to look deeper before investing effort.")
para([("In the data — the frequency ladder (computed from Book6):", True)], size=10.5, after=3)

# ---- comparison table ----
rows = [
 ("Root","Gloss","Ayahs (n)","/1k ayahs","/1k roots","Band"),
 ("كفر","disbelief","465","74.6","10.29","pervasive"),
 ("ظلم","injustice","290","46.5","6.17","pervasive"),
 ("هدي","guidance","268","43.0","6.19","pervasive"),
 ("ضلل","going astray","170","27.3","3.74","pervasive"),
 ("صبر","patience","93","14.9","2.02","mid"),
 ("شكر","gratitude","69","11.1","1.47","mid"),
 ("يسر","ease","40","6.4","0.86","mid"),
 ("عدل","justice","24","3.8","0.55","rare"),
 ("عسر","hardship","12","1.9","0.24","rare"),
]
tbl = doc.add_table(rows=len(rows), cols=6)
tbl.style = "Light Grid Accent 1"
for ri,row in enumerate(rows):
    for ci,val in enumerate(row):
        c=tbl.cell(ri,ci); c.text=""
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
        run=p.add_run(val); run.font.size=Pt(9.5); setcs(run)
        if ri==0: run.bold=True
para("", after=2)
beat("Live walk-through.", "Retrieve كفر — 465 ayahs (74.6/1k), the most-named root in our set — and عسر — 12 ayahs (1.9/1k), the rarest; show the ~39× span between them.")
beat("Predict-then-check #1.", "Ease vs hardship: ask the room which is named more, then retrieve — يسر 40 vs عسر 12, about 3.3× more “ease.”")
beat("Reading habit.", "Always pair a root's rate with its raw count and its size-true rate before saying anything.")
if os.path.exists(FIG):
    doc.add_picture(FIG, width=Inches(5.6))
    cap=doc.paragraphs[-1]; cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para([("Figure 1. Themed-root frequency ladder, per 1,000 ayahs (size-true per-1k-roots rate and raw n alongside).",False)],
         size=8.5, after=4, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
beat("Takeaway.", "The bands are descriptive, not a law of the text — they orient attention, nothing more.")
beat("Bridge.", "One of these numbers should already unsettle a common assumption — that's next.")

# ===== MODULE 5 — HEADLINE & UNLEARN =====
module("Module 5 · The headline finding & the “unlearn”", "30–35 min")
beat("What it is.", "A predict-then-check that overturns an intuition.")
beat("Why we do it.", "The strongest test of a method is when it contradicts what we expected.")
beat("How it's done.", "Ask the room which is named more — عدل (justice) or ظلم (injustice) — then retrieve both in the app.")
beat("What we get.", "ظلم 290 vs عدل 24 — about 12× more.")
beat("Why it matters.", "Intuition usually says justice; the corpus names the violation far more than the ideal.")
beat("In the data.", "ظلم → 290 ayahs (98.6th percentile, “ubiquitous”);  عدل → 24 ayahs.")
beat("Takeaway.", "A genuine unlearning — but this is presence, not endorsement: ظلم is named overwhelmingly to condemn it (polarity is verified later).")
beat("Bridge.", "Those per-root counts are really vectors of numbers — and that idea is the door to everything machines do with text.")

# ===== MODULE 6 — VECTOR-SPACE CODA (NEW) =====
module("Module 6 · Frequency as a vector — the vector-space coda  [preview]", "35–39 min")
beat("What it is.", "The realization that a root's counts form a vector, turning texts and roots into points in a space.")
beat("Why we do it.", "It reveals that today's simple counting is the foundation of modern text analysis.")
beat("How it's done.", "List a root's counts across contexts → that ordered list is a vector → texts and roots become points; nearness = a similar frequency profile.")
beat("What we get.", "A geometric picture in which “related” becomes “close.”")
beat("Why it matters.", "Every later measure — partners, co-occurrence, networks — is geometry on these vectors. This is the vector space model behind how machines “read” text.")
beat("In the data.", "The same per-root counts we just tabulated are the raw coordinates — nothing new is computed here, the numbers are only re-seen.")
beat("Takeaway (labeled preview).", "This is motivation, not a result we test today; richer versions — embeddings, skip-gram — are further-study, not course claims, and are never asserted against Book6.")
beat("Bridge.", "Before trusting any geometry, we must be honest about what the count already threw away.")

# ===== MODULE 7 — ADVANTAGES & LIMITS =====
module("Module 7 · Advantages, disadvantages & limits", "39–43 min")
beat("What it is.", "An honest ledger of what frequency gives and what it loses.")
beat("Why we do it.", "A method is only usable if you know its blind spots.")
beat("How it's done.", "Weigh the advantages — objective & reproducible, scalable, comparable via rates, good for scoping — against what is lost: context.")
beat("What we get.", "Criteria for when a frequency claim is safe and when it misleads.")
beat("Why it matters.", "It stops us from over-reading a count as a verdict.")
beat("In the data.", "كفر is counted whether belief is affirmed or disbelief is quoted-and-refuted; ظلم (injustice) and ظلمات (darkness) collapse into one; speaker (God, prophet, opponents) and syntax (who does what to whom) are gone.")
beat("Takeaway.", "Frequency scopes inquiry; it does not settle it.")
beat("Bridge.", "So any number we report must be paired with a labeled reading — the discipline.")

# ===== MODULE 8 — FACT VS INTERPRETATION + WRAP =====
module("Module 8 · Fact vs. interpretation — the discipline & wrap", "43–47 min")
beat("What it is.", "The two-sentence discipline, and the term's through-line.")
beat("Why we do it.", "It keeps computed facts and human readings from blurring into each other.")
beat("How it's done.", "Write one sentence of fact (with its normalization), then one separately labeled sentence of interpretation.")
beat("What we get.", "A defensible reading a peer can audit — and tell which sentence is which.")
beat("Why it matters.", "The whole course is graded on keeping these two apart; it is also how a community reads itself — its moral vocabulary shows in what it names most, and what it rarely names.")
beat("In the data.", "Fact: “ظلم is named in 290 ayahs (46.5/1k ayahs; 6.17/1k roots), about 12× عدل.”  Interpretation: “I read this as the text foregrounding the diagnosis of wrong over the abstract ideal.”")
beat("Takeaway.", "Presence is a rate, not a count; it can defy intuition, but it is context-blind.")
beat("Bridge.", "Frequency says how much; next week asks where — distribution & concentration. Your task: your assigned root + surah, record both rates, one fact + one labeled interpretation.")

doc.save(OUT)
print("saved", OUT)
print("paragraphs:", len(doc.paragraphs))
