# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WK="/sessions/kind-compassionate-feynman/mnt/RootCourse/week01"
OUT=os.path.join(WK,"Week1_Instructor_Script.docx")
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55); CUE=RGBColor(0x8A,0x4B,0x08)

doc=Document()
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(11)
rf=st.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.0
sec=doc.sections[0]
sec.left_margin=sec.right_margin=Inches(0.9); sec.top_margin=sec.bottom_margin=Inches(0.8)
ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ftr.add_run("Week 1 — Instructor Script (v3) · page "); r.font.size=Pt(8); r.font.color.rgb=GREY
run=ftr.add_run()
for t,txt in (("begin",None),("instr","PAGE"),("separate",None),("end",None)):
    e=OxmlElement("w:fldChar") if t!="instr" else OxmlElement("w:instrText")
    if t=="instr": e.set(qn("xml:space"),"preserve"); e.text="PAGE"
    else: e.set(qn("w:fldCharType"),t)
    run._r.append(e)

def setcs(run):
    rPr=run._r.get_or_add_rPr(); rff=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rff.set(qn(a),"Arial")
def para(segs,size=11,after=4,before=0,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=1.05
    if isinstance(segs,str): segs=[(segs,False)]
    for text,bold in segs:
        rr=p.add_run(text); rr.bold=bold; rr.font.size=Pt(size); setcs(rr)
        if color: rr.font.color.rgb=color
    return p
def cue(text): para([("▸ "+text,False)],size=9.5,after=3,color=CUE)
def marker(time,title): para([(time+"  ",True),(title,True)],size=12,before=10,after=3,color=ACCENT)
def say(text): para(text,size=11,after=5)

# TITLE
para([("Week 1 — Instructor Lecture Script (Frequency)",True)],size=18,after=2,color=ACCENT)
para("Spoken script, ~45 minutes, frequency-only (distribution → Week 2). Mirrors the v3 lecture notes: nine modules, "
     "each carrying what · why · how · what-we-get · why-it-matters · in-the-data · takeaway · bridge. "
     "▸ lines are delivery cues. Time markers are cumulative. Every value is computed from Book6.",
     size=9.5,after=8,color=GREY)

# M0
marker("0:00","Module 0 · Opening & course frame")
cue("Slides 1–2. Ask before you explain.")
say("Open with a question. If I asked which idea matters most in the Qur’an, you’d answer from memory and feeling — and we’d have no way to check each other. So today we take a smaller, humbler question, one we can actually settle: not what matters most, but how often a root appears.")
say("Here is the one rule we set today and keep all term: we state a computed fact, and then, separately and clearly labeled, an interpretation. We never smuggle an opinion in as a fact. We measure not to replace reading, but to make our claims about the text auditable — a count can be wrong, but it can be checked. By the end of the hour you’ll defend a claim like “ظلم is named in 290 ayahs,” and keep it cleanly apart from what you think it means. And today we claim presence only — not importance, not meaning, not relationships.")

# M1
marker("3:00","Module 1 · What frequency analysis is")
cue("Slide 3.")
say("Frequency analysis means counting how often a meaning-bearing unit occurs in a text — the word-frequency, or “bag-of-words,” tradition in corpus linguistics. We temporarily set order and grammar aside and ask only: what is present, and how often?")
say("The working assumption is modest — how often a text names something is a signal of emphasis, a measurable trace, not proof. We start here because frequency is the most objective, reproducible layer; distribution, partners, co-occurrence and motifs all build on these counts. And the reason we let a machine do it: Book6 has 6,236 ayahs and about 1,701 roots — no reader could tally that by hand, but the method scales to all of it. Hold onto the takeaway: frequency measures presence, not importance. But before we can count, we have to decide what counts as a unit.")

# M2
marker("8:00","Module 2 · From text to countable units")
cue("Slide 4. Walk the pipeline on the board.")
say("The same idea wears many clothes, and we want to count the idea, not the spelling. So we run a pipeline. First we tokenize each ayah into word-tokens. Then we reduce each word to its three-consonant root, so كَتَبَ, يَكْتُبُ and كِتاب all collapse to the one root كتب. We normalize letters and diacritics so the same root matches everywhere, and we drop the function words — في، من، الذي — which carry grammar, not theme, and mostly have no triliteral root anyway.")
say("Our counting unit is once per ayah — document frequency — so one long verse can’t dominate just by repeating a root. After this reduction the corpus holds 51,024 root-tokens. Be honest about it: this reduction is lossy on purpose, we’re trading nuance for comparability, and we’ll face exactly what we discarded later. Raw counts still aren’t comparable across texts of different sizes — so next we turn a count into a rate.")

# M3
marker("14:00","Module 3 · Normalization to a rate")
cue("Slides 6–7. Two formulas on the board.")
say("A bigger text yields bigger counts regardless of emphasis, so a raw count alone tells us nothing comparable. We fix that with two rates. The rate per 1,000 ayahs is ayah-frequency divided by 6,236, times 1,000 — the share of verses a root touches. The rate per 1,000 roots is term-frequency divided by 51,024, times 1,000 — the size-true share of all root-occurrences.")
say("Why bother with two? Because the denominator can flip the ranking, and that is not cosmetic. Watch ظلم and هدي: per 1,000 ayahs ظلم leads, 46.5 to 43.0; but per 1,000 roots هدي edges ahead, 6.19 to 6.17, because هدي repeats slightly more and sits in denser verses. So always report both rates — presence is a rate, not a count. Now we can finally read the data honestly.")

# M4
marker("18:00","Module 4 · The Qur’anic data — reading it")
cue("Slide 8 — the themed-root ladder with numbers. App live.")
say("A single number means little; a ladder of them shows the shape of the vocabulary. Here’s our themed set, ranked, each with its raw count, its rate per 1,000 ayahs, and its size-true rate per 1,000 roots. Read the span: at the top كفر in 465 ayahs, 74.6 per 1,000 — the most-named root in our set; at the bottom عسر in just 12 ayahs, 1.9 per 1,000. That’s about a thirty-nine-fold spread. Three descriptive bands fall out — pervasive, mid, rare — but they only orient attention; they are not a law of the text.")
cue("Go to the app. Type كفر → Analyze → Per-Root Profile; then عسر.")
say("Let’s retrieve them live so you trust the numbers. Type كفر, Analyze, open the Per-Root Profile — 465 ayahs. Now عسر — 12. Same tool, two ends of the ladder.")
cue("Predict, then check — ease vs hardship.")
say("Quick prediction before we look: which is named more, ease or hardship — يسر or عسر? Check it: يسر 40 ayahs, عسر 12 — about three and a third times more “ease.” Notice the habit I just used: I paired the rate with the raw count before saying a word. Always do that. And one of these numbers is about to unsettle a very common assumption — that’s next.")

# M5
marker("30:00","Module 5 · The headline finding & the unlearn")
cue("Slide 9. Predict, then check — then type both in the app.")
say("The strongest test of a method is when it contradicts what we expected. So predict with me: which does the Qur’an name more often — عدل, justice, or ظلم, injustice? Most rooms say justice. Let’s check. ظلم: 290 ayahs, 98.6th percentile, the app calls it ubiquitous. عدل: 24 ayahs. That’s about twelve times more.")
say("That is a genuine unlearning — the corpus names the violation far more than the ideal. But stay disciplined: this is presence, not endorsement. ظلم is named overwhelmingly in order to condemn it, and we verify polarity later in the term. Hold that thought, because those per-root counts are really vectors of numbers — and that idea is the doorway to everything machines do with text.")

# M6 (NEW coda)
marker("35:00","Module 6 · Frequency as a vector — the vector-space coda  [preview]")
cue("New slide — vector-space. Flag clearly: this is preview, not tested today.")
say("Here’s the idea that makes today bigger than it looks. Take a root and list its counts across contexts — that ordered list of numbers is a vector. Do it for every root, and each root becomes a point in a space; roots with similar frequency profiles sit close together. That is the vector space model, and it is the foundation of how machines “read” text. Everything we do in later weeks — partners, co-occurrence, networks — is just geometry on these vectors.")
say("And notice: I computed nothing new just now. The very counts we tabulated a moment ago are the raw coordinates — we only re-saw them. Be explicit with the room: this is motivation, a preview, not a result we test today; the richer versions you may have heard of — embeddings, skip-gram — are further-study, and we never assert them against Book6. Before we trust any geometry, though, we have to be honest about what the plain count already threw away.")

# M7
marker("39:00","Module 7 · Advantages, disadvantages & limits")
cue("Slides 10–11 (dark). Be explicit about the cost.")
say("A method is only usable if you know its blind spots. The advantages are real: frequency is objective and reproducible — anyone re-running it gets the same numbers; it’s scalable to all 6,236 ayahs; it’s comparable once we use rates; and it’s excellent for scoping — showing where to look deeper before you spend effort.")
say("But what it loses is context, and that’s a lot. كفر is counted whether belief is affirmed or disbelief is quoted and refuted — stance is gone. Speaker is gone: God, prophet, opponents, all undistinguished. Polysemy collapses — ظلم, injustice, and ظلمات, darkness, count as one. Syntax is gone — who does what to whom. So the rule of thumb: frequency scopes inquiry; it does not settle it. Which means every number we report has to be paired with a labeled reading — the discipline we close on.")

# M8
marker("43:00","Module 8 · Fact vs. interpretation — the discipline, wrap & exercise")
cue("Slide 16. Model one aloud, slowly.")
say("Here is the form the whole course is graded on. One sentence of computed fact, with its normalization stated. Then one separately labeled sentence of interpretation. Fact: “ظلم is named in 290 ayahs — 46.5 per 1,000 ayahs, 6.17 per 1,000 roots — about twelve times عدل.” Interpretation: “I read this as the text foregrounding the diagnosis of wrong over the abstract ideal.” A classmate should be able to tell instantly which sentence is which; if they can’t, rewrite. This is also how a community reads itself — its moral vocabulary shows in what it names most, and what it rarely names.")
cue("Slide 17 — assignment table. Slide 18 — close.")
say("For the exercise: find your member number. You have one short surah to count by hand and one root to profile in the app, recording both rates each time, plus one fact and one labeled interpretation for each. Submit one screenshot and your two readings the night before next class — that submission is what lets you join the debrief.")
say("To close: we refused the big question and took a small one, and by normalizing it, distrusting it where it’s thin, and labeling what we add, we earned a finding we can defend. Presence is a rate, not a count — it can defy intuition, but it is context-blind. Next week we ask not how much, but where: distribution and concentration. See you then.")

doc.save(OUT)
print("saved",OUT,"| paragraphs:",len(doc.paragraphs),"| words:",sum(len(p.text.split()) for p in doc.paragraphs))
