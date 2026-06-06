# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
WK="/sessions/kind-compassionate-feynman/mnt/RootCourse/week04"; FD=os.path.join(WK,"figs")
OUT=os.path.join(WK,"Week4_Lecture_Notes.docx")
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55)
doc=Document()
stl=doc.styles["Normal"]; stl.font.name="Arial"; stl.font.size=Pt(11)
rf=stl.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
stl.paragraph_format.space_after=Pt(4); stl.paragraph_format.line_spacing=1.04
sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(0.9); sec.top_margin=sec.bottom_margin=Inches(0.8)
ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ftr.add_run("Week 4 — Co-occurrence · Lecture Notes (v1) · page "); r.font.size=Pt(8); r.font.color.rgb=GREY
run=ftr.add_run()
for t in ("begin","instr","separate","end"):
    e=OxmlElement("w:instrText") if t=="instr" else OxmlElement("w:fldChar")
    if t=="instr": e.set(qn("xml:space"),"preserve"); e.text="PAGE"
    else: e.set(qn("w:fldCharType"),t)
    run._r.append(e)
def cs(run):
    rPr=run._r.get_or_add_rPr(); rff=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rff.set(qn(a),"Arial")
def para(segs,size=11,after=4,before=0,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=1.04
    if isinstance(segs,str): segs=[(segs,False)]
    for tx,b in segs:
        rr=p.add_run(tx); rr.bold=b; rr.font.size=Pt(size); cs(rr)
        if color: rr.font.color.rgb=color
    return p
def beat(label,text): para([(label+"  ",True),(text,False)],size=10.5,after=4)
def module(title,minutes): para([(title,True),("   ("+minutes+")",False)],size=13,before=11,after=4,color=ACCENT)
def figure(name,cap,width=5.8):
    p=os.path.join(FD,name)
    if os.path.exists(p):
        doc.add_picture(p,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        para([(cap,False)],size=8.5,after=5,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER)

para([("Week 4 — Co-occurrence",True)],size=20,after=2,color=ACCENT)
para([("Lecture Notes · v1 — deep (with analogies)",True)],size=12,after=2,color=GREY)
para("A 45-minute lecture, eight modules, each carrying the eight beats — what it is · why · how · what we get · why it matters · in the data · takeaway · bridge. This is our first PAIR measure and our first FIND task: given a target root, which other root shares the most ayahs with it — and why the raw count will try to fool us. Worked target: صلو (prayer). Scope: co-occurrence as observed-vs-expected (a first, frequency-based control); the formal length-aware null, Monte-Carlo validation, and lift TIERS are Week 5, named here only as a preview. Every value computed from Book6.",size=9.5,after=8,color=GREY)

module("Module 1 · Opening & recap — from one root to a pair","0–5 min")
beat("What it is.","Until now every measure looked at one root at a time: Week 1 asked how often a root is named, Week 2 where it sits and how it spreads, Week 3 what shapes it wears and what company it keeps. Today we change the unit entirely and look at a PAIR — two roots side by side — and ask how often they appear in the same ayah. That is co-occurrence, and it is the first step from describing single words to mapping the structure that connects them.")
beat("Why we do it.","Last week's partners told us a root keeps company; this week we make that precise and turn it into a contest: among several candidate roots, which one really shares the most with a target? It is the difference between saying “they're friends” and being able to prove who the closest friend is.")
beat("How it's done.","Pick a target root, line up a slate of candidate roots, and for each candidate count the ayahs the two share. Then — and this is the whole lesson — we will check whether that count means anything.")
beat("What we get.","A ranked answer to a “find” question — who is the closest companion — and, just as important, a hard-won caution about how to rank, because the obvious way to rank, by raw shared count, is not just imperfect but systematically wrong. We will replace it with a way of ranking that asks a better question.")
beat("Why it matters.","Co-occurrence is the gateway to structure: which roots travel together is how we will eventually map clusters, hubs, and three-root motifs in the weeks ahead — the whole back half of the course is built on this one idea. But it is also the single most over-read number in all of text analysis, the place where confident-sounding nonsense is most often manufactured, so we have to learn to read it honestly before we are allowed to trust it.")
beat("In the data.","Our worked target is صلو, prayer (90 ayahs). By the end you will know its true closest companion — and it will not be the one the raw count nominates.")
beat("Takeaway.","Co-occurrence asks a new question — not how much or where, but WHO WITH — and, like every measure before it, the honest answer needs a control before it can be trusted.")
beat("Bridge.","Start with the raw ingredient: the count of shared ayahs.")

module("Module 2 · Counting shared ayahs (the joint count)","5–11 min")
beat("What it is.","The joint count of two roots — the number of ayahs that contain BOTH.")
beat("Why we do it.","It is the raw material of every co-occurrence measure ever devised; everything we build today — expectation, ratio, next week's null and tiers — is just a way of putting this one number into context so it can be trusted.")
beat("How it's done.","Take the set of ayahs containing the target and the set containing the candidate, and count the overlap — the intersection of two circles in a Venn diagram. Note what we deliberately ignore: we do not care where in the ayah they sit, or in what grammatical relation; only that both roots are present in the same verse. It is the bluntest possible notion of “together,” and we sharpen it later.")
beat("What we get.","A single number per pair — the size of the overlap. For صلو and زكو (zakat), that number is 28: twenty-eight ayahs name both prayer and almsgiving, two acts the Qur'an seems reluctant to separate.")
figure("fig_cooccurrence_concept.png","Figure 1. Co-occurrence is the overlap: 28 ayahs contain both صلو (prayer) and زكو (zakat).",width=5.4)
beat("Why it matters.","The joint count feels like the answer — surely, the intuition runs, the root that shares the most ayahs is the closest companion. It is the same intuition that would tell you your closest friend is whoever appears in the most of your photographs. The rest of the lecture is about why that intuition fails, and what to put in its place.")
beat("In the data.","صلو shares 52 ayahs with قوم (establish), 44 with ءله (God), 28 with زكو (zakat), 27 with ءتي (give). On raw counts, قوم and ءله lead the field and zakat trails in third. If we stopped here we would declare that prayer's closest companion is “establish” or “God” — a conclusion that is both unremarkable and, as we will see, an artefact of counting.")
beat("Takeaway.","The joint count is real and we will keep it, but it is an ingredient, not a verdict — the flour, not the bread.")
beat("Bridge.","Because that raw ranking is about to mislead us — and we should see exactly how.")

module("Module 3 · Why the raw count misleads — the frequency confound","11–18 min")
beat("What it is.","The frequency confound: a root that appears everywhere will share many ayahs with everything, whether or not there is any real bond.")
beat("Why we do it.","If we do not name this trap, we will crown the wrong companion every time — and the trap is the same shape as Week 2's length confound, just with frequency instead of surah size.")
beat("How it's done.","Here is the analogy to keep. Imagine a celebrity who is photographed at every event in town. You will find hundreds of photos of that celebrity standing next to your friend — not because they are close, but because the celebrity is in every photo. Now imagine an ordinary person who appears in only fifty photos all year, and your friend is in twenty-eight of them. THAT is closeness — twenty-eight of fifty is a relationship; two hundred of a celebrity's thousands is noise. The root ءله (God) is the celebrity of this corpus: it is named in 1,879 ayahs, more than a quarter of the whole book, so of course it shares 44 ayahs with prayer — it shares dozens with nearly everything. زكو, by contrast, appears in only 56 ayahs in the entire Qur'an, and 28 of them — exactly half — fall beside prayer. The small number is the loud one.")
figure("fig_raw_joint.png","Figure 2. Raw co-occurrence with prayer (صلو): the frequent roots قوم and ءله top the list — but they top everyone's list.",width=6.0)
beat("What we get.","A diagnosis: the raw ranking measures the candidate's overall fame at least as much as its bond with the target, and usually more. The louder a root is everywhere, the higher it floats on every root's raw list — which makes the raw list nearly useless for telling pairs apart.")
beat("Why it matters.","This is the course's spine again, in a fourth costume — control before you conclude. In Week 1 we divided a raw count by the size of the text to get a fair rate; in Week 2 we divided by surah length to find a root's true home; today we must divide out the candidate's sheer frequency to find a target's true companion. Same discipline, a new thing to divide out each time — that is the through-line of the whole course, and once you see it here you will expect it everywhere.")
beat("In the data.","قوم (establish) shares 52 ayahs with صلو — the highest raw count — and the phrase أقيموا الصلاة, “establish prayer,” is indeed common; but قوم appears in 597 ayahs and establishes a great many things besides prayer, so the bond is real yet diluted. ءله shares 44, but God is named beside almost everything in the book, so 44 is unremarkable. The raw leaders are exactly the roots whose leadership means the least.")
beat("Takeaway.","A big shared count can be nothing more than the candidate being everywhere.")
beat("Bridge.","So we need a yardstick for “more than you'd expect from a root that common.”")

module("Module 4 · The fix — observed vs expected-by-chance","18–25 min")
beat("What it is.","A control that compares the shared count we OBSERVED to the count we'd EXPECT if the two roots were sprinkled through the Qur'an independently, with no bond at all.")
beat("Why we do it.","Expectation gives us a fair baseline because it already has the frequency confound built in: it knows a candidate named in a quarter of the book will land beside the target dozens of times by sheer chance, so it sets the bar that high, and only what clears the bar counts as signal. The frequent root no longer gets credit for showing up everywhere.")
beat("How it's done.","The expected overlap is just the two frequencies multiplied and scaled to the corpus — freq(target) × freq(candidate) ÷ 6,236 ayahs. Worked for prayer and God: 90 × 1,879 ÷ 6,236 ≈ 27, so we would EXPECT about 27 shared ayahs from two roots that common even if nothing tied them together — and we observe 44, only a little above. For prayer and zakat: 90 × 56 ÷ 6,236 ≈ 0.8 expected — and we observe 28. Then the measure is simply how many TIMES the observed count beats that expectation. Think of a crowded reception where everyone mills at random: two guests who keep ending up in the same corner far more often than random milling would predict are clearly seeking each other out — and the rarer the guest, the more telling it is.")
figure("fig_observed_vs_expected.png","Figure 3. Observed vs expected-by-chance. God's 44 shared ayahs barely beat the ~27 you'd expect from a root that frequent; zakat's 28 tower over an expectation of less than 1.",width=5.8)
beat("What we get.","A ratio — “× more than chance” — that is fair across candidates of wildly different frequencies.")
beat("Why it matters.","It flips the picture entirely. ءله's 44 shared ayahs are only about 1.6× what chance predicts — essentially the background hum you would get from a root named in a quarter of the book, doing nothing in particular. قوم's 52 fall to ×6. زكو's 28, against an expectation below one, stand at 34.6× chance. The candidate that was third on the raw list is first by a landslide once the field is levelled.")
beat("In the data.","Expected overlap for صلو–ءله is ~27 (observed 44 → ×1.6); for صلو–زكو it is under 1 (observed 28 → ×34.6). The frequent root collapses; the specific one soars.")
beat("Takeaway.","Ask not “how many ayahs do they share?” but “how many MORE than you'd expect by chance?”")
beat("Bridge.","Apply that question to prayer's whole slate, and a companion no one nominated steps forward.")

module("Module 5 · The headline & the “unlearn”","25–32 min")
beat("What it is.","A corpus-level reversal: once we control for frequency, prayer's closest companion changes completely.")
beat("Why we do it.","To make vivid that the raw count and the controlled measure can name DIFFERENT winners — and only one of them is trustworthy.")
beat("How it's done.","Rank صلو's candidates twice: once by raw shared ayahs, once by times-over-chance, and lay the two lists side by side.")
beat("What we get.","Raw says قوم (52) and God (44). Controlled says زكو — zakat — at 34.6× chance, far ahead of everything.")
figure("fig_controlled_ratio.png","Figure 4. Controlled for frequency, zakat is prayer's true companion — ×34.6 over chance, dwarfing the frequent roots.",width=5.8)
figure("fig_flip.png","Figure 5. The flip: قوم and God top the raw list (left); zakat tops once you control for frequency (right).",width=6.2)
beat("Why it matters.","This is one of the Qur'an's signature pairings — أقيموا الصلاة وآتوا الزكاة, “establish prayer and give zakat.” Half of all the verses that mention zakat at all mention it beside prayer. The bond a careful reader feels is now a measured fact — but only the controlled measure could see it, because zakat is too rare to win a raw shouting match.")
beat("In the data.","Of زكو's 56 total ayahs, 28 — exactly half — are shared with صلو; you almost cannot mention almsgiving in the Qur'an without prayer in the same breath. No frequent root comes close to that density of bond: قوم shares more ayahs in absolute terms, but only a tiny fraction of قوم's 597 ayahs are with prayer, so its bond is thin and wide where zakat's is narrow and deep.")
beat("Takeaway (the unlearn).","Prayer's closest companion is not “establish” or “God,” which merely sit near everything; it is charity — and you can only see it once you stop rewarding roots for being common.")
beat("Bridge.","Now the skill itself: how to read a candidate slate and pick the real winner.")

module("Module 6 · Reading a candidate slate — the find-task","32–38 min")
beat("What it is.","The week's “find” skill: given a target and a slate of candidates, rank them controlled and name the true companion.")
beat("Why we do it.","This is exactly the exercise each member will run on their own root, so the method has to be a habit, not a one-off.")
beat("How it's done.","For each candidate compute the times-over-chance ratio; ignore any pair with fewer than five shared ayahs, because two or three overlaps are too little evidence to trust (the small-sample caution we met with the support floor in Week 2 returns here); then read off the top of the controlled list — never the raw list. The discipline is a single habit: rank by surprise, not by size.")
beat("What we get.","A defensible answer for any target, and a second and third worked case to prove the method is not a one-off trick tuned to prayer and zakat.")
figure("fig_kayl_wazn.png","Figure 6. A second case — “measure” (كيل) is bonded to “weight” (وزن) at ×137: the honest scales of the marketplace (the المطففين passage).",width=5.8)
beat("Why it matters.","The pattern holds across the corpus with almost monotonous reliability: take any target, and its raw winner is almost always one of a handful of ubiquitous roots — God, say, establish — while its controlled winner is a specific bond that actually teaches you something about the text. The raw list is the same boring names every time; the controlled list is where the Qur'an's own associations live.")
figure("fig_targets_overview.png","Figure 7. Across twelve targets, the raw winner is a generic frequent root; the controlled winner is the specific bond every time.",width=6.0)
beat("In the data.","كيل (measure) and وزن (weight) share only six ayahs, but at ×137 over chance — the Qur'an's warning to the المطففين, those who give short measure and short weight in the marketplace. Covenant (عهد) bonds to breaking (نقض) at ×96 — the Qur'an speaks of the covenant chiefly to condemn its breach. Prostration (سجد) bonds to bowing (ركع) at ×46, the two postures of the prayer itself; and the “goodly loan,” قرض, bonds to حسن at ×30 — قرضًا حسنًا. In each case a rare, specific pair the raw count would have buried.")
beat("Takeaway.","To find a root's true companion, rank by times-over-chance and read the top; treat the raw list as a decoy that will always nominate the loudest root in the room.")
beat("Bridge.","But even a controlled ranking has limits we must state plainly.")

module("Module 7 · Advantages, limits & what co-occurrence can't say","38–43 min")
beat("What it is.","An honest ledger of what observed-vs-expected gives us and what it still cannot.")
beat("Why we do it.","A measure is only safe in the hands of someone who knows its blind spots; a tool you trust blindly is more dangerous than no tool at all, because it lends false confidence to a wrong answer.")
beat("How it's done.","Weigh the advantages — it is fair across frequencies, it surfaces specific bonds, it scales to the whole corpus — against three real limits.")
beat("What we get.","Criteria for a co-occurrence claim you can defend, and a clear map of the three things this measure still cannot do — each of which is a later week.")
beat("Why it matters.","First limit: our control is deliberately rough — it divides out each candidate's frequency but still ignores ayah LENGTH, and a long ayah packed with twenty roots hands out co-occurrences far more freely than a three-root ayah. That is the very same length confound we met in Week 2, now waiting one level down, and Week 5 fixes it: it replaces our back-of-envelope expectation with a length-aware null model, checks it by running thousands of shuffled simulations, and only then sorts pairs into tiers of real, borderline, and spurious. Second limit: co-occurrence is symmetric — it tells you prayer and zakat travel together, but not whether mentioning prayer pulls in zakat or the reverse; that asymmetry, P(A given B) versus P(B given A), is Week 6. Third, and most important to say aloud: association is not cause, and not even agreement — two roots can share verses because the text praises them together, or because it contrasts them, as we saw with antonym partners last week.")
figure("fig_paradise_cluster.png","Figure 8. Co-occurrence also reveals clusters: paradise (جنن) sits with Eden, rivers, and eternity — a constellation, not a single pair.",width=5.6)
beat("In the data.","جنن (the garden) bonds to عدن (Eden, ×32), to تحت and جري (the rivers that flow beneath, جنات تجري من تحتها الأنهار), and to خلد (eternity) — a whole constellation that co-occurrence lights up at once.")
beat("Takeaway.","Co-occurrence maps who travels with whom; it does not yet tell us how strongly to trust it, which way it points, or why.")
beat("Bridge.","Trust comes next week: a rigorous, length-aware null and the tiers that grade a bond as real, borderline, or spurious.")

module("Module 8 · Fact vs. interpretation — the discipline & wrap","43–47 min")
beat("What it is.","The two-sentence discipline applied to a co-occurrence claim, and the term's through-line.")
beat("Why we do it.","A vivid pairing like prayer-and-charity begs to be preached; the discipline keeps the measured fact and the meaning visibly apart.")
beat("How it's done.","One sentence of fact — a controlled ratio, with its support noted — and one separately labeled sentence of interpretation, beginning “I read this as …”. The fact is what the corpus forces on anyone who counts; the interpretation is what you, a reader, make of it.")
beat("What we get.","A reading a classmate can audit at a glance: they can point to the sentence a computer could reproduce and the sentence that is yours alone. If they cannot tell the two apart, the reading has blurred and must be rewritten — the same peer-test that has governed every week.")
beat("Why it matters.","The course is graded on this separation, and co-occurrence is where it is most tempting to break — because the bonds are so evocative that the number feels like it already contains the sermon.")
beat("In the data.","Fact: “Of zakat's 56 ayahs, 28 are shared with prayer — 34.6 times the chance rate, the highest controlled bond on prayer's slate.”  Interpretation: “I read prayer and charity in the Qur'an as a single act of devotion with two faces, vertical and horizontal.”")
beat("Takeaway.","Frequency told us how much, distribution where, forms and partners what-shape and who-with; co-occurrence now measures who-with rigorously — all computed, all needing a labeled reading.")
beat("Bridge.","Next week we make the trust rigorous: a pair's lift, a length-aware null tested by simulation, and the tiers that tell us which bonds are real.")

doc.save(OUT); print("saved | words:",sum(len(p.text.split()) for p in doc.paragraphs))
