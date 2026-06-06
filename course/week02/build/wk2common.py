# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55); TEAL=RGBColor(0x0E,0x6D,0x5C)
def new_doc(footer_label):
    doc=Document()
    st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(11)
    rf=st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
    st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.0
    sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(0.9); sec.top_margin=sec.bottom_margin=Inches(0.8)
    ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=ftr.add_run(footer_label+" · page "); r.font.size=Pt(8); r.font.color.rgb=GREY
    run=ftr.add_run()
    for t in ("begin","instr","separate","end"):
        e=OxmlElement("w:instrText") if t=="instr" else OxmlElement("w:fldChar")
        if t=="instr": e.set(qn("xml:space"),"preserve"); e.text="PAGE"
        else: e.set(qn("w:fldCharType"),t)
        run._r.append(e)
    return doc
def setcs(run):
    rPr=run._r.get_or_add_rPr(); rff=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rff.set(qn(a),"Arial")
def P(doc,segs,size=11,after=4,before=0,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=1.05
    if isinstance(segs,str): segs=[(segs,False)]
    for text,bold in segs:
        rr=p.add_run(text); rr.bold=bold; rr.font.size=Pt(size); setcs(rr)
        if color: rr.font.color.rgb=color
    return p
def H(doc,text,size=15,color=ACCENT,before=10,after=4): return P(doc,[(text,True)],size=size,color=color,before=before,after=after)
def TITLE(doc,t,sub):
    P(doc,[(t,True)],size=19,after=2,color=ACCENT); P(doc,[(sub,False)],size=10,after=8,color=GREY)
def bullet(doc,segs,size=10.5):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.05
    if isinstance(segs,str): segs=[(segs,False)]
    for text,bold in segs:
        rr=p.add_run(text); rr.bold=bold; rr.font.size=Pt(size); setcs(rr)
    return p
def table(doc,rows,header=True,fontsize=9.5,widths=None):
    t=doc.add_table(rows=len(rows),cols=len(rows[0])); t.style='Light Grid Accent 1'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.cell(ri,ci); c.text=""
            par=c.paragraphs[0]; par.paragraph_format.space_after=Pt(0); par.paragraph_format.line_spacing=1.0
            rr=par.add_run(str(val)); rr.font.size=Pt(fontsize); setcs(rr)
            if header and ri==0: rr.bold=True
    return t

# member data: root, translit, breadth, top3, gini, raw(name,tok,stok,p), home(name,tok,stok,p)
MEMBERS=[
(1,'ظلم','ẓulm',59,21.7,0.738,('al-Baqara',31,3966,7.8),('Ibrahim',9,568,15.8)),
(2,'نفس','nafs',63,24.1,0.733,('al-Baqara',35,3966,8.8),('al-Infitar',3,51,58.8)),
(3,'هدي','hudā',62,22.4,0.730,('al-Baqara',34,3966,8.6),('as-Sajda',5,240,20.8)),
(4,'ضلل','ḍalāl',56,19.4,0.702,('al-Anʿam',12,1968,6.1),('Nuh',3,157,19.1)),
(5,'رزق','rizq',44,21.1,0.748,('al-Baqara',12,3966,3.0),('at-Talaq',3,204,14.7)),
(6,'صبر','ṣabr',45,24.7,0.754,('al-Baqara',9,3966,2.3),('at-Tur',3,205,14.6)),
(7,'شكر','šukr',35,23.2,0.793,('al-Baqara',7,3966,1.8),('al-Insan',3,179,16.8)),
(8,'صرط','ṣirāṭ',25,24.4,0.832,('al-Anʿam',5,1968,2.5),('Ya-Sin',3,443,6.8)),
(9,'يسر','yusr',27,27.5,0.823,('al-Baqara',5,3966,1.3),('al-Layl',3,48,62.5)),
(10,'عدل','ʿadl',11,50.0,0.935,('al-Anʿam',6,1968,3.0),('al-Anʿam',6,1968,3.0)),
(11,'قسط','qisṭ',15,36.4,0.899,('an-Nisaʾ',3,2496,1.2),('Yunus',3,1147,2.6)),
(12,'رشد','rušd',9,57.9,0.945,('al-Kahf',4,1080,3.7),('al-Jinn',4,182,22.0)),
]
WK="/sessions/kind-compassionate-feynman/mnt/RootCourse/week02"
