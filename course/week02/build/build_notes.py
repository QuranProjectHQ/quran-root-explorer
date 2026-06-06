# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
WK="/sessions/kind-compassionate-feynman/mnt/RootCourse/week02"; OUT=os.path.join(WK,"Week2_Lecture_Notes.docx")
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55)
doc=Document()
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(11)
rf=st.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.0
sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(0.9); sec.top_margin=sec.bottom_margin=Inches(0.8)
ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ftr.add_run("Week 2 — Distribution & Concentration · Lecture Notes (v2) · page "); r.font.size=Pt(8); r.font.color.rgb=GREY
run=ftr.add_run()
for t in ("begin","instr","separate","end"):
    e=OxmlElement("w:instrText") if t=="instr" else OxmlElement("w:fldChar")
    if t=="instr": e.set(qn("xml:space"),"preserve"); e.text="PAGE"
    else: e.set(qn("w:fldCharType"),t)
    run._r.append(e)
def setcs(run):
    rPr=run._r.get_or_add_rPr(); rff=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rff.set(qn(a),"Arial")
def para(segs,size=11,after=4,before=0,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=1.0
    if isinstance(segs,str): segs=[(segs,False)]
    for text,bold in segs:
        rr=p.add_run(text); rr.bold=bold; rr.font.size=Pt(size); setcs(rr)
        if color: rr.font.color.rgb=color
    return p
def beat(label,text): para([(label+"  ",True),(text,False)],size=10.5,after=3)
def module(title,minutes): para([(title,True),("   ("+minutes+")",False)],size=13,before=10,after=4,color=ACCENT)
def figure(path,caption,width=5.6):
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        para([(caption,False)],size=8.5,after=4,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER)

para([("Week 2 — Distribution & Concentration",True)],size=20,after=2,color=ACCENT)
para([("Lecture Notes · v2 — modular skeleton",True)],size=12,after=2,color=GREY)
para("45-minute lecture. Same eight beats per module as Week 1. Scope: distribution & concentration only "
     "(partners, co-occurrence, lift and motifs are Weeks 3–5). Size-true rule (locked): every density is normalized "
     "to per 1,000 ROOT-TOKENS — never per ayah, because ayahs vary in length just as surahs do. Every value computed from Book6.",
     size=9.5,after=8,color=GREY)

module("Module 0 · Opening & recap","0–3 min")
beat("What it is.","A bridge from Week 1: frequency told us how much a root is named; today we ask where it sits and how it spreads.")
beat("Why we do it.","A root's total says nothing about whether it is woven through the whole Qur'an or pooled in a few surahs.")
beat("How it's done.","We move from one number (the count) to the shape of a root across all 114 surahs.")
beat("What we get.","A second coordinate on every root — its geography, not just its size.")
beat("Why it matters.","Two roots with nearly identical totals can have opposite footprints.")
beat("In the data.","ظلم is named in 290 ayahs — but in how many of the 114 surahs, and are they evenly spread? Frequency cannot say.")
beat("Takeaway.","How much is not where; presence has a geography.")
beat("Bridge.","Start with the simplest geographic question — in how many surahs?")

module("Module 1 · What distribution & concentration are","3–9 min")
beat("What it is.","Distribution = how a root's occurrences spread across the 114 surahs; concentration = how unevenly they pile up.")
beat("Why we do it.","Emphasis can be broad (a theme everywhere) or pooled (a topic of a few surahs); the raw count hides which.")
beat("How it's done.","Take the root's per-surah counts — a vector across 114 surahs (the same vector idea from Week 1) — and summarize its spread and its inequality.")
beat("What we get.","Two new lenses: breadth (reach) and concentration (inequality).")
beat("Why it matters.","Broad versus pooled changes the interpretation entirely.")
beat("In the data.","ظلم's 290 ayahs fall across 59 surahs — and not evenly: its three busiest surahs alone hold 21.7% of them.")
beat("Takeaway.","Distribution is the shape of presence; concentration is its inequality.")
beat("Bridge.","Begin with the simplest shape measure — breadth.")

module("Module 2 · Breadth — in how many surahs","9–14 min")
beat("What it is.","The count of distinct surahs (out of 114) in which a root appears.")
beat("Why we do it.","It is the crudest, most intuitive measure of reach.")
beat("How it's done.","Count the surahs containing at least one ayah of the root.")
beat("What we get.","A reach number from 0 to 114.")
beat("Why it matters.","It separates ubiquitous roots from niche ones, independent of total count.")
beat("In the data.","Narrow: عسر 9 surahs, رشد 9. Broad: كفر 77, علم 85. A root's reach is not predicted by its frequency.")
beat("Takeaway.","Breadth is reach, not amount — a rare root can be wide, a frequent root narrow.")
beat("Bridge.","Breadth can't tell even spread from lumpy — that needs concentration.")

module("Module 3 · Concentration — Lorenz, Gini, top-3 share","14–21 min")
beat("What it is.","How unequally a root's occurrences pile into a few surahs.")
beat("Why we do it.","Two roots of equal breadth can be evenly spread or dominated by one surah.")
beat("How it's done.","Read the top-3 share (intuitive), look at the Lorenz curve (the picture), and report the Gini coefficient — one number from 0 (perfectly even) to 1 (all in one surah). The app computes the Gini; we read it, not hand-calculate it.")
beat("What we get.","A 0–1 inequality score plus an intuitive percentage.")
beat("Why it matters.","Concentration tells you whether a root is a whole-Qur'an theme or a few-surahs topic.")
para([("In the data — a concentrated root vs a spread root:",True)],size=10.5,after=3)
figure(os.path.join(WK,"fig_concentration.png"),
       "Figure 1. Lorenz curves. rushd (رشد) bows hard to the corner — Gini 0.95, only 9 surahs, top-3 share 57.9% (concentrated). kufr (كفر) stays nearer the diagonal — Gini 0.69, 77 surahs (spread).",width=4.6)
beat("Takeaway.","High Gini = pooled in a few surahs; low Gini = woven throughout. Concentration is not importance.")
beat("Bridge.","Now the sharpest question — which single surah is a root's home?")

module("Module 4 · The home surah — normalize by ROOT-TOKENS, not ayahs","21–30 min")
beat("What it is.","The surah where a root is densest — measured size-true, as a share of that surah's root-tokens.")
beat("Why we do it.","Two traps stack here. First, the surah with the most raw hits is usually just the longest surah. Second — and this is the subtle one — dividing by the surah's AYAH count is still wrong, because ayahs vary in length just as surahs do; a surah of long ayahs has more room per verse.")
beat("How it's done.","Use prevalence = the root's tokens in the surah ÷ the surah's total root-tokens × 1,000 (per 1,000 root-tokens), then take the maximum — the same size-true unit as Week 1, now applied within a surah.")
beat("What we get.","A home that is robust to both surah length and ayah length.")
beat("Why it matters.","The answer can change at every level of normalization — raw, per-ayah, and per-roots may each name a different surah.")
figure(os.path.join(WK,"fig_normalization_levels.png"),
       "Figure 2. ṣabr (صبر): the home surah moves at every level — raw → al-Baqara (longest surah), per-AYAH → al-Kahf, per-ROOT-TOKENS → at-Tur. Only the per-root-tokens answer is size-true.",width=6.4)
figure(os.path.join(WK,"fig_home_flip.png"),
       "Figure 3. ẓulm (ظلم): raw busiest surah is al-Baqara (most ayahs AND most tokens); size-true, the home is Ibrahim at 15.8 per 1,000 root-tokens.",width=6.0)
beat("In the data.","ظلم: raw al-Baqara → size-true home Ibrahim (15.8/1k root-tokens). صبر: raw al-Baqara → per-ayah al-Kahf → per-roots at-Tur — three different surahs.")
beat("Takeaway.","Normalize by root-tokens, never by ayah-count: ayahs are no more equal than surahs.")
beat("Bridge.","How often does length fool us across the whole corpus? That is the headline.")

module("Module 5 · The headline finding & the “unlearn”","30–35 min")
beat("What it is.","A corpus-wide test of the length confound.")
beat("Why we do it.","To see whether al-Baqara's apparent dominance is real or a size artifact.")
beat("How it's done.","Take the 50 most frequent roots; find each one's raw busiest surah, its per-ayah home, and its size-true (per-root-tokens) home.")
beat("What we get.","Raw busiest surah = al-Baqara for 30 of the 50; per-ayah, 0 of 50; per-root-tokens, 0 of 50.")
beat("Why it matters.","The single longest surah masquerades as the home of most frequent roots — and normalizing dethrones it completely.")
beat("In the data.","30/50 → 0/50. Both normalizations remove al-Baqara, but they do not always agree with each other (recall صبر: al-Kahf vs at-Tur) — and only the per-root-tokens answer is size-true.")
beat("Takeaway (the unlearn).","“This root belongs to al-Baqara” is a length illusion; and even per-ayah normalization can mislead — trust per-root-tokens.")
beat("Bridge.","But normalization opens a new trap — tiny surahs.")

module("Module 6 · The support floor — small-sample reliability","35–39 min")
beat("What it is.","A minimum-evidence rule before we trust a size-true home.")
beat("Why we do it.","A tiny surah can show a sky-high prevalence off just one or two tokens.")
beat("How it's done.","Require count ≥ 3 in the surah AND the surah ≥ 30 root-tokens. If no surah qualifies, report “insufficient support” rather than a number.")
beat("What we get.","Protection against small-sample noise.")
beat("Why it matters.","Without a floor, the rarest roots get the most confident — and most wrong — homes.")
figure(os.path.join(WK,"fig_support_floor.png"),
       "Figure 4. ʿusr (عسر). Its highest per-root-tokens prevalence is ash-Sharh — 2 tokens in 16 = 125 per 1k — but every surah it appears in fails the floor. Honest verdict: ʿusr has no reliable home surah.",width=6.0)
beat("In the data.","عسر's apparent “home,” ash-Sharh (2 of 16 root-tokens), is an artifact; with the floor applied, عسر has insufficient support for any home.")
beat("Takeaway.","A rate from too little data is a guess; sometimes the honest answer is “not enough evidence.”")
beat("Bridge.","With breadth, concentration, and a guarded home in hand — what does distribution still NOT tell us?")

module("Module 7 · Advantages, limits & what distribution loses","39–43 min")
beat("What it is.","An honest ledger of what distribution and concentration give and what they miss.")
beat("Why we do it.","A method is only usable if you know its blind spots.")
beat("How it's done.","Weigh the advantages — reveals a root's geography, is size-true (per root-tokens), and separates broad from pooled — against the limits.")
beat("What we get.","Criteria for when a distribution claim is safe.")
beat("Why it matters.","Concentration is not importance, and spread is not meaning.")
beat("In the data.","كفر spreading across 77 surahs says it is a pervasive theme — not that it is the most “central” root; and distribution still says nothing about which roots travel together (Weeks 3 onward).")
beat("Takeaway.","Distribution maps where a root lives; it does not rank importance or reveal relationships.")
beat("Bridge.","So we pair every distribution number with a labeled reading — the discipline.")

module("Module 8 · Fact vs. interpretation — the discipline & wrap","43–47 min")
beat("What it is.","The two-sentence discipline applied to a distribution claim, and the term's through-line.")
beat("Why we do it.","To keep computed geography apart from how we read it.")
beat("How it's done.","One sentence of fact (with its size-true normalization and the support floor), one separately labeled sentence of interpretation.")
beat("What we get.","An auditable distribution reading a peer can check.")
beat("Why it matters.","The course is graded on keeping these apart; a community, too, can be read by where its key terms cluster.")
beat("In the data.","Fact: “ظلم spans 59 surahs (Gini 0.74); its size-true home is Ibrahim at 15.8 per 1,000 root-tokens — not al-Baqara, which only leads on raw count because it is the longest surah.”  Interpretation: “I read ظلم as a whole-Qur'an concern, not the topic of one surah.”")
beat("Takeaway.","Where a root lives is a size-true (per-root-tokens), support-checked claim — never its raw busiest surah.")
beat("Bridge.","Distribution says where; next week asks who with — a root's partners and its morphological forms.")
doc.save(OUT); print("saved",OUT,"| paragraphs:",len(doc.paragraphs))
