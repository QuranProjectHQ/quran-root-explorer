# -*- coding: utf-8 -*-
"""Universal Persian-paper normalizer: B Nazanin complex-script font, JUSTIFY body
(keep centered titles/captions), bidi paragraphs + rtl runs. Idempotent."""
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
PERSIAN="Arial"
def ensure_bidi(p):
    pPr=p._p.get_or_add_pPr()
    if pPr.find(qn('w:bidi')) is None: pPr.append(OxmlElement('w:bidi'))
def ensure_rtl(r):
    rPr=r._r.get_or_add_rPr()
    if rPr.find(qn('w:rtl')) is None:
        e=OxmlElement('w:rtl'); e.set(qn('w:val'),'1'); rPr.append(e)
def set_cs(r):
    rPr=r._r.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.insert(0,rf)
    rf.set(qn('w:cs'),PERSIAN); rf.set(qn('w:eastAsia'),PERSIAN)
def fix_para(p):
    ensure_bidi(p)
    if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: set_cs(r); ensure_rtl(r)
def fix_doc(path):
    d=Document(path)
    for p in d.paragraphs: fix_para(p)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs: fix_para(p)
    d.save(path); return path
for f in sys.argv[1:]:
    try: fix_doc(f); print("fixed:",f)
    except Exception as e: print("ERR",f,e)
