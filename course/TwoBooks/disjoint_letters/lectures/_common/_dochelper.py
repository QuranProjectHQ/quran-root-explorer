# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55); TEAL=RGBColor(0x0E,0x6E,0x63); RED=RGBColor(0xA2,0x3B,0x3B)
def newdoc(footer):
    doc=Document()
    stl=doc.styles["Normal"]; stl.font.name="Arial"; stl.font.size=Pt(11)
    rf=stl.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
    stl.paragraph_format.space_after=Pt(4); stl.paragraph_format.line_spacing=1.05
    sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(0.9); sec.top_margin=sec.bottom_margin=Inches(0.8)
    ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=ftr.add_run(footer+" · page "); r.font.size=Pt(8); r.font.color.rgb=GREY
    run=ftr.add_run()
    for t in ("begin","instr","separate","end"):
        e=OxmlElement("w:instrText") if t=="instr" else OxmlElement("w:fldChar")
        if t=="instr": e.set(qn("xml:space"),"preserve"); e.text="PAGE"
        else: e.set(qn("w:fldCharType"),t)
        run._r.append(e)
    return doc
def cs(run):
    rPr=run._r.get_or_add_rPr(); rff=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rff.set(qn(a),"Arial")
def P(doc,segs,size=11,after=4,before=0,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if isinstance(segs,str): segs=[(segs,False)]
    for tx,b in segs:
        rr=p.add_run(tx); rr.bold=b; rr.font.size=Pt(size); cs(rr)
        if color: rr.font.color.rgb=color
    return p
def H(doc,t,size=15,color=ACCENT,before=10,after=4): return P(doc,[(t,True)],size=size,after=after,before=before,color=color)
def bullet(doc,segs,size=10.5,after=3):
    p=doc.add_paragraph(style=None); p.paragraph_format.left_indent=Inches(0.28); p.paragraph_format.space_after=Pt(after)
    r0=p.add_run("•  "); r0.font.size=Pt(size); cs(r0)
    if isinstance(segs,str): segs=[(segs,False)]
    for tx,b in segs:
        rr=p.add_run(tx); rr.bold=b; rr.font.size=Pt(size); cs(rr)
    return p
def table(doc,rows,widths=None,head=True,size=9.5):
    t=doc.add_table(rows=0,cols=len(rows[0])); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""
            pp=cells[ci].paragraphs[0]; pp.paragraph_format.space_after=Pt(1)
            rr=pp.add_run(str(val)); rr.font.size=Pt(size); cs(rr)
            if head and ri==0:
                rr.bold=True; rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),"1F4E79"); cells[ci]._tc.get_or_add_tcPr().append(sh)
            if widths and ci<len(widths): cells[ci].width=Inches(widths[ci])
    return t
