# -*- coding: utf-8 -*-
import os, glob, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
NAVY=RGBColor(0x1E,0x27,0x61); INK=RGBColor(0x1c,0x1c,0x1c); GREY=RGBColor(0x60,0x60,0x60)
ROOT="/sessions/stoic-serene-wozniak/mnt/Downloads/RootCourse/TwoBooks"
LATIN="Liberation Serif"; ARAB="DejaVu Sans"; HEAD="Liberation Sans"
PERSIAN="Arial"   # user-chosen Persian font (universal, has Persian glyphs)
FA=False              # set per-document in build(); routes cs font to PERSIAN for Persian papers

def set_fonts(run, latin=LATIN):
    rPr=run._r.get_or_add_rPr(); rf=rPr.get_or_add_rFonts()
    rf.set(qn("w:ascii"),latin); rf.set(qn("w:hAnsi"),latin)
    cs = PERSIAN if FA else ARAB
    rf.set(qn("w:cs"),cs); rf.set(qn("w:eastAsia"),cs)
def make_rtl(p):
    pPr=p._p.get_or_add_pPr(); pPr.append(OxmlElement("w:bidi"))
def run_rtl(run):
    rPr=run._r.get_or_add_rPr(); el=OxmlElement("w:rtl"); el.set(qn("w:val"),"1"); rPr.append(el)
def add_bottom_border(p,color="1E2761",sz="6",space="4"):
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),sz); bot.set(qn("w:space"),space); bot.set(qn("w:color"),color)
    pbdr.append(bot); pPr.append(pbdr)
def add_page_numbers(section):
    ftr=section.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=ftr.add_run()
    for t in ("begin","instr","sep","end"):
        if t=="instr":
            e=OxmlElement("w:instrText"); e.set(qn("xml:space"),"preserve"); e.text="PAGE"
        else:
            e=OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),{"begin":"begin","sep":"separate","end":"end"}[t])
        run._r.append(e)
    run.font.size=Pt(9); run.font.color.rgb=GREY; set_fonts(run,HEAD)

FIGCAP={
"biology":["Figure 1. Combinatorial amplification (Pillar 1): a tiny alphabet yields a vast lexicon in both Books (log scale).",
 "Figure 2. Order beats chance (Pillar 2): the Qur'anic anagram-sibling rate (54.2%) exceeds a frequency-matched null at p<0.003.",
 "Figure 3. Expression (Pillar 3): surface forms per root, mean 4.7 (heavy-tailed).",
 "Figure 4. Pillar 3: only 9.4% of possible three-letter combinations are realized as roots.",
 "Figure 5. Order makes meaning: number of distinct roots generated from a single letter-set (anagram families).",
 "Figure 6. Fidelity differs (audit X): biological copying error rates vs the claim of perfect preservation.",
 "Figure 7. Genome side: units expand from bases to proteins (log scale).",
 "Figure 8. The Qur'anic corpus by the numbers (log scale).",
 "Figure 9. The 54.2% anagram rate against a 20,000-draw frequency-matched null (p<0.003).",
 "Figure 10. Heavy-tailed (Zipf-like) root-frequency spectrum.",
 "Figure 11. Expression echo: ~4.7 forms per Qur'anic root vs ~4-5 protein isoforms per gene."],
"signal":["Figure 1. Verse 112:1 as a root-frequency signal: [1722, 2848, 153].",
 "Figure 2. The root anchor cleans the null: mean correlation of random verse pairs 0.18 (surface) -> 0.04 (roots).",
 "Figure 3. Autocorrelation of Surat ar-Rahman's refrain: a period-2 signature (+0.75 at lag 2).",
 "Figure 4. Exact-verse repetition beats the natural-language baseline (7.1% vs 0.81%, ~8.8x).",
 "Figure 5. Ayah-length distribution (median 7 root-tokens, heavy tail).",
 "Figure 6. Top root frequencies (transliterated) - heavy-tailed.",
 "Figure 7. Scale rule: the lag-1 autocorrelation estimate stabilizes only with length.",
 "Figure 8. Root null: 20,000 random 7-root verse pairs (mean r ~ 0.04).",
 "Figure 9. Fourier spectrum of the refrain indicator: a line at period ~ 2 verses.",
 "Figure 10. Embeddings recover meaning: nearest-root cosine (RHM->GHFR, etc.).",
 "Figure 11. Generic to Arabic: Qur'anic Zipf slope and top-10 share match the baseline."],
"disjoint_letters":["Figure 1. Disjoint-letter families cluster in mushaf order (label-permutation p=2x10^-5).",
 "Figure 2. The Hawamim (HM) were revealed in seven consecutive slots (60-66).",
 "Figure 3. Disjoint-letter suras are the long ones (median 85 vs 26 verses, p=2x10^-5).",
 "Figure 4. The frequency claim collapses under the cross-sura baseline (own ~ others; 0/29).",
 "Figure 5. Per-family contiguity significance (-log10 p).",
 "Figure 6. Revelation-phase mapping: mean nuzul slot by tag (simple early, families late).",
 "Figure 7. Label-permutation null vs the observed within-family distance (p=2x10^-5).",
 "Figure 8. No content link by tag: within-family ~ cross-family root similarity (p=0.27).",
 "Figure 9. Disjoint-letter family sizes across the 29 suras.",
 "Figure 10. Same-tag families also cluster in revelation order.",
 "Figure 11. The frequency claim: spectacular under a weak null, zero under the right baseline."]}
TABLE={
"biology":("Table 1. The three-pillar thesis, dual-domain.",[["Pillar","Qur'an (computed)","Genome (textbook)"],["1 generativity","27 letters -> 1,702 roots -> 7,236 words","4 bases -> 64 codons -> ~90,000 proteins"],["2 order","54.2% of roots have an anagram sibling (p<0.003)","frameshift / point change alters function"],["3 expression","9.4% of triplets realized; 4.7 forms/root","~4-5 isoforms/gene; differential expression"]]),
"signal":("Table 1. Key results and the validation gauntlet.",[["Claim","Outcome"],["Verse-level content (freq, length, entropy)","generic to Arabic X"],["Surface vs root null (mean r)","0.18 -> 0.04"],["Exact-verse repetition (refrains)","7.1% vs 0.81%, ~8.8x (p~0.03) OK"],["Ar-Rahman period","autocorr +0.75 lag 2; FFT period 2.05 OK"]]),
"disjoint_letters":("Table 1. Summary of tests (all 29 muqatta'at suras).",[["Test","Result"],["Tag -> mushaf contiguity","p = 2x10^-5 OK"],["Tag -> revelation-order contiguity","p = 2x10^-5 OK"],["Disjoint-letter suras are long","median 85 vs 26, p = 2x10^-5 OK"],["Shared content/theme by tag","p = 0.27 (none) X"],["Frequency-enrichment claim","0/29 (refuted) X"]]),
}
def topic_of(p): return "biology" if "/biology/" in p else ("signal" if "/signal/" in p else "disjoint_letters")
def build(md_path):
    global FA
    fa="Persian" in md_path; FA=fa; topic=topic_of(md_path)
    text=open(md_path,encoding="utf-8").read().split("\n")
    doc=Document(); st=doc.styles["Normal"]; st.font.size=Pt(11); st.font.name=LATIN
    sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(1.0); sec.top_margin=sec.bottom_margin=Inches(1.0)
    add_page_numbers(sec)
    def para(txt,size=11,bold=False,color=INK,align=None,italic=False,sa=8,sb=0,latin=LATIN):
        p=doc.add_paragraph()
        if fa: make_rtl(p)
        p.alignment=align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY  # body justified (RTL when fa)
        p.paragraph_format.space_after=Pt(sa); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.line_spacing=1.0
        for seg in re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)",txt):
            if not seg: continue
            b=bold; it=italic; s=seg
            if seg.startswith("**") and seg.endswith("**"): b=True; s=seg[2:-2]
            elif seg.startswith("*") and seg.endswith("*") and len(seg)>2: it=True; s=seg[1:-1]
            r=p.add_run(s); r.font.size=Pt(size); r.font.bold=b; r.font.italic=it; r.font.color.rgb=color; set_fonts(r,latin)
            if fa: run_rtl(r)
        return p
    first=True
    for line in text:
        ln=line.rstrip()
        if not ln.strip(): continue
        if ln.startswith("# "):
            t=ln[2:].strip()
            if first:
                p=para(t,size=20,bold=True,color=NAVY,align=WD_ALIGN_PARAGRAPH.CENTER,sa=10,latin=HEAD); add_bottom_border(p,sz="12",space="6"); first=False
            else:
                p=para(t,size=14,bold=True,color=NAVY,align=(WD_ALIGN_PARAGRAPH.RIGHT if fa else None),sb=12,sa=5,latin=HEAD); add_bottom_border(p)
        elif ln.startswith("## "):
            p=para(ln[3:].strip(),size=14,bold=True,color=NAVY,align=(WD_ALIGN_PARAGRAPH.RIGHT if fa else None),sb=12,sa=5,latin=HEAD); add_bottom_border(p)
        elif ln.startswith("### "):
            para(ln[4:].strip(),size=12,bold=True,color=NAVY,align=(WD_ALIGN_PARAGRAPH.RIGHT if fa else None),sb=8,sa=3,latin=HEAD)
        elif ln.startswith("**") and ln.endswith("**") and len(ln)<200:
            para(ln,size=11.5,italic=True,color=GREY,align=(WD_ALIGN_PARAGRAPH.RIGHT if fa else WD_ALIGN_PARAGRAPH.CENTER))
        else:
            para(ln)
    # table
    cap,rows=TABLE[topic]; pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run(cap); rc.font.bold=True; rc.font.size=Pt(10); rc.font.color.rgb=NAVY; set_fonts(rc,HEAD)
    tb=doc.add_table(rows=0,cols=len(rows[0])); tb.style="Light Grid Accent 1"; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,row in enumerate(rows):
        cells=tb.add_row().cells
        for ci,val in enumerate(row):
            cells[ci].text=""; rr=cells[ci].paragraphs[0].add_run(str(val)); rr.font.size=Pt(9.5); set_fonts(rr); rr.font.bold=(ri==0)
    doc.add_paragraph()
    hp=doc.add_paragraph(); hr=hp.add_run("Figures"); hr.font.bold=True; hr.font.size=Pt(14); hr.font.color.rgb=NAVY; set_fonts(hr,HEAD)
    figs=sorted(glob.glob(os.path.join(ROOT,topic,"figs","*.png")))
    for i,fp in enumerate(figs):
        doc.add_picture(fp,width=Inches(5.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cr=cp.add_run(FIGCAP[topic][i] if i<len(FIGCAP[topic]) else "Figure %d."%(i+1)); cr.font.size=Pt(9); cr.font.italic=True; cr.font.color.rgb=GREY; set_fonts(cr,HEAD)
        cp.paragraph_format.space_after=Pt(10)
    out=md_path.replace(".md",".docx"); doc.save(out); return out
mds=sorted(glob.glob(ROOT+"/*/PAPER_*.md")); built=0
for m in mds:
    wc=len(open(m,encoding="utf-8").read().split())
    try:
        build(m); built+=1; print("docx:",os.path.basename(m),"(%d words)"%wc)
    except PermissionError:
        print("LOCKED:",os.path.basename(m).replace(".md",".docx"))
print("built",built)
