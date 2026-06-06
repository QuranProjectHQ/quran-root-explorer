# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
WK="/sessions/kind-compassionate-feynman/mnt/RootCourse/week05"; FD=os.path.join(WK,"figs")
OUT=os.path.join(WK,"Week5_Lecture_Notes.docx")
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55)
doc=Document()
stl=doc.styles["Normal"]; stl.font.name="Arial"; stl.font.size=Pt(11)
rf=stl.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
stl.paragraph_format.space_after=Pt(4); stl.paragraph_format.line_spacing=1.04
sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(0.9); sec.top_margin=sec.bottom_margin=Inches(0.8)
ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ftr.add_run("Week 5 — Lift & Tiers · Lecture Notes (v1) · page "); r.font.size=Pt(8); r.font.color.rgb=GREY
run=ftr.add_run()
for t in ("begin","instr","separate","end"):
    e=OxmlElement("w:instrText") if t=="instr" else OxmlElement("w:fldChar")
    if t=="instr": e.set(qn("xml:space"),"preserve"); e.text="PAGE"
    else: e.set(qn("w:fldCharType"),t)
    run._r.append(e)
def cs(run):
    rPr=run._r.get_or_add_rPr(); rff=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rff.set(qn(a),"Arial")
def para(segs,size=11,after=4,before=0,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=1.04
    if isinstance(segs,str): segs=[(segs,False)]
    for tx,b in segs:
        rr=p.add_run(tx); rr.bold=b; rr.font.size=Pt(size); cs(rr)
        if color: rr.font.color.rgb=color
    return p
def beat(label,text): para([(label+"  ",True),(text,False)],size=10.5,after=4)
def module(title,minutes): para([(title,True),("   ("+minutes+")",False)],size=13,before=11,after=4,color=ACCENT)
def figure(name,cap,width=5.8):
    p=os.path.join(FD,name)
    if os.path.exists(p):
        doc.add_picture(p,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        para([(cap,False)],size=8.5,after=5,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER)

para([("Week 5 — Lift & Tiers",True)],size=20,after=2,color=ACCENT)
para([("Lecture Notes · v1 — deep (with analogies)",True)],size=12,after=2,color=GREY)
para("A 45-minute lecture, eight modules, each carrying the eight beats. Week 4 ranked a target's companions; this week we put each bond on trial and deliver a VERDICT — is it real? We sharpen co-occurrence with a length-aware null, test it with Monte-Carlo simulation, and sort every pair into a tier: structural, borderline, or spurious. Worked pair: صلو ↔ زكو (prayer ↔ zakat). Scope: lift, the null, significance, and tiers; DIRECTION — which root leads — is Week 6, named here only as a preview. Every value computed from Book6.",size=9.5,after=8,color=GREY)

module("Module 1 · Opening & recap — from ranking to judging","0–5 min")
beat("What it is.","Last week we ranked a target's candidate companions and learned to trust the controlled ranking over the raw count. This week we stop ranking and start JUDGING: for a single pair of roots, is the bond between them REAL, or could it be an accident?")
beat("Why we do it.","A ranking tells you who is closest; it does not tell you whether even the closest is anything more than chance. Two roots can top each other's lists and still be strangers who happen to share a few crowded verses — the way two names can keep appearing in the same news articles simply because both are in the news, not because they are connected. A ranking is a leaderboard; what we need now is a verdict, the way a court does not merely rank suspects but decides guilt.")
beat("How it's done.","We move from a single ratio to a courtroom. We state the strength of the bond (its lift), we test that strength against the possibility of coincidence (a simulation that asks how often chance alone would do this), and we deliver a verdict (a tier). Strength of evidence, test against the null hypothesis, then a labelled judgment — accusation, cross-examination, verdict — it is the logic of a trial, applied to a pair of words. And like a trial, it can acquit: most pairs are not guilty of any real bond, and saying so plainly is as much a finding as convicting the few that are.")
beat("What we get.","For any pair, three things that together form a judgment: a lift (how much stronger than chance), a significance (how sure we can be it is not coincidence), and a tier label — structural, borderline, or spurious — that condenses both into a word a reader can act on.")
beat("Why it matters.","This is where co-occurrence becomes trustworthy, and trustworthiness is not optional. Without it, the back half of the course — clusters, hubs, three-root motifs — would be a tower built on bonds we never checked, and a single spurious bond near the base would tilt everything above it. Better to find the rot now, with one pair at a time, than to discover it after we have drawn a whole network.")
beat("In the data.","Our worked pair is صلو ↔ زكو, prayer and almsgiving, which shared 28 ayahs at a high controlled ratio in Week 4. By the end we'll have its verdict — and the verdict of a pair that shares FOUR TIMES as many ayahs yet means nothing.")
beat("Takeaway.","Ranking asks who is closest; judging asks whether closeness is real. The first is a comparison; the second is a verdict — and only the verdict can be trusted on its own.")
beat("Bridge.","To judge fairly, we first have to admit two problems Week 4 left unsolved.")

module("Module 2 · Two leftover problems","5–11 min")
beat("What it is.","An honest accounting of where last week's simple control still falls short — because a method's gaps, named plainly, are the map of what the new machinery must repair.")
beat("Why we do it.","Naming the gaps is what tells us exactly what the new machinery has to fix.")
beat("How it's done.","Lay the two problems side by side.")
figure("fig_two_problems.png","Figure 1. Week 4 ranked bonds, but two problems remain: long-ayah inflation, and big shared counts that are nonetheless chance.",width=5.6)
beat("What we get.","A precise to-do list for the week: fix the long-ayah inflation with a fairer baseline, and fix the chance-disguised-as-size problem with a test of significance.")
beat("Why it matters.","Problem one: roots cluster in LONG ayahs. A verse packed with twenty roots is like a crowded party where everyone brushes shoulders with everyone; a three-root verse is a quiet room where being together actually means something. Two roots that both happen to favour long, crowded verses will share many of them with no real bond at all — and last week's flat expectation never knew the difference between the party and the quiet room. Problem two: a big shared count can still be pure chance. Two ubiquitous roots will share dozens of ayahs simply because each is nearly everywhere; the impressive-looking count is just the arithmetic of two large numbers overlapping, and it means nothing about a bond.")
beat("In the data.","قول (say) and شيء (thing) share 113 ayahs — more than any bond in our entire set, and more than three times our worked pair — yet both are among the commonest roots in the language, so 113 is no more than the overlap you'd expect by chance alone. The most-shared pair in the room is the emptiest.")
beat("Takeaway.","Length inflates bonds, and size disguises chance; the raw and even the simple-controlled numbers can't see either.")
beat("Bridge.","Fix the first problem with a smarter baseline — the length-aware null.")

module("Module 3 · The length-aware null — a fairer baseline","11–18 min")
beat("What it is.","An expectation that already knows roots cluster in long ayahs, so it raises the bar accordingly.")
beat("Why we do it.","If the baseline accounts for length, then any overlap ABOVE it is real signal, not a long-ayah artefact.")
beat("How it's done.","Instead of the flat expectation of Week 4, we compute the overlap we'd expect if roots were scattered through the book in proportion to ayah LENGTH — long verses more likely, by their sheer size, to hold any given root. Lift is then the observed overlap divided by this fairer expectation. Two analogies, both useful. It is a handicap in golf: we adjust each score for the difficulty of the course before we compare players, so that a good round on a hard course is not beaten by a mediocre round on an easy one. It is also grading on a curve that already knows some exams were easier: a high mark on the hard exam counts for more.")
beat("What we get.","A lift that is honest about length — and, because it now subtracts the crowding that long verses hand out for free, it is always somewhat lower than last week's simple ratio. How much lower is itself informative: a bond that barely changes lived in short verses; a bond that collapses lived only in the crowd.")
figure("fig_lift_deflate.png","Figure 2. The length-aware null deflates every lift, because roots really do cluster in long ayahs — but the strong bonds stay strong.",width=6.0)
beat("Why it matters.","Every bond shrinks a little under the fairer baseline, which is exactly as it should be: some part of last week's ratio was always just long-ayah crowding, credit the bond never earned. صلو ↔ زكو falls from ×34.6 (the Week-4 simple ratio) to ×23.6 under the length-aware null — and ×23.6 is still enormous, twenty-three times what chance would manage. Real bonds survive the handicap with room to spare; the deflation only trims the unearned part, and a bond that vanishes under a fair handicap was never real to begin with.")
beat("In the data.","كيل ↔ وزن goes ×137 → ×93; عهد ↔ نقض ×96 → ×66; جنن ↔ نهر ×13 → ×9. The ranking barely changes; the inflation is trimmed.")
beat("Takeaway.","A fair baseline lowers every lift a little and leaves the real bonds standing.")
beat("Bridge.","Lift tells us how surprising a bond is; now we ask how CONFIDENT we can be.")

module("Module 4 · Monte-Carlo significance — testing against chance","18–25 min")
beat("What it is.","A simulation that asks one disciplined question: if there were no real bond between these two roots — if they were merely scattered through the book under its real constraints — how often would blind chance alone produce an overlap as large as the one we actually see?")
beat("Why we do it.","Lift measures size; significance measures certainty, and the two can disagree. A huge lift resting on only two or three shared ayahs could still be a fluke — a coincidence dressed up as a discovery by a small denominator — and the only way to know is to ask chance directly. A bond worth trusting must be both large and unlikely-by-accident.")
beat("How it's done.","We reshuffle. Keeping each ayah's length fixed and each root's overall frequency fixed — so every root still appears exactly as often, and every verse is still as long — we scatter the roots across the verses at random, then count how many verses our two roots now share. Do that three thousand times and you have a full picture of what blind chance produces under exactly the corpus's real constraints. The analogy is a banquet: take the seating chart of a great hall, keep the table sizes and the guest list fixed, and reshuffle who sits where three thousand times. If two particular guests keep ending up at the same table far more often than any reshuffle manages, they are not there by luck — they are choosing each other.")
figure("fig_montecarlo.png","Figure 3. Monte-Carlo for صلو ↔ زكو: across 3,000 reshuffles, chance produces 0–3 shared ayahs; the observed 28 is off the chart (p < 0.001).",width=5.6)
beat("What we get.","A p-value: the fraction of shuffles that matched or beat the observed overlap. A small p — say, below one in a thousand — means chance, given the corpus's real constraints, almost never produces what we actually see, so the bond is real. A large p means chance produces it constantly, so we have learned nothing. The p-value is, quite literally, a measure of how surprised we should be.")
beat("Why it matters.","For صلو ↔ زكو, three thousand reshuffles never come close to 28 shared verses; chance, under the real constraints, gives between zero and about three. So the p-value — the share of reshuffles that match or beat the real overlap — is essentially zero: chance simply does not do this. But for قول ↔ شيء, the reshuffles routinely produce 113 shared verses or more, because two such common roots overlap heavily no matter how you scatter them; chance does this all the time, so the p-value is about 0.99 — almost every reshuffle is as “impressive” as the real data.")
beat("In the data.","صلو ↔ زكو: p < 0.001. قول ↔ شيء: p ≈ 0.99. علم ↔ رحم (36 shared): p ≈ 0.99. Even صلو ↔ ءله (44 shared) fails, at p ≈ 0.06.")
beat("Takeaway.","Lift says how big; the p-value says how sure — and a big overlap with a big p-value is just noise.")
beat("Bridge.","Two numbers — lift and significance — combine into one verdict: the tier.")

module("Module 5 · The tiers — a verdict, not a number","25–31 min")
beat("What it is.","A three-way classification that turns three quantities — lift, significance, and support — into a single, defensible label a reader can carry forward. It is the moment the analysis stops producing numbers and starts producing judgments.")
beat("Why we do it.","People cannot hold three numbers in their heads while reading a whole table of pairs; a tier is the headline those three numbers earn, the one word that lets a reader move on without losing the verdict. It is the same reason a court returns “guilty” or “not guilty” rather than reciting the evidence each time the case is mentioned.")
beat("How it's done.","Combine the evidence the way a court combines it: not by averaging the numbers, but by requiring each kind of evidence to clear its own bar before a strong verdict is allowed.")
figure("fig_tier_ladder.png","Figure 4. The three tiers and their criteria — structural, borderline, spurious.",width=6.0)
beat("What we get.","Tier 1 — STRUCTURAL: lift ≥ 3 AND p < 0.001 AND at least 5 shared ayahs — proven beyond reasonable doubt. Tier 2 — BORDERLINE: significant but with a modest lift or thin support — the balance of probabilities. Tier 3 — SPURIOUS: not significant, or lift below chance level — not proven.")
beat("Why it matters.","The tier protects you from both of the week's traps at once. A Tier-1 verdict demands real lift AND real significance AND enough shared verses to stand on — three independent hurdles — so a length artefact (which fails on lift) and a lucky handful of verses (which fails on significance or support) are each turned away at a different gate. It is harder to fake a verdict than to fake any single number.")
beat("In the data.","Six of our twelve pairs are Tier 1 (صلو↔زكو, كيل↔وزن, عهد↔نقض, سجد↔ركع, قرض↔حسن, جنن↔نهر); one is Tier 2 (عبد↔رزق — worship and provision, real but modest); five are Tier 3.")
beat("Takeaway.","A tier is a verdict: structural is proven, borderline is plausible, spurious is not proven.")
beat("Bridge.","Now the headline — and the unlearn that the tiers make unmissable.")

module("Module 6 · The headline & the “unlearn”","31–37 min")
beat("What it is.","A side-by-side of the biggest shared count in our set and the bond we most trust.")
beat("Why we do it.","To make permanent the lesson that counting is not judging.")
beat("How it's done.","Put قول ↔ شيء and صلو ↔ زكو on the same axis — by raw count, then by tier.")
figure("fig_unlearn.png","Figure 5. A big shared count is not a bond: قول↔شيء (113 shared) is Tier 3 spurious; صلو↔زكو (28 shared) is Tier 1 structural.",width=6.0)
beat("What we get.","A reversal that should unsettle the intuition for good: the pair with the MOST shared ayahs in our set is spurious, and a pair with barely a quarter as many shared ayahs is the most structural bond we have. If the raw count and the verdict disagree, the verdict wins — every time.")
beat("Why it matters.","قول and شيء share 113 ayahs — more than any pair we have studied — and it means nothing: two of the commonest words in the language will keep landing in the same verse for no reason but their ubiquity. صلو and زكو share 28, and it is one of the deepest bonds in the whole book. The number of shared verses pointed exactly the wrong way: the larger count was the emptier bond. It is the difference between the 113 acquaintances whose faces you pass in a crowded market and the 28 friends who would come to your home — and if you counted only heads, you would mistake the crowd for the friendship.")
figure("fig_verdict_card.png","Figure 6. The verdict for the worked pair — صلو ↔ زكو: lift ×23.6, p < 0.001, Tier 1 structural.",width=5.0)
beat("In the data.","قول↔شيء: 113 shared, lift ×0.8, p ≈ 0.99 → spurious. صلو↔زكو: 28 shared, lift ×23.6, p < 0.001 → structural.")
beat("Takeaway (the unlearn).","A big shared count is not a bond; the verdict comes from lift and significance together, never from the raw number of shared verses.")
beat("Bridge.","How do we know our tier thresholds are set right? We calibrate.")

module("Module 7 · Calibration & limits","37–43 min")
beat("What it is.","Checking the tier rule against pairs whose answer we already know — the way you check a new thermometer against boiling and freezing water before you trust it on a fever — and then naming plainly what the verdict still cannot do.")
beat("Why we do it.","A classifier you never check is a classifier you can't trust; and a measure whose limits you ignore will mislead you.")
beat("How it's done.","Run the whole twelve-pair set through the rule and confirm that the bonds we already know are real — prayer and zakat, measure and weight — land in Tier 1, while the pairs we know are generic — say and thing, knowledge and mercy — land in Tier 3. If a known-real bond came out spurious, or a known-generic pair came out structural, the thresholds would be wrong and we would move them; here they hold.")
figure("fig_calibration_table.png","Figure 7. The 12-pair calibration: known-real bonds land Tier 1; known-generic pairs (قول↔شيء, علم↔رحم) land Tier 3 — the thresholds hold.",width=6.0)
beat("What we get.","Confidence that the thresholds are neither too loose — letting a lucky or length-inflated pair slip into Tier 1 — nor too strict — demoting a genuine bond to borderline. Calibration is the quiet discipline that keeps a classifier honest: you do not get to choose where the line sits without checking it against cases whose answer you already know.")
figure("fig_scatter.png","Figure 8. Every pair placed by lift and significance: real bonds (teal) sit top-right, spurious (red) bottom-left, borderline (amber) between.",width=5.8)
beat("Why it matters.","The limits are real. The verdict tells you a bond is real, but not WHICH root leads the other — whether mentioning prayer pulls in zakat or the reverse; that asymmetry is next week. It tells you a bond is significant, not what it MEANS. And a structural tier is a strong invitation to interpret, never the interpretation itself.")
beat("In the data.","On the lift-versus-significance map, the six structural pairs cluster in the trustworthy top-right; قول↔شيء and علم↔رحم sit in the bottom-left where chance lives; عبد↔رزق hovers on the borderline.")
beat("Takeaway.","Calibrate the verdict against known cases, and remember it judges reality, not direction or meaning.")
beat("Bridge.","So the verdict, like every number this term, is paired with a labeled reading.")

module("Module 8 · Fact vs. interpretation — the discipline & wrap","43–47 min")
beat("What it is.","The two-sentence discipline applied to a tiered bond, and the term's through-line.")
beat("Why we do it.","A Tier-1 verdict is the most persuasive number we've produced — and therefore the easiest to over-read into a sermon.")
beat("How it's done.","One sentence of fact — lift, significance, tier — and one separately labeled sentence of interpretation.")
beat("What we get.","A reading a classmate can audit at a glance, and a defensible, reproducible verdict standing behind it — the statistics on one side of the line, the meaning on the other, and a visible boundary between them.")
beat("Why it matters.","The course is graded on this line, and Week 5 is its hardest test: when the statistics say a bond is structural beyond reasonable doubt, the temptation to present your reading of it as part of the finding is strongest.")
beat("In the data.","Fact: “صلو ↔ زكو share 28 ayahs at lift ×23.6 over a length-aware null, p < 0.001 — Tier 1, structural.”  Interpretation: “I read prayer and almsgiving as inseparable halves of one devotion, the vertical and the horizontal.”")
beat("Takeaway.","Frequency, distribution, forms, partners, co-occurrence, and now a tiered verdict — each computed, each needing a labeled reading; the statistics earn trust, not meaning.")
beat("Bridge.","Next week — direction and networks: which root LEADS the other (P(A|B) vs P(B|A)), and how bonds assemble into hubs.")

doc.save(OUT); print("saved | words:",sum(len(p.text.split()) for p in doc.paragraphs))
