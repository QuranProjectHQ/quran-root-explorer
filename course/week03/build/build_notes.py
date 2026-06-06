# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
WK="/sessions/kind-compassionate-feynman/mnt/RootCourse/week03"; FD=os.path.join(WK,"figs")
OUT=os.path.join(WK,"Week3_Lecture_Notes.docx")
ACCENT=RGBColor(0x1F,0x4E,0x79); GREY=RGBColor(0x55,0x55,0x55)
doc=Document()
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(11)
rf=st.element.get_or_add_rPr().get_or_add_rFonts()
for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rf.set(qn(a),"Arial")
st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.04
sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(0.9); sec.top_margin=sec.bottom_margin=Inches(0.8)
ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ftr.add_run("Week 3 — Partners & Forms · Lecture Notes (v2, expanded) · page "); r.font.size=Pt(8); r.font.color.rgb=GREY
run=ftr.add_run()
for t in ("begin","instr","separate","end"):
    e=OxmlElement("w:instrText") if t=="instr" else OxmlElement("w:fldChar")
    if t=="instr": e.set(qn("xml:space"),"preserve"); e.text="PAGE"
    else: e.set(qn("w:fldCharType"),t)
    run._r.append(e)
def cs(run):
    rPr=run._r.get_or_add_rPr(); rff=rPr.get_or_add_rFonts()
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"): rff.set(qn(a),"Arial")
def para(segs,size=11,after=4,before=0,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=1.04
    if isinstance(segs,str): segs=[(segs,False)]
    for tx,b in segs:
        rr=p.add_run(tx); rr.bold=b; rr.font.size=Pt(size); cs(rr)
        if color: rr.font.color.rgb=color
    return p
def beat(label,text): para([(label+"  ",True),(text,False)],size=10.5,after=4)
def module(title,minutes): para([(title,True),("   ("+minutes+")",False)],size=13,before=11,after=4,color=ACCENT)
def figure(name,cap,width=5.8):
    p=os.path.join(FD,name)
    if os.path.exists(p):
        doc.add_picture(p,width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        para([(cap,False)],size=8.5,after=5,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER)

para([("Week 3 — Partners & Forms",True)],size=20,after=2,color=ACCENT)
para([("Lecture Notes · v2 — expanded (deep, with analogies)",True)],size=12,after=2,color=GREY)
para("A 45-minute lecture, eight modules, each carrying the eight beats — what it is · why · how · what we get · why it matters · in the data · takeaway · bridge. Two new lenses this week: Arabic root-and-pattern morphology (a root's internal company — its forms) and collocation (its external company — its partners). Worked root: ءمن (believe). Scope: forms in depth and partners read descriptively; the co-occurrence MEASURE (which root shares the most ayahs, why raw counts mislead) is Week 4, and lift / the length-aware null / tiers are Week 5. Every value computed from Book6.",size=9.5,after=8,color=GREY)

module("Module 1 · The root-and-pattern system (الجذر والوزن)","0–6 min")
beat("What it is.","Arabic builds words by pouring a three-consonant root into a pattern — a wazn, or template. The root carries the core meaning; the pattern carries the grammatical role. Think of it as a jeweller's workshop: the root is the raw metal, the pattern is the mould you press it into. The same gram of silver becomes a ring, a chain, or a coin depending only on the mould.")
beat("Why we do it.","For two weeks we counted roots, and to count them we did something deliberately violent — we melted every word down to its three consonants and kept only the metal. That let us count concepts, but it threw away the shapes. This week we pick the shapes back up, because the shape of a word tells you whether it names a doer, a deed, an attribute, or a command.")
beat("How it's done.","Take one root and walk it through its moulds. From ك-ت-ب, “to write,” the pattern فاعِل gives كاتب (a writer, the doer); مفعول gives مكتوب (written, the thing done); فِعال gives كِتاب (a book); the present-tense mould gives يكتب (he writes); and the passive past gives كُتِب (it was decreed). One root, one meaning-core, but six different jobs. The mould even carries grammatical voice: فاعِل is the one who acts (كاتب, the writer) while مفعول is the one acted upon (مكتوب, the written) — same root, but who-does-what-to-whom flipped by the pattern alone.")
beat("What we get.","From a single root, a whole family of words spanning agent, object, act, instrument, and even divine decree — a family tree we can read.")
beat("Why it matters.","Meaning is split between the root and the pattern, like a skeleton and the clothing it wears. You cannot read a person's profession from their bones alone, and you cannot read a word's grammatical force from its bare root alone — كُتِب (it was decreed, as in “fasting was prescribed for you”) and كِتاب (a book) share every consonant and could not be further apart in use.")
figure("fig_root_pattern.png","Figure 1. One root (ك-ت-ب) poured into many moulds yields a family of words — writer, written, book, he-writes, was-decreed.",width=5.6)
beat("In the data.","Across our twelve roots this week, a single root wears between 12 and 27 distinct surface forms in the Qur'an — a dozen to two dozen outfits cut from the same cloth.")
beat("Takeaway.","The root is the lexical atom — the WHAT; the pattern is the grammar that assigns its role — the HOW.")
beat("Bridge.","If a root wears many forms, those forms are not noise — they are a distribution, and a distribution can be read.")

module("Module 2 · Reading a form distribution","6–12 min")
beat("What it is.","The share of a root's tokens taken by each surface form — a census of the faces a word wears.")
beat("Why we do it.","Just as a frequency rate told us which roots a text dwells on (Week 1) and distribution told us where they sit (Week 2), a form distribution tells us which FACE of a root the text dwells on — and that is often more revealing than the bare count. Frequency asked which words a text dwells on; this asks which FACE of a word it dwells on — the difference between knowing a person is mentioned often and knowing they are always mentioned as a leader, never as a follower.")
beat("How it's done.","Tally each surface form and rank by share. Book6 makes this exact: every root-token is aligned to its actual word-form, so the counts are not guesses — they are the text itself, sorted.")
beat("What we get.","A profile. For ءمن: the verb آمن at 41%, the participle مؤمنين at 16.5%, the present-tense يؤمن at 16%, and the abstract noun إيمان trailing at 4% — 27 forms in all, across 879 tokens.")
beat("Why it matters.","A root is not used evenly across its forms, any more than a coin lands evenly on its faces. One or two forms carry most of the weight, and which ones they are is a fact about how the concept is handled.")
figure("fig_amn_forms.png","Figure 2. Surface forms of ءمن. The single most common form is the verb آمن (“believed”) — far ahead of the abstract noun إيمان.",width=4.7)
beat("In the data.","ءمن's most common face is the verb آمن — the act of believing — not the noun إيمان we reach for when we translate “faith” into English. Its second face is the people — مؤمنين, the believers — so the root is busiest naming the act and the actors, and least busy naming the abstraction.")
beat("Takeaway.","Read the dominant form before you pronounce on what a root “means.”")
beat("Bridge.","But the raw spellings still hide a deeper layer — the moulds, the patterns, behind them.")

module("Module 3 · Pattern families & the “unlearn”","12–19 min")
beat("What it is.","Grouping a root's many forms into a few pattern families — sorting the outfits by what kind of garment they are, not by their colour.")
beat("Why we do it.","Individual spellings are too fine-grained; the pattern family is where the grammatical meaning actually lives. آمن and يؤمن differ only in tense, but both are the act; مؤمن is a different thing entirely — a person.")
beat("How it's done.","Collapse the forms into families: the verb (آمن، يؤمن — the act, in its tenses and moods), the active participle (مؤمن — the agent, “the believer”), the masdar (إيمان — the act named as an abstract thing), and here a fourth, the security branch (أمن — a second sense the root carries).")
beat("What we get.","For ءمن: 61% verb, 26% participle, 5% masdar, 5% security — a clear, lopsided shape.")
beat("Why it matters.","The shape of that distribution is not decoration; it is an argument about how the concept is held. A verb-heavy root is a concept the text treats as action; a noun-heavy root is one it treats as a thing or an essence.")
figure("fig_amn_patterns.png","Figure 3. ءمن by pattern family — faith is overwhelmingly a verb form (an act), not the abstract noun إيمان.",width=6.2)
beat("In the data.","Faith in the Qur'an is overwhelmingly a VERB — something done — and only about 5% of the time the tidy abstract noun إيمان. The grammar is doing theology: belief is staged as an act before it is a possession. Contrast a root like خلق, “create,” whose forms cluster tightly on the deed itself — different roots are framed by their grammar as actions, agents, or essences, and the form distribution is how you tell which.")
beat("Takeaway (the unlearn).","“Faith” is most often named as something you DO, not something you HAVE — a claim invisible at the root layer, visible only once we reopen the forms.")
beat("Bridge.","One pattern family is special enough to deserve its own module — the intensive mould that names God.")

module("Module 4 · The intensive forms & the Divine Names","19–26 min")
beat("What it is.","The intensive adjectival patterns — فعيل and فعّال — which take a quality and turn its dial up to the maximum.")
beat("Why we do it.","Some of the most important vocabulary in the Qur'an — the Names of God — turns out to be a single grammatical move repeated, and seeing the move lets you read the Names as a family rather than a list to memorise.")
beat("How it's done.","Pour a root into the intensive mould and read the result: mercy (رحم) becomes رحيم; hearing (سمع) becomes سميع; seeing (بصر) becomes بصير; forgiving (غفر) becomes غفور; wisdom (حكم) becomes حكيم. The same root can even yield two intensities — from رحم come both رحيم and the still-stronger رحمن (a related فعلان mould), the pair that opens almost every chapter.")
beat("What we get.","A set of forms that function, across the Qur'an, as Names of God — each one the same intensive pattern pressed over a different root.")
beat("Why it matters.","The Divine Names are not coined ad hoc, one by one; many are the SAME mould applied to many metals. The attributes are grammatically marked as superlatives — the language itself says “to the highest degree.” When you meet a new فعيل form — قدير, خبير, عليم — you can already guess you are looking at an attribute pressed to its maximum, very often a Name.")
figure("fig_divine_names.png","Figure 4. Divine Names as intensive forms of these roots — رحيم, حكيم, غفور, رحمن, سميع, بصير — bars show occurrence counts.",width=6.0)
beat("In the data.","Six roots, one intensive pattern — رحيم, حكيم, غفور, رحمن, سميع, بصير — together among the most frequent names in the text (occurrence counts in Figure 4).")
beat("Takeaway.","Learn one pattern and a whole class of the Divine Names lines up — the grammar is a key, not a wall.")
beat("Bridge.","Patterns also explain something stranger: how one root can carry two opposite meanings.")

module("Module 5 · Polysemy by shared root, and morphological richness","26–32 min")
beat("What it is.","Polysemy — when one root branches into distinct meanings across its forms, like a river that forks into two valleys.")
beat("Why we do it.","The root layer of Week 1 merged these branches into a single count; reopening the forms recovers the meanings it silently fused.")
beat("How it's done.","Trace the branches. By sense: ءمن flows into إيمان (faith) and into أمن (security) and أمين (trustworthy) — to believe, to be safe, to be relied upon. كتب flows into كِتاب (a book) and كُتِب (was decreed). And most strikingly, by moral valence: the root ك-ث-ر, simply “to be many,” yields كوثر — the river of blessed abundance given to the Prophet (108:1) — and تكاثر — the blameworthy rivalry of piling up wealth that distracts a person until the grave (102:1). One neutral root, “muchness,” pulled to opposite moral poles by its mould and its context — and 57:20 spells the negative pole out, listing تكاثر in wealth and children among the fleeting ornaments of worldly life. (We met this once before: ظلم, injustice, and ظلمات, darkness, are the same root wearing two faces.)")
beat("What we get.","A map of a root's semantic range — and a sense of how wide that range is.")
beat("Why it matters.","Counting a root as a single concept (Week 1) can quietly weld two ideas together; the forms pry them back apart. كوثر and تكاثر share every consonant, yet one is praise and the other is blame.")
figure("fig_polysemy.png","Figure 5. One root, divergent forms: ك-ث-ر splits by moral valence (كوثر blessed abundance vs تكاثر blameworthy rivalry); ءمن and كتب split by sense.",width=6.0)
figure("fig_form_richness.png","Figure 6. Morphological richness varies by root — ءمن is richest, خلق leanest; bars give the form counts.",width=4.8)
beat("In the data.","Richness ranges widely — ءمن is the richest at 27 forms, شهد and ذكر at 23, down to خلق at only 12 (78% of which is a single form). A rich root is one the text turns over in the hand many ways; a lean one it uses for a single, tight job.")
beat("Takeaway.","A single root can split by sense (faith vs security) or even by moral valence (blessed abundance vs blameworthy hoarding); the form, not the root, tells you which is meant.")
beat("Bridge.","Forms are a root's company on the inside. Now its company on the outside — its partners.")

module("Module 6 · Partners — a root's external company","32–38 min")
beat("What it is.","A partner is a root that appears in the same ayah as the target root more often than chance would predict — a habitual companion.")
beat("Why we do it.","There is an old line in linguistics: you shall know a word by the company it keeps. A concept's neighbours sketch its lived context as surely as a person's friends sketch theirs.")
beat("How it's done.","Read the app's significant-partners list. Crucially, it is already controlled for length and frequency — graded on a curve — so that a ubiquitous root like “God” or “say” does not show up as everyone's partner just by being everywhere. How that curve is set is Weeks 4–5; today we simply trust the cleaned list, the way you might trust a thermostat without yet knowing the wiring. The point of the control is fairness: without it, every root would look like a partner of “God” and “say,” and we would learn nothing.")
beat("What we get.","For ءمن, its specific partners: صلح (righteous deeds) far in front, then عمل (works), then رسل (messengers) and قلب (heart).")
beat("Why it matters.","Faith is almost never named alone in the Qur'an — it walks arm in arm with action. آمنوا وعملوا الصالحات, “those who believe and do righteous deeds,” is one of the most repeated phrases in the whole book, and the partner list surfaces that bond as a number — turning a phrase a reader feels into a measurement a reader can check.")
figure("fig_amn_partners.png","Figure 7. Partners of ءمن (length-controlled). Its strongest specific partner is صلح (righteous deeds); عمل (works) follows.",width=6.0)
beat("In the data.","ءمن↔صلح and ءمن↔عمل are the strongest specific partners (significance z = 14 and 11); and tellingly, كفر appears too — the antonym.")
beat("Takeaway.","A root's partners sketch its lived context — and for faith, that context is action.")
beat("Bridge.","But a partner list can fool you if you read it like a verdict.")

module("Module 7 · Reading partners honestly — antonyms as partners","38–43 min")
beat("What it is.","The discipline of reading a partner list without over-reading it — taking each partner as a lead to follow, not a conclusion to announce.")
beat("Why we do it.","Co-occurrence is association, and association is the most over-interpreted number in all of text analysis. Two roots sharing a verse means they were spoken together — not that one causes the other, points to the other, or agrees with the other.")
beat("How it's done.","Hold three cautions. First, significant is not the same as meaningful — two common roots can keep crossing paths with nothing between them. Second, the list is controlled, but it is still association, with no direction and no cause. Third, the counter-intuitive one: a partner can be an ANTONYM — opposites are defined together, like rivals who are always in the same room.")
beat("What we get.","Criteria for a partner claim that is safe to make — and a guard against the most common mistake, reading nearness as alignment.")
beat("Why it matters.","If you forget that antonyms collocate, you will badly misread the data: a root's closest company can include its very opposite, not because they agree but because the text forever sets them face to face. A careless reader might announce that “the Qur'an associates faith approvingly with disbelief” — when the two share verses only to be contrasted, the way “light” and “dark” share a sentence precisely because it is about their opposition.")
figure("fig_antonym_partners.png","Figure 8. Antonyms are partners too — ءمن with كفر, هدي with ضلل, ذكر with نسي; bars give the shared-ayah counts.",width=5.6)
beat("In the data.","ءمن and its antonym كفر share 126 ayahs — among ءمن's very closest company — because belief and disbelief are constantly set against each other. The same holds for هدي and ضلل (guidance and going-astray) and for ذكر and نسي (remembering and forgetting). And some roots are each other's strongest partner — a mutual bond: رحم↔غفر (mercy and forgiveness), سمع↔بصر (the All-Hearing and the All-Seeing), حكم↔عزز (the Mighty and the Wise). When two concepts name each other, the text is telling you they are spoken in one breath.")
beat("Takeaway.","A partner is a lead, not a verdict; even an opposite can be a close companion.")
beat("Bridge.","So every forms-or-partners claim, like every claim in this course, gets paired with a labeled reading.")

module("Module 8 · Fact vs. interpretation — the discipline & wrap","43–47 min")
beat("What it is.","The two-sentence discipline applied to a forms-or-partners claim, and the thread that runs through the whole term.")
beat("Why we do it.","Morphology and collocation are seductive — they hand you a vivid story, and the temptation is to present the story as a finding. The discipline keeps the computed fact and the human reading visibly apart.")
beat("How it's done.","Write one sentence of fact — a form share or a partner's significance — and then one separately labeled sentence of interpretation, beginning “I read this as …”.")
beat("What we get.","A reading a classmate can audit — they can point to which sentence the computer could check and which sentence is yours. That peer-test is the whole method in miniature: if a classmate cannot tell your fact from your reading, the two have blurred, and the reading has to be rewritten until they separate cleanly.")
beat("Why it matters.","The course is graded on keeping these two apart; and beyond the grade, it is the difference between analysis and projection. The grammar of a word can encode how a community holds a concept — but only if we don't smuggle our conclusions into the count.")
beat("In the data.","Fact: “61% of ء-م-ن's tokens are verb forms, and its strongest partner is صلح (significance z = 14).”  Interpretation: “I read faith here as an enacted commitment, bound to righteous works, rather than an abstract belief.”")
beat("Takeaway.","The root tells you WHAT, the pattern tells you HOW, the partners tell you WHO WITH — all computed, all needing a labeled reading.")
beat("Bridge.","Next week we turn the partners idea into a measure: which root shares the MOST ayahs with a target, and why the raw count will try to fool us.")

doc.save(OUT); print("saved",OUT,"| paragraphs:",len(doc.paragraphs))
